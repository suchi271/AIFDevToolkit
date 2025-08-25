"""
Assessment Report Generator for Azure Migration Project
Generates application assessment reports based on transcript and Q&A analysis using AI-driven content generation
"""

from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dataclasses import dataclass, field
import os
import re
import json
import openai
from dotenv import load_dotenv

from .StateBase import QuestionAnswer, AzureMigrateServer, TargetArchitecture, SubnetRecommendation, NSGRule, LoadBalancerConfig, LoadBalancingRule

# Load environment variables from .env file
load_dotenv()


@dataclass
class AssessmentReportData:
    """Data structure for Assessment Report"""
    
    application_name: str = "Application Name"
    environments: List[str] = field(default_factory=lambda: ["Production", "Development", "Pre-Production"])
    questions_answers: List[QuestionAnswer] = field(default_factory=list)
    security_considerations: List[Dict[str, str]] = field(default_factory=list)
    network_requirements: List[Dict[str, str]] = field(default_factory=list)
    identity_providers: List[Dict[str, str]] = field(default_factory=list)
    automation_details: List[Dict[str, str]] = field(default_factory=list)
    customer_impact: List[Dict[str, str]] = field(default_factory=list)
    operational_concerns: List[Dict[str, str]] = field(default_factory=list)
    observability: Dict[str, Any] = field(default_factory=dict)
    supporting_documents: List[Dict[str, str]] = field(default_factory=list)
    architecture_heatmap: List[Dict[str, str]] = field(default_factory=list)
    application_allocation: Dict[str, Any] = field(default_factory=dict)
    target_architecture: Any = None  # Can be TargetArchitecture object or dict
    additional_backlog: List[Dict[str, str]] = field(default_factory=list)
    rbac_information: List[Dict[str, str]] = field(default_factory=list)
    azure_tagging: List[Dict[str, str]] = field(default_factory=list)
    source_delivery_info: List[Dict[str, str]] = field(default_factory=list)
    target_delivery_info: List[Dict[str, str]] = field(default_factory=list)


class ApplicationAssessmentReportGenerator:
    """Generates comprehensive application assessment reports for Azure migration using AI-driven content generation."""
    
    def __init__(self, llm_client=None):
        self.template_path = None
        self.llm_client = llm_client
        
        # Load configuration from .env file (similar to MigrationPlanGenerator pattern)
        self.config = self._load_config()
        
        # Initialize AI client for content generation
        if self.llm_client is None:
            self.llm_client = self._initialize_ai_client()
        
    def _load_config(self) -> Dict[str, Any]:
        """Load configuration settings from .env file following MigrationPlanGenerator pattern."""
        return {
            # AI Settings - optimized for speed
            "ai_temperature": float(os.getenv("AI_TEMPERATURE", os.getenv("AZURE_OPENAI_TEMPERATURE", "0.3"))),  # Lower for speed
            "ai_max_tokens": int(os.getenv("AI_MAX_TOKENS", "500")),  # Reduced for speed
            "ai_timeout_seconds": int(os.getenv("AI_TIMEOUT_SECONDS", "30")),  # Reduced timeout
            "ai_model": os.getenv("OPENAI_MODEL", os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4")),
            "ai_api_version": os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-01"),
            
            # Speed optimizations
            "enable_full_ai_generation": os.getenv("ENABLE_FULL_AI_GENERATION", "false").lower() == "true",
            "ai_generation_mode": os.getenv("AI_GENERATION_MODE", "fast"),  # fast, balanced, comprehensive
            
            # Content Generation Preferences
            "content_style": os.getenv("CONTENT_STYLE", "professional"),
            "detail_level": os.getenv("DETAIL_LEVEL", "comprehensive"),
            "include_technical_details": os.getenv("INCLUDE_TECHNICAL_DETAILS", "true").lower() == "true",
            "include_cost_analysis": os.getenv("INCLUDE_COST_ANALYSIS", "true").lower() == "true",
            "include_risk_assessment": os.getenv("INCLUDE_RISK_ASSESSMENT", "true").lower() == "true",
            
            # Assessment Report Customization
            "default_project_name": os.getenv("DEFAULT_PROJECT_NAME", "Azure Migration Project"),
            "organization_name": os.getenv("ORGANIZATION_NAME", "Your Organization"),
            "include_architecture_analysis": os.getenv("INCLUDE_ARCHITECTURE_ANALYSIS", "true").lower() == "true",
            "generate_cost_estimates": os.getenv("GENERATE_COST_ESTIMATES", "true").lower() == "true",
        }
    
    def _initialize_ai_client(self):
        """Initialize AI client based on .env configuration following MigrationPlanGenerator pattern."""
        try:
            # Try Azure OpenAI first
            azure_api_key = os.getenv('AZURE_OPENAI_API_KEY')
            azure_endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
            azure_api_version = os.getenv('AZURE_OPENAI_API_VERSION', '2024-02-01')
            
            if azure_api_key and azure_endpoint:
                return openai.AzureOpenAI(
                    api_key=azure_api_key,
                    api_version=azure_api_version,
                    azure_endpoint=azure_endpoint
                )
            
            # Fall back to standard OpenAI
            openai_api_key = os.getenv('OPENAI_API_KEY')
            if openai_api_key:
                return openai.OpenAI(api_key=openai_api_key)
                
        except Exception as e:
            print(f"Warning: Could not initialize AI client: {e}")
            
        return None
        
    def _generate_ai_content(self, prompt: str, context_data: Dict[str, Any] = None) -> str:
        """
        Generate content using AI based on prompt and context data following MigrationPlanGenerator pattern.
        
        Args:
            prompt: The prompt for content generation
            context_data: Additional context data to include in the prompt
            
        Returns:
            Generated content as string
        """
        if not self.llm_client:
            return f"[AI Content Generation Unavailable - Please configure AI client]\n{prompt}"
        
        try:
            # Prepare context
            context_str = ""
            if context_data:
                context_str = f"\n\nContext Data:\n{json.dumps(context_data, indent=2, default=str)}"
            
            # Customize prompt based on configuration
            style_instruction = self._get_style_instruction()
            detail_instruction = self._get_detail_instruction()
            
            full_prompt = f"""
You are an expert Azure migration consultant and application assessment specialist creating professional assessment report content.

{style_instruction}

{detail_instruction}

{prompt}

Requirements:
- Generate {self.config['content_style']} content suitable for enterprise application assessment reports
- Use specific data from the context when available
- Include concrete details, insights, and actionable recommendations
- Structure content with clear sections and bullet points where appropriate
- Focus on practical, implementation-ready guidance for Azure migration
- Organization: {self.config['organization_name']}
- Base recommendations on actual conversation data and technical findings

{context_str}

Generate the content:
"""
            
            # Handle different LLM client types
            if hasattr(self.llm_client, 'chat') and hasattr(self.llm_client.chat, 'completions'):
                # OpenAI style client
                response = self.llm_client.chat.completions.create(
                    model=self.config['ai_model'],
                    messages=[
                        {"role": "system", "content": "You are an expert Azure migration consultant and application assessment specialist with deep knowledge of enterprise application architecture and cloud migration strategies."},
                        {"role": "user", "content": full_prompt}
                    ],
                    max_tokens=self.config['ai_max_tokens'],
                    temperature=self.config['ai_temperature']
                )
                return response.choices[0].message.content.strip()
            elif hasattr(self.llm_client, 'invoke'):
                # LangChain style client
                response = self.llm_client.invoke(full_prompt)
                if hasattr(response, 'content'):
                    return response.content.strip()
                else:
                    return str(response).strip()
            else:
                # Fallback for unknown client types
                return str(self.llm_client(full_prompt)).strip()
            
        except Exception as e:
            return f"[AI Content Generation Error: {str(e)}]\n\nFallback content for: {prompt}"
    
    def _get_style_instruction(self) -> str:
        """Get style instruction based on configuration following MigrationPlanGenerator pattern."""
        style_instructions = {
            "professional": "Write in a professional, business-appropriate tone suitable for enterprise stakeholders and technical teams.",
            "technical": "Focus on technical details and implementation specifics for IT professionals and architects.",
            "executive": "Write in a high-level, strategic tone suitable for executive leadership and decision makers."
        }
        return style_instructions.get(self.config['content_style'], style_instructions['professional'])
    
    def _get_detail_instruction(self) -> str:
        """Get detail level instruction based on configuration following MigrationPlanGenerator pattern."""
        detail_instructions = {
            "summary": "Provide concise, high-level information focusing on key assessment points only.",
            "standard": "Provide balanced detail with essential assessment information and supporting technical details.",
            "comprehensive": "Provide thorough, detailed analysis with comprehensive coverage of all assessment aspects."
        }
        return detail_instructions.get(self.config['detail_level'], detail_instructions['comprehensive'])
    
    def _llm_analyze(self, prompt: str, fallback_response: Any = None) -> Any:
        """Central LLM analysis method with fallback handling - Legacy method, use _generate_ai_content for new implementations."""
        if not self.llm_client:
            return fallback_response
            
        try:
            # Handle different LLM client types
            if hasattr(self.llm_client, 'chat') and hasattr(self.llm_client.chat, 'completions'):
                # OpenAI style client
                response = self.llm_client.chat.completions.create(
                    model=self.config['ai_model'],
                    messages=[
                        {"role": "system", "content": "You are an expert Azure migration consultant and application assessment specialist."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=self.config['ai_max_tokens'],
                    temperature=self.config['ai_temperature']
                )
                content = response.choices[0].message.content.strip()
            elif hasattr(self.llm_client, 'invoke'):
                # LangChain style client
                response = self.llm_client.invoke(prompt)
                if hasattr(response, 'content'):
                    content = response.content.strip()
                else:
                    content = str(response).strip()
            else:
                # Fallback for unknown client types
                content = str(self.llm_client(prompt)).strip()
            
            try:
                return json.loads(content)
            except json.JSONDecodeError:
                return content
        except Exception as e:
            print(f"LLM analysis failed: {e}")
            return fallback_response
    
    def _format_target_architecture_content(self, env_name: str, target_architecture: TargetArchitecture) -> str:
        """Format target architecture data for Word document inclusion."""
        content = f"**Network Traffic Analysis for {env_name} Environment**\n\n"
        
        # Generate comprehensive network analysis in the requested format
        content += self._generate_comprehensive_network_analysis(target_architecture)
        
        return content
    
    def _generate_comprehensive_network_analysis(self, target_architecture: TargetArchitecture) -> str:
        """Generate comprehensive network analysis using optimized LLM prompts for faster generation."""
        
        # Quick check for valid data
        network_connections = getattr(target_architecture, 'network_connections', []) or []
        if not network_connections:
            return self._generate_fallback_network_analysis(target_architecture)
        
        # Prepare minimal context data for faster processing
        context_data = self._prepare_minimal_network_context(target_architecture)
        
        # Use a simplified, faster prompt
        network_analysis_prompt = f"""Generate a concise Low Level Design - Network Traffic Analysis section for Azure migration.

Based on {len(network_connections)} network connections analyzed, create a professional section with:

1. **🏗️ Network Architecture Analysis** - List subnet recommendations, NSG rules, and load balancer configs
2. **💻 Compute Recommendations** - Azure service recommendations for discovered applications  
3. **📋 Network Diagram Description** - Brief architecture overview

Keep it concise and technical. Use the actual numbers: {context_data['summary_stats']}

Generate 200-300 words maximum focusing on actionable Azure recommendations."""

        try:
            # Generate with shorter timeout and fallback
            generated_content = self._generate_ai_content_fast(network_analysis_prompt, context_data)
            
            # If generation is successful and reasonable length, use it
            if generated_content and len(generated_content) > 100 and not generated_content.startswith("[AI Content Generation"):
                return generated_content
            else:
                # Quick fallback
                return self._generate_fallback_network_analysis(target_architecture)
                
        except Exception as e:
            print(f"LLM generation failed, using fallback: {e}")
            return self._generate_fallback_network_analysis(target_architecture)
    
    def _prepare_minimal_network_context(self, target_architecture: TargetArchitecture) -> Dict[str, Any]:
        """Prepare minimal context data for faster LLM processing."""
        
        network_connections = getattr(target_architecture, 'network_connections', []) or []
        subnets = getattr(target_architecture, 'subnet_recommendations', []) or []
        nsg_rules = getattr(target_architecture, 'nsg_rules', []) or []
        load_balancers = getattr(target_architecture, 'load_balancer_config', []) or []
        
        # Quick stats calculation
        unique_ports = set()
        unique_apps = set()
        
        for conn in network_connections[:20]:  # Limit to first 20 for speed
            if hasattr(conn, 'destination_port') and conn.destination_port:
                unique_ports.add(str(conn.destination_port))
            if hasattr(conn, 'source_application') and conn.source_application:
                unique_apps.add(conn.source_application)
            if hasattr(conn, 'destination_application') and conn.destination_application:
                unique_apps.add(conn.destination_application)
        
        summary_stats = f"{len(network_connections)} connections, {len(subnets)} subnets, {len(nsg_rules)} NSG rules, {len(load_balancers)} load balancers, {len(unique_ports)} unique ports, {len(unique_apps)} applications"
        
        return {
            "summary_stats": summary_stats,
            "connection_count": len(network_connections),
            "subnet_count": len(subnets),
            "nsg_count": len(nsg_rules),
            "lb_count": len(load_balancers),
            "unique_ports": list(unique_ports)[:10],  # First 10 ports
            "unique_apps": list(unique_apps)[:10]     # First 10 apps
        }
    
    def _generate_ai_content_fast(self, prompt: str, context_data: Dict[str, Any] = None) -> str:
        """Generate AI content with optimized settings for faster response."""
        if not self.llm_client:
            return f"[AI Content Generation Unavailable]\n{prompt}"
        
        try:
            # Simplified prompt for faster processing
            simple_prompt = f"""Create a technical Low Level Design section for Azure migration:

{prompt}

Context: {context_data.get('summary_stats', '')}

Be concise and technical."""
            
            response = self.llm_client.chat.completions.create(
                model=self.config['ai_model'],
                messages=[
                    {"role": "system", "content": "You are a concise Azure migration expert. Provide brief, technical responses."},
                    {"role": "user", "content": simple_prompt}
                ],
                max_tokens=800,  # Reduced from 2000
                temperature=0.3  # Lower temperature for faster, more deterministic responses
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            return f"[AI Content Generation Error: {str(e)}]"
        """Prepare comprehensive context data for LLM network analysis generation."""
        
        # Safely extract network connection details
        connections_data = []
        unique_applications = set()
        unique_ports = set()
        unique_source_ips = set()
        unique_dest_ips = set()
        
        # Handle network connections safely
        network_connections = getattr(target_architecture, 'network_connections', []) or []
        
        for conn in network_connections:
            conn_data = {
                "source_ip": getattr(conn, 'source_ip', 'Unknown'),
                "destination_ip": getattr(conn, 'destination_ip', 'Unknown'),
                "destination_port": getattr(conn, 'destination_port', 'Unknown'),
                "source_application": getattr(conn, 'source_application', ''),
                "destination_application": getattr(conn, 'destination_application', '')
            }
            connections_data.append(conn_data)
            
            # Collect unique values for analysis
            if conn_data["source_application"]:
                unique_applications.add(conn_data["source_application"])
            if conn_data["destination_application"]:
                unique_applications.add(conn_data["destination_application"])
            if conn_data["destination_port"] != 'Unknown':
                unique_ports.add(str(conn_data["destination_port"]))
            if conn_data["source_ip"] != 'Unknown':
                unique_source_ips.add(conn_data["source_ip"])
            if conn_data["destination_ip"] != 'Unknown':
                unique_dest_ips.add(conn_data["destination_ip"])
        
        # Safely prepare subnet recommendations data
        subnets_data = []
        subnet_recommendations = getattr(target_architecture, 'subnet_recommendations', []) or []
        for subnet in subnet_recommendations:
            subnets_data.append({
                "name": getattr(subnet, 'name', 'Unknown'),
                "purpose": getattr(subnet, 'purpose', 'Unknown'),
                "address_space": getattr(subnet, 'address_space', 'Auto-allocated'),
                "security_level": getattr(subnet, 'security_level', 'Standard')
            })
        
        # Safely prepare NSG rules data
        nsg_rules_data = []
        nsg_rules = getattr(target_architecture, 'nsg_rules', []) or []
        for rule in nsg_rules:
            nsg_rules_data.append({
                "name": getattr(rule, 'name', 'Unknown'),
                "direction": getattr(rule, 'direction', 'Inbound'),
                "priority": getattr(rule, 'priority', 1000),
                "access": getattr(rule, 'access', 'Allow'),
                "protocol": getattr(rule, 'protocol', 'TCP'),
                "source": getattr(rule, 'source', '*'),
                "destination": getattr(rule, 'destination', '*'),
                "destination_port": getattr(rule, 'destination_port', '*'),
                "description": getattr(rule, 'description', '')
            })
        
        # Safely prepare load balancer data
        load_balancers_data = []
        load_balancer_config = getattr(target_architecture, 'load_balancer_config', []) or []
        for lb in load_balancer_config:
            lb_rules = []
            for lb_rule in getattr(lb, 'rules', []):
                lb_rules.append({
                    "name": getattr(lb_rule, 'name', 'Unknown'),
                    "frontend_port": getattr(lb_rule, 'frontend_port', 80),
                    "backend_port": getattr(lb_rule, 'backend_port', 80),
                    "protocol": getattr(lb_rule, 'protocol', 'TCP')
                })
            
            load_balancers_data.append({
                "name": getattr(lb, 'name', 'Unknown'),
                "type": getattr(lb, 'type', 'Application'),
                "sku": getattr(lb, 'sku', 'Standard'),
                "rules": lb_rules
            })
        
        # Analyze application patterns
        application_patterns = self._analyze_application_patterns(unique_applications, unique_ports)
        
        # Calculate network metrics
        network_metrics = {
            "total_connections": len(network_connections),
            "unique_applications": len(unique_applications),
            "unique_ports": len(unique_ports),
            "unique_source_ips": len(unique_source_ips),
            "unique_destination_ips": len(unique_dest_ips),
            "subnet_count": len(subnet_recommendations),
            "nsg_rules_count": len(nsg_rules),
            "load_balancer_count": len(load_balancer_config)
        }
        
        # Port analysis
        port_analysis = self._analyze_port_usage(unique_ports)
        
        return {
            "network_metrics": network_metrics,
            "connections_sample": connections_data[:20],  # Sample for context
            "unique_applications": list(unique_applications),
            "unique_ports": list(unique_ports),
            "subnets": subnets_data,
            "nsg_rules": nsg_rules_data[:10],  # Sample of rules
            "load_balancers": load_balancers_data,
            "application_patterns": application_patterns,
            "port_analysis": port_analysis,
            "security_insights": self._generate_security_insights(target_architecture),
            "modernization_opportunities": self._identify_modernization_opportunities(unique_applications, unique_ports)
        }
    
    def _analyze_application_patterns(self, applications: set, ports: set) -> Dict[str, Any]:
        """Analyze application patterns for Azure service recommendations."""
        patterns = {
            "web_applications": [],
            "database_applications": [],
            "monitoring_tools": [],
            "file_services": [],
            "custom_applications": []
        }
        
        for app in applications:
            if not app or app in ['-', '']:
                continue
                
            app_lower = app.lower()
            if any(web in app_lower for web in ['apache', 'nginx', 'iis', 'tomcat', 'web', 'http']):
                patterns["web_applications"].append(app)
            elif any(db in app_lower for db in ['oracle', 'mysql', 'postgres', 'sql', 'database']):
                patterns["database_applications"].append(app)
            elif any(monitor in app_lower for monitor in ['agent', 'monitor', 'splunk', 'qualys']):
                patterns["monitoring_tools"].append(app)
            elif any(file_svc in app_lower for file_svc in ['file', 'storage', 'backup', 'ftp']):
                patterns["file_services"].append(app)
            else:
                patterns["custom_applications"].append(app)
        
        return patterns
    
    def _analyze_port_usage(self, ports: set) -> Dict[str, Any]:
        """Analyze port usage patterns for security and service recommendations."""
        port_categories = {
            "web_ports": [],
            "database_ports": [],
            "management_ports": [],
            "custom_ports": []
        }
        
        for port in ports:
            try:
                port_num = int(port)
                if port_num in [80, 443, 8080, 8443, 9000, 8000]:
                    port_categories["web_ports"].append(port_num)
                elif port_num in [1433, 3306, 5432, 1521, 27017]:
                    port_categories["database_ports"].append(port_num)
                elif port_num in [22, 3389, 5985, 5986]:
                    port_categories["management_ports"].append(port_num)
                else:
                    port_categories["custom_ports"].append(port_num)
            except ValueError:
                continue
        
        return {
            "categories": port_categories,
            "security_recommendations": self._get_port_security_recommendations(port_categories),
            "azure_service_mapping": self._map_ports_to_azure_services(port_categories)
        }
    
    def _get_port_security_recommendations(self, port_categories: Dict[str, list]) -> List[str]:
        """Generate security recommendations based on port usage."""
        recommendations = []
        
        if port_categories.get("web_ports"):
            recommendations.append("Implement Azure Application Gateway with WAF for web traffic protection")
            recommendations.append("Use Azure Front Door for global load balancing and DDoS protection")
        
        if port_categories.get("database_ports"):
            recommendations.append("Implement private endpoints for database connectivity")
            recommendations.append("Use Azure SQL Database with Advanced Threat Protection")
        
        if port_categories.get("management_ports"):
            recommendations.append("Implement Azure Bastion for secure RDP/SSH access")
            recommendations.append("Disable direct internet access for management ports")
        
        if port_categories.get("custom_ports"):
            recommendations.append("Review custom port usage for potential service consolidation")
            recommendations.append("Implement just-in-time (JIT) access for non-standard ports")
        
        return recommendations
    
    def _map_ports_to_azure_services(self, port_categories: Dict[str, list]) -> Dict[str, str]:
        """Map discovered ports to recommended Azure services."""
        service_mapping = {}
        
        if port_categories.get("web_ports"):
            service_mapping["Web Services"] = "Azure App Service, Azure Container Apps, Azure Application Gateway"
        
        if port_categories.get("database_ports"):
            service_mapping["Database Services"] = "Azure SQL Database, Azure Database for MySQL/PostgreSQL, Azure Cosmos DB"
        
        if port_categories.get("management_ports"):
            service_mapping["Management Services"] = "Azure Bastion, Azure Virtual Machines, Azure Arc"
        
        return service_mapping
    
    def _generate_security_insights(self, target_architecture: TargetArchitecture) -> List[str]:
        """Generate security insights based on network analysis."""
        insights = []
        
        # Safely analyze open ports for security risks
        open_ports = set()
        nsg_rules = getattr(target_architecture, 'nsg_rules', []) or []
        for rule in nsg_rules:
            if getattr(rule, 'access', '').lower() == 'allow':
                port = getattr(rule, 'destination_port', '')
                if port:
                    open_ports.add(str(port))
        
        if '3389' in open_ports or '22' in open_ports:
            insights.append("Management ports (RDP/SSH) detected - recommend Azure Bastion implementation")
        
        if len(open_ports) > 10:
            insights.append(f"Multiple open ports ({len(open_ports)}) detected - review for principle of least privilege")
        
        # Check for database ports
        db_ports = {'1433', '3306', '5432', '1521', '27017'}
        if db_ports.intersection(open_ports):
            insights.append("Database ports detected - implement private endpoints and network isolation")
        
        # Load balancer security
        load_balancer_config = getattr(target_architecture, 'load_balancer_config', []) or []
        if len(load_balancer_config) > 0:
            insights.append("Load balancers configured - ensure SSL termination and health monitoring")
        
        return insights
    
    def _identify_modernization_opportunities(self, applications: set, ports: set) -> List[str]:
        """Identify modernization opportunities based on discovered applications and ports."""
        opportunities = []
        
        # Container opportunities
        web_apps = [app for app in applications if any(web in app.lower() for web in ['apache', 'nginx', 'iis', 'tomcat']) if app]
        if web_apps:
            opportunities.append(f"Containerization opportunity for web applications: {', '.join(web_apps[:3])}")
        
        # Database modernization
        db_apps = [app for app in applications if any(db in app.lower() for db in ['mysql', 'postgres', 'sql']) if app]
        if db_apps:
            opportunities.append(f"Database modernization opportunity with Azure PaaS: {', '.join(db_apps[:3])}")
        
        # Serverless opportunities
        if '80' in ports or '443' in ports:
            opportunities.append("Azure Functions opportunity for lightweight HTTP-based services")
        
        # API Management
        if len([p for p in ports if p in ['80', '443', '8080', '9000']]) > 1:
            opportunities.append("Azure API Management opportunity for API consolidation and governance")
        
        return opportunities
    
    def _generate_fallback_network_analysis(self, target_architecture: TargetArchitecture) -> str:
        """Generate fallback network analysis when LLM is unavailable."""
        analysis = f"**Low Level Design - Network Traffic Analysis**\n\n"
        analysis += f"Based on analysis of {len(target_architecture.network_connections)} network connections, the following Azure architecture recommendations have been generated:\n\n"
        
        # 🏗️ Network Architecture Analysis
        analysis += "🏗️ **Network Architecture Analysis:**\n"
        
        # Subnet Recommendations
        subnet_count = len(target_architecture.subnet_recommendations)
        analysis += f"  • Subnet Recommendations: {subnet_count}\n"
        for subnet in target_architecture.subnet_recommendations:
            analysis += f"    - {subnet.name}: {subnet.purpose}\n"
        
        # NSG Rules
        nsg_count = len(target_architecture.nsg_rules)
        analysis += f"  • NSG Rules Generated: {nsg_count}\n"
        for rule in target_architecture.nsg_rules[:5]:  # Show first 5 rules
            analysis += f"    - {rule.name}: Port {rule.destination_port}\n"
        if nsg_count > 5:
            analysis += f"    - ... and {nsg_count - 5} more rules\n"
        
        # Load Balancer Recommendations
        lb_count = len(target_architecture.load_balancer_config)
        analysis += f"  • Load Balancer Recommendations: {lb_count}\n"
        for lb in target_architecture.load_balancer_config:
            analysis += f"    - {lb.name}: {lb.type} load balancer\n"
        
        analysis += "\n"
        
        # 💻 Compute Recommendations
        compute_recommendations = self._extract_compute_recommendations_from_connections(target_architecture.network_connections)
        analysis += f"💻 **Compute Recommendations: {len(compute_recommendations)}**\n"
        for rec in compute_recommendations:
            analysis += f"  • {rec['application']}: {rec['recommendation']}\n"
        
        analysis += "\n"
        
        # 📋 Network Diagram Description
        analysis += "📋 **Network Diagram Description:**\n"
        analysis += self._generate_network_diagram_description(target_architecture)
        
        return analysis
    
    def _extract_compute_recommendations_from_connections(self, network_connections: List) -> List[Dict[str, str]]:
        """Extract compute service recommendations from network connections."""
        recommendations = []
        applications_seen = set()
        
        # Extract unique applications from network connections
        for conn in network_connections:
            apps_to_analyze = []
            
            if hasattr(conn, 'source_application') and conn.source_application and conn.source_application.strip():
                apps_to_analyze.append(conn.source_application.strip())
            
            if hasattr(conn, 'destination_application') and conn.destination_application and conn.destination_application.strip():
                apps_to_analyze.append(conn.destination_application.strip())
            
            for app in apps_to_analyze:
                if app and app not in applications_seen and app != '-' and len(app) > 2:
                    applications_seen.add(app)
                    
                    # Generate Azure service recommendation based on application type
                    recommendation = self._get_azure_service_recommendation(app, conn)
                    recommendations.append({
                        'application': app,
                        'recommendation': recommendation
                    })
        
        # Add standard connectivity recommendations
        if len(network_connections) > 50:  # If substantial network traffic
            recommendations.extend([
                {
                    'application': 'Hybrid Connectivity',
                    'recommendation': 'Azure ExpressRoute or Site-to-Site VPN'
                },
                {
                    'application': 'Service-to-Service Communication', 
                    'recommendation': 'Azure Private Endpoints and Service Endpoints'
                }
            ])
        
        return recommendations[:12]  # Limit to 12 recommendations as shown in example
    
    def _get_azure_service_recommendation(self, application: str, connection) -> str:
        """Get Azure service recommendation based on application name and connection details."""
        app_lower = application.lower()
        
        # Database applications
        if any(db in app_lower for db in ['oracle', 'mysql', 'postgres', 'sql', 'database', 'db']):
            return 'Azure Database Service (SQL/MySQL/PostgreSQL)'
        
        # Web applications
        if any(web in app_lower for web in ['apache', 'nginx', 'iis', 'tomcat', 'web', 'http']):
            return 'Azure App Service or Azure Container Apps'
        
        # Monitoring and agents
        if any(monitor in app_lower for monitor in ['agent', 'monitor', 'splunk', 'omsagent', 'qualys']):
            return 'Azure Virtual Machine or Azure Container Apps'
        
        # File and storage applications
        if any(storage in app_lower for storage in ['file', 'storage', 'backup', 'ftp']):
            return 'Azure Files or Azure Blob Storage'
        
        # Default recommendation based on port
        if hasattr(connection, 'destination_port'):
            port = str(connection.destination_port)
            if port in ['80', '443', '8080', '9000']:
                return 'Azure App Service or Azure Container Apps'
            elif port in ['1433', '3306', '5432', '1521']:
                return 'Azure Database Service'
            elif port in ['22', '3389']:
                return 'Azure Virtual Machine with Bastion Host'
        
        return 'Azure Virtual Machine or Azure Container Apps'
        """Extract compute service recommendations from network connections."""
        recommendations = []
        applications_seen = set()
        
        # Extract unique applications from network connections
        for conn in network_connections:
            apps_to_analyze = []
            
            if hasattr(conn, 'source_application') and conn.source_application and conn.source_application.strip():
                apps_to_analyze.append(conn.source_application.strip())
            
            if hasattr(conn, 'destination_application') and conn.destination_application and conn.destination_application.strip():
                apps_to_analyze.append(conn.destination_application.strip())
            
            for app in apps_to_analyze:
                if app and app not in applications_seen and app != '-' and len(app) > 2:
                    applications_seen.add(app)
                    
                    # Generate Azure service recommendation based on application type
                    recommendation = self._get_azure_service_recommendation(app, conn)
                    recommendations.append({
                        'application': app,
                        'recommendation': recommendation
                    })
        
        # Add standard connectivity recommendations
        if len(network_connections) > 50:  # If substantial network traffic
            recommendations.extend([
                {
                    'application': 'Hybrid Connectivity',
                    'recommendation': 'Azure ExpressRoute or Site-to-Site VPN'
                },
                {
                    'application': 'Service-to-Service Communication', 
                    'recommendation': 'Azure Private Endpoints and Service Endpoints'
                }
            ])
        
        return recommendations[:12]  # Limit to 12 recommendations as shown in example
    
    def _get_azure_service_recommendation(self, application: str, connection) -> str:
        """Get Azure service recommendation based on application name and connection details."""
        app_lower = application.lower()
        
        # Database applications
        if any(db in app_lower for db in ['oracle', 'mysql', 'postgres', 'sql', 'database', 'db']):
            return 'Azure Database Service (SQL/MySQL/PostgreSQL)'
        
        # Web applications
        if any(web in app_lower for web in ['apache', 'nginx', 'iis', 'tomcat', 'web', 'http']):
            return 'Azure App Service or Azure Container Apps'
        
        # Monitoring and agents
        if any(monitor in app_lower for monitor in ['agent', 'monitor', 'splunk', 'omsagent', 'qualys']):
            return 'Azure Virtual Machine or Azure Container Apps'
        
        # File and storage applications
        if any(storage in app_lower for storage in ['file', 'storage', 'backup', 'ftp']):
            return 'Azure Files or Azure Blob Storage'
        
        # Default recommendation based on port
        if hasattr(connection, 'destination_port'):
            port = str(connection.destination_port)
            if port in ['80', '443', '8080', '9000']:
                return 'Azure App Service or Azure Container Apps'
            elif port in ['1433', '3306', '5432', '1521']:
                return 'Azure Database Service'
            elif port in ['22', '3389']:
                return 'Azure Virtual Machine with Bastion Host'
        
        return 'Azure Virtual Machine or Azure Container Apps'
    
    def _generate_network_diagram_description(self, target_architecture: TargetArchitecture) -> str:
        """Generate network diagram description."""
        description = "**Target Network Architecture Overview**\n\n"
        description += f"The proposed Azure network architecture is designed based on analysis of {len(target_architecture.network_connections)} identified network connections from the dependency analysis.\n\n"
        
        description += "**Key Components:**\n"
        description += "1. **Virtual Network (VNet)**: Primary network container with multiple subnets for security isolation\n"
        description += f"2. **Subnets**: {len(target_architecture.subnet_recommendations)} application-specific subnets based on traffic patterns\n"
        description += f"3. **Network Security Groups**: {len(target_architecture.nsg_rules)} rules based on discovered ports and protocols\n"
        description += f"4. **Load Balancers**: {len(target_architecture.load_balancer_config)} load balancing configurations for high availability\n"
        description += "5. **Private Endpoints**: Service endpoints for secure Azure service connectivity\n"
        description += "6. **Hybrid Connectivity**: ExpressRoute or VPN Gateway for on-premises integration\n\n"
        
        # Add specific network insights
        if target_architecture.network_connections:
            unique_ports = set()
            unique_ips = set()
            
            for conn in target_architecture.network_connections:
                if hasattr(conn, 'destination_port') and conn.destination_port:
                    unique_ports.add(str(conn.destination_port))
                if hasattr(conn, 'source_ip') and conn.source_ip:
                    unique_ips.add(conn.source_ip)
            
            description += "**Network Traffic Insights:**\n"
            description += f"• Discovered {len(unique_ports)} unique ports requiring firewall rules\n"
            description += f"• Identified {len(unique_ips)} unique IP addresses for subnet planning\n"
            description += f"• Analyzed {len(target_architecture.network_connections)} connection patterns for security and performance optimization\n"
        
        return description
    
    def _format_network_analysis_summary(self, assessment_data: AssessmentReportData) -> str:
        """Format network analysis summary for AI context."""
        if not hasattr(assessment_data, 'target_architecture') or not assessment_data.target_architecture:
            return "No network traffic analysis available."
        
        try:
            target_arch = assessment_data.target_architecture
            
            # Safely get network connections
            network_connections = getattr(target_arch, 'network_connections', []) or []
            summary = f"Network Connections Analyzed: {len(network_connections)}\n"
            
            # Unique ports discovered
            ports = set()
            applications = set()
            source_ips = set()
            dest_ips = set()
            
            for conn in network_connections:
                # Safely access connection attributes
                dest_port = getattr(conn, 'destination_port', None)
                if dest_port:
                    ports.add(str(dest_port))
                    
                src_app = getattr(conn, 'source_application', None)
                if src_app:
                    applications.add(src_app)
                    
                dest_app = getattr(conn, 'destination_application', None)
                if dest_app:
                    applications.add(dest_app)
                    
                src_ip = getattr(conn, 'source_ip', None)
                if src_ip:
                    source_ips.add(src_ip)
                    
                dest_ip = getattr(conn, 'destination_ip', None)
                if dest_ip:
                    dest_ips.add(dest_ip)
            
            summary += f"Unique Ports: {sorted(list(ports))}\n"
            summary += f"Applications Discovered: {sorted(list(applications))}\n"
            summary += f"Source IP Ranges: {len(source_ips)} unique IPs\n"
            summary += f"Destination IP Ranges: {len(dest_ips)} unique IPs\n"
            
            # Safely get subnet recommendations
            subnet_recs = getattr(target_arch, 'subnet_recommendations', []) or []
            summary += f"Subnets Recommended: {len(subnet_recs)}\n"
            
            return summary
            
        except Exception as e:
            print(f"Warning: Error formatting network analysis summary: {e}")
            return "Network traffic analysis summary unavailable due to processing error."
        summary += f"NSG Rules Generated: {len(target_arch.nsg_rules) if target_arch.nsg_rules else 0}"
        
        return summary
    
    def generate_assessment_report(
        self,
        questions_answers: List[QuestionAnswer],
        azure_migrate_data: Any,
        dependency_analysis: Any = None,
        project_name: str = "Azure Migration Project",
        template_path: Optional[str] = None,
        llm_client=None
    ) -> AssessmentReportData:
        """
        Generate comprehensive assessment report using transcript, Azure Migrate data, and dependency analysis.
        
        Args:
            questions_answers: List of question-answer pairs from transcript analysis
            azure_migrate_data: Azure Migrate report data
            dependency_analysis: Azure Migrate dependency analysis data
            project_name: Name of the migration project
            template_path: Optional path to template file
            llm_client: LLM client for intelligent analysis
            
        Returns:
            AssessmentReportData object containing structured assessment information
        """
        
        # Use provided LLM client or instance client
        if llm_client:
            self.llm_client = llm_client
        
        assessment_data = AssessmentReportData()
        
        # Store questions_answers for formatting methods
        assessment_data.questions_answers = questions_answers
        
        # Extract application name from Q&A or use project name
        assessment_data.application_name = self._extract_application_name(questions_answers, project_name)
        
        # Extract environment information for dynamic content generation
        assessment_data.environments = self._extract_environments(questions_answers)
        
        # Process Q&A data to populate assessment sections using comprehensive analysis
        assessment_data.security_considerations = self._extract_security_considerations(questions_answers)
        assessment_data.network_requirements = self._extract_network_requirements_enhanced(questions_answers, dependency_analysis)
        assessment_data.identity_providers = self._extract_identity_providers(questions_answers)
        assessment_data.automation_details = self._extract_automation_details(questions_answers)
        assessment_data.customer_impact = self._extract_customer_impact(questions_answers)
        assessment_data.operational_concerns = self._extract_operational_concerns(questions_answers)
        assessment_data.observability = self._extract_observability_info(questions_answers)
        
        # Process Azure Migrate data if available
        if azure_migrate_data:
            assessment_data.architecture_heatmap = self._generate_architecture_heatmap_enhanced(azure_migrate_data, questions_answers, dependency_analysis)
            assessment_data.application_allocation = self._generate_application_allocation(azure_migrate_data)
        
        # Generate target architecture recommendations based on network traffic analysis
        assessment_data.target_architecture = self._generate_target_architecture(questions_answers, azure_migrate_data, dependency_analysis)
        
        # Generate supporting documentation list
        assessment_data.supporting_documents = self._generate_supporting_documents()
        
        return assessment_data
    
    def _determine_migration_approach(self, questions_answers: List[QuestionAnswer]) -> Dict[str, str]:
        """Centrally determine the migration approach and justification to ensure consistency throughout the document."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return {
                "approach": "Replatform",
                "justification": "Balanced approach recommended for typical application modernization, providing cloud optimization benefits while maintaining reasonable migration complexity and timeline."
            }
        
        context_data = {
            "qa_content": qa_text[:3000],
            "total_qa_pairs": len(questions_answers)
        }
        
        migration_prompt = f"""Analyze this application assessment conversation and determine the single most appropriate migration approach.

Choose ONE approach from: Rehost, Replatform, or Refactor

ANALYSIS CRITERIA:
1. **Rehost (Lift-and-Shift)**: Minimal changes, fastest migration, maintain current architecture
   - Choose if: Time pressures, minimal disruption required, legacy systems, conservative approach needed
   
2. **Replatform (Lift-Tinker-and-Shift)**: Some optimization, moderate changes, leverage cloud services
   - Choose if: Moderate timeline, some cloud optimization desired, modern technology stack, balanced approach
   
3. **Refactor (Re-architect)**: Significant changes, cloud-native benefits, architectural modernization
   - Choose if: Long-term transformation goals, modern development practices, significant cloud optimization needed

Based on the conversation, provide:
1. The single recommended approach
2. Detailed justification explaining why this approach was chosen over the other two options

Return as JSON:
{{
    "approach": "Replatform|Rehost|Refactor",
    "justification": "Detailed explanation of why this approach is optimal, including what factors led to choosing it over the other two approaches. Include specific references to requirements, constraints, or goals mentioned in the conversation."
}}

Ensure the justification explains why the other approaches were not selected."""

        ai_response = self._generate_ai_content(migration_prompt, context_data)
        
        # Parse AI response
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                approach = parsed_response.get('approach', '')
                justification = parsed_response.get('justification', '')
                
                # Validate approach is one of the expected values
                if approach in ['Rehost', 'Replatform', 'Refactor'] and justification:
                    return {
                        "approach": approach,
                        "justification": justification
                    }
                    
        except json.JSONDecodeError:
            pass
        
        # Fallback to intelligent default based on technology analysis
        tech_stack = self._analyze_technology_stack(questions_answers)
        
        if tech_stack.get('containers') or tech_stack.get('cloud_ready'):
            return {
                "approach": "Replatform",
                "justification": "Replatform approach recommended based on existing modern technology stack and containerization readiness. This approach enables leveraging Azure managed services while maintaining existing application architecture, providing optimal balance between migration speed and cloud optimization benefits. Rehost was not selected due to missed opportunities for cloud optimization, while Refactor was deemed unnecessary given the existing modern architecture."
            }
        elif any('legacy' in qa.answer.lower() or 'old' in qa.answer.lower() for qa in questions_answers if qa.answer and qa.answer != "Not addressed in transcript"):
            return {
                "approach": "Rehost",
                "justification": "Rehost (Lift-and-Shift) approach recommended due to legacy technology constraints and need for minimal disruption during migration. This approach prioritizes speed and risk mitigation over optimization. Replatform was not selected due to technology constraints, while Refactor would require excessive time and resources for legacy application modernization."
            }
        else:
            return {
                "approach": "Replatform",
                "justification": "Replatform approach recommended as the optimal balance between migration speed and cloud optimization benefits. This approach enables adoption of Azure managed services while preserving core application architecture, reducing operational overhead and improving scalability. Rehost was not selected to avoid missing cloud optimization opportunities, while Refactor would extend timeline without proportional benefits for the current application architecture."
            }
    
    def export_to_word(
        self,
        assessment_data: AssessmentReportData,
        output_path: str,
        template_path: Optional[str] = None
    ) -> bool:
        """
        Export assessment report data to Word document.
        
        Args:
            assessment_data: Assessment report data
            output_path: Path to save the Word document
            template_path: Ignored - template is now embedded
            
        Returns:
            True if successful, False otherwise
        """
        
        try:
            # Create document from embedded template structure
            doc = self._create_embedded_template(assessment_data)
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error exporting assessment report: {e}")
            return False
    
    def _extract_application_name(self, questions_answers: List[QuestionAnswer], project_name: str) -> str:
        """Extract application name from Q&A data."""
        
        # Look for questions about application name, project name, or system name
        name_keywords = ['application name', 'app name', 'system name', 'project name', 'service name', 'what is the name']
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in name_keywords):
                    # Clean up the answer and extract just the name
                    app_name = qa.answer.strip()
                    
                    # Try to extract just the application name from common patterns
                    
                    # Pattern 1: "The application name mentioned in the conversation is "Name""
                    match = re.search(r'(?:application name|name).*?is\s*["\']?([^"\'.\n,]+)["\']?', app_name, re.IGNORECASE)
                    if match:
                        extracted_name = match.group(1).strip()
                        # Further clean the extracted name
                        clean_name = re.sub(r'[^\w\s-]', '', extracted_name).strip()
                        if len(clean_name) > 1 and len(clean_name) < 50:
                            return clean_name
                    
                    # Pattern 2: Look for quoted names
                    match = re.search(r'["\']([^"\']+)["\']', app_name)
                    if match:
                        extracted_name = match.group(1).strip()
                        clean_name = re.sub(r'[^\w\s-]', '', extracted_name).strip()
                        if len(clean_name) > 1 and len(clean_name) < 50:
                            return clean_name
                    
                    # Pattern 3: Look for capitalized words that could be application names
                    words = app_name.split()
                    for word in words:
                        # Skip common filler words
                        if word.lower() not in ['the', 'application', 'name', 'mentioned', 'in', 'conversation', 'is', 'called', 'system', 'project', 'service', 'transcript', 'discussion', 'meeting']:
                            clean_word = re.sub(r'[^\w-]', '', word)
                            # Look for words that start with capital letter or are all caps (likely names)
                            if len(clean_word) > 1 and len(clean_word) < 50 and (clean_word[0].isupper() or clean_word.isupper()):
                                return clean_word
                    
                    # Fallback: use the answer if it's reasonable length
                    if len(app_name) > 3 and len(app_name) < 100:
                        return app_name
        
        return project_name
    
    def _extract_environments(self, questions_answers: List[QuestionAnswer]) -> List[str]:
        """Extract environment information from Q&A data."""
        
        environments = []
        environment_keywords = [
            'environment', 'environments', 'env', 'prod', 'production', 'dev', 'development', 
            'test', 'testing', 'staging', 'uat', 'user acceptance', 'pre-prod', 'pre-production',
            'qa', 'quality assurance', 'demo', 'sandbox'
        ]
        
        # Common environment mappings
        env_mappings = {
            'prod': 'Production',
            'production': 'Production',
            'dev': 'Development', 
            'development': 'Development',
            'test': 'Testing',
            'testing': 'Testing',
            'qa': 'QA',
            'quality assurance': 'QA',
            'staging': 'Staging',
            'uat': 'UAT',
            'user acceptance': 'UAT',
            'pre-prod': 'Pre-Production',
            'pre-production': 'Pre-Production',
            'demo': 'Demo',
            'sandbox': 'Sandbox'
        }
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                answer_lower = qa.answer.lower()
                
                # Check if question is about environments
                if any(keyword in question_lower for keyword in environment_keywords):
                    # Extract environment names from the answer
                    for env_name, standard_name in env_mappings.items():
                        if env_name in answer_lower and standard_name not in environments:
                            environments.append(standard_name)
                
                # Also check answers for environment mentions
                elif any(keyword in answer_lower for keyword in environment_keywords):
                    for env_name, standard_name in env_mappings.items():
                        if env_name in answer_lower and standard_name not in environments:
                            environments.append(standard_name)
        
        # If no environments found, use default
        if not environments:
            environments = ["Production", "Development", "Pre-Production"]
        
        # Ensure Production is always first if present
        if "Production" in environments:
            environments.remove("Production")
            environments.insert(0, "Production")
        
        return environments
    
    def _generate_dynamic_toc(self, environments: List[str]) -> str:
        """Generate table of contents with dynamic environment sections and proper page alignment."""
        
        # Base static sections with proper alignment using dots
        toc_lines = [
            "Introduction" + "." * 70 + "3",
            "1\tApplication Overview" + "." * 60 + "4",
            "1.1\tKey Business Drivers" + "." * 55 + "4",
            "1.2\tKey Contacts" + "." * 65 + "4",
            "1.3\tMigration Strategy" + "." * 60 + "4",
            "1.3.1\tMigration Pattern and Complexity" + "." * 40 + "4",
            "1.3.2\tTechnology Selection" + "." * 50 + "5",
            "1.3.3\tIndicative Azure Cost" + "." * 50 + "5",
            "1.4\tDatabase Information" + "." * 55 + "5",
            "1.5\tMacro Dependencies" + "." * 55 + "6",
            "1.6\tSecurity Considerations" + "." * 50 + "6",
            "1.7\tResiliency Configuration" + "." * 50 + "6",
            "1.8\tNetwork Access Requirements" + "." * 45 + "7",
            "1.9\tIdentity Providers" + "." * 55 + "7",
            "1.10\tAutomation" + "." * 65 + "7",
            "1.11\tCustomer Impact" + "." * 60 + "8",
            "1.12\tOperational Concerns" + "." * 50 + "8",
            "1.13\tMigration Acceptance Tests" + "." * 45 + "8",
            "1.14\tObservability" + "." * 60 + "9",
            "2\tSupporting Documents" + "." * 55 + "10",
            "3\tCurrent Logical Architecture" + "." * 45 + "11"
        ]
        
        page_num = 11
        
        # Add dynamic environment sections for current architecture
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"3.{i+1}\t{env} Logical Architecture"))
            toc_lines.append(f"3.{i+1}\t{env} Logical Architecture" + dots + f"{page_num}")
        
        # Application Network Flow section
        page_num += 1
        toc_lines.append(f"4\tApplication Network Flow" + "." * 50 + f"{page_num}")
        
        # Add dynamic environment sections for network flow
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"4.{i+1}\t{env} Application Network Flow"))
            toc_lines.append(f"4.{i+1}\t{env} Application Network Flow" + dots + f"{page_num}")
        
        # Proposed Architecture section
        page_num += 1
        toc_lines.append(f"5\tProposed Architecture in Azure" + "." * 40 + f"{page_num}")
        
        # Add dynamic environment sections for proposed architecture
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"5.{i+1}\t{env} Proposed Architecture"))
            toc_lines.append(f"5.{i+1}\t{env} Proposed Architecture" + dots + f"{page_num}")
        
        # Add Low Level Design section
        page_num += 1
        toc_lines.append(f"5.4\tLow Level Design - Network Traffic Analysis" + "." * 25 + f"{page_num}")
        
        # Continue with static sections
        page_num += 1
        toc_lines.extend([
            f"6\tArchitecture Heatmap" + "." * 55 + f"{page_num}",
            f"7\tDecision Matrix" + "." * 60 + f"{page_num + 1}",
            f"8\tApplication Allocation and Scheduling" + "." * 35 + f"{page_num + 2}",
            f"9\tAppendix" + "." * 70 + f"{page_num + 3}",
            f"9.1\tAdditional Backlog Items" + "." * 45 + f"{page_num + 3}",
            f"9.2\tApplication and Infrastructure RBAC Information" + "." * 20 + f"{page_num + 3}"
        ])
        
        page_num += 3
        
        # Add dynamic environment sections for RBAC
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"9.2.{i+1}\t{env} Application and Infrastructure RBAC"))
            toc_lines.append(f"9.2.{i+1}\t{env} Application and Infrastructure RBAC" + dots + f"{page_num}")
        
        # Azure Services RBAC
        page_num += 1
        toc_lines.append(f"9.3\tAzure Services RBAC Information" + "." * 35 + f"{page_num}")
        
        # Add dynamic environment sections for Azure Services RBAC
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"9.3.{i+1}\t{env} Azure Services RBAC"))
            toc_lines.append(f"9.3.{i+1}\t{env} Azure Services RBAC" + dots + f"{page_num}")
        
        # Azure Tagging
        page_num += 1
        toc_lines.append(f"9.4\tAzure Tagging" + "." * 60 + f"{page_num}")
        
        # Add dynamic environment sections for Azure Tagging
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"9.4.{i+1}\t{env} Azure Tagging"))
            toc_lines.append(f"9.4.{i+1}\t{env} Azure Tagging" + dots + f"{page_num}")
        
        # Source Migration Delivery
        page_num += 1
        toc_lines.append(f"9.5\tSource Migration Delivery Information" + "." * 30 + f"{page_num}")
        
        # Add dynamic environment sections for Source Migration
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"9.5.{i+1}\t{env} Source Delivery Information"))
            toc_lines.append(f"9.5.{i+1}\t{env} Source Delivery Information" + dots + f"{page_num}")
        
        # Target Migration Delivery
        page_num += 1
        toc_lines.append(f"9.6\tTarget Migration Delivery Information" + "." * 30 + f"{page_num}")
        
        # Add dynamic environment sections for Target Migration
        for i, env in enumerate(environments):
            page_num += 1
            dots = "." * max(5, 70 - len(f"9.6.{i+1}\t{env} Target Delivery Information"))
            toc_lines.append(f"9.6.{i+1}\t{env} Target Delivery Information" + dots + f"{page_num}")
        
        return "\n".join(toc_lines)
    
    def _generate_source_delivery_requirements(self, env_name: str, assessment_data: AssessmentReportData) -> Dict[str, str]:
        """Generate intelligent source migration delivery requirements using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return self._generate_default_source_requirements(env_name, assessment_data)
        
        context_data = {
            "application_name": assessment_data.application_name,
            "environment": env_name,
            "qa_content": qa_text[:2500],
            "total_qa_pairs": len(assessment_data.questions_answers)
        }
        
        source_prompt = f"""Analyze this application assessment and generate specific source migration delivery requirements for the {env_name} environment.

Based on the conversation, generate detailed source delivery requirements covering:

1. **Server Specifications**: Current infrastructure details mentioned
2. **Authentication Systems**: Access control and authentication mechanisms discussed  
3. **Backup and Recovery**: Current backup strategies and requirements
4. **Network Configuration**: Connectivity and integration requirements
5. **Application Deployment**: Current deployment processes and binary management
6. **Configuration Management**: Settings and configuration file management
7. **Data Migration**: Database and data migration requirements

Provide specific, actionable requirements based on what was discussed in the assessment.

Return as JSON:
{{
    "source_requirements": {{
        "Server Specifications": "specific current server details and infrastructure requirements",
        "Authentication Systems": "current authentication and access control mechanisms",
        "Backup and Recovery": "current backup strategies and recovery requirements",
        "Network Configuration": "current network connectivity and integration needs",
        "Application Deployment": "current deployment processes and binary management",
        "Configuration Management": "current configuration and settings management",
        "Data Migration Requirements": "database and data migration specifications"
    }}
}}

Base all requirements on actual information discussed in the conversation. If specific details aren't available, provide logical requirements based on the application context."""

        ai_response = self._generate_ai_content(source_prompt, context_data)
        
        # Parse AI response
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                source_reqs = parsed_response.get('source_requirements', {})
                
                if source_reqs and len(source_reqs) > 0:
                    return source_reqs
                    
        except json.JSONDecodeError:
            pass
        
        # Fallback to context-based requirements if AI parsing fails
        return self._generate_default_source_requirements(env_name, assessment_data)
    
    def _generate_default_source_requirements(self, env_name: str, assessment_data: AssessmentReportData) -> Dict[str, str]:
        """Generate default source requirements based on application context."""
        
        # Analyze technology stack for better requirements
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        
        requirements = {}
        
        # Server specifications based on technology stack
        if tech_stack.get('containers'):
            requirements['Server Specifications'] = f"Container runtime environment documentation for {assessment_data.application_name}, including Docker/Kubernetes configurations and resource requirements"
        elif tech_stack.get('cloud_ready'):
            requirements['Server Specifications'] = f"Current application server specifications for {assessment_data.application_name}, including OS version, runtime dependencies, and resource utilization"
        else:
            requirements['Server Specifications'] = f"Legacy server infrastructure documentation for {assessment_data.application_name}, including OS, middleware, and hardware specifications"
        
        # Authentication based on detected systems
        auth_info = self._extract_authentication_info(assessment_data.questions_answers)
        if auth_info and 'active directory' in auth_info.lower():
            requirements['Authentication Systems'] = "Active Directory integration documentation, including domain configuration, service accounts, and authentication flows"
        else:
            requirements['Authentication Systems'] = f"Current authentication mechanism documentation for {assessment_data.application_name}, including user management and access control"
        
        # Environment-specific backup requirements
        if env_name.lower() == 'production':
            requirements['Backup and Recovery'] = "Production backup strategy documentation, including RTO/RPO requirements, backup schedules, and recovery procedures"
        else:
            requirements['Backup and Recovery'] = f"Backup and recovery procedures for {env_name} environment, including data protection and restoration processes"
        
        # Network configuration based on architecture
        requirements['Network Configuration'] = f"Current network architecture for {assessment_data.application_name}, including firewall rules, load balancer configuration, and integration endpoints"
        
        # Deployment based on technology stack
        if tech_stack.get('containers'):
            requirements['Application Deployment'] = "Container deployment documentation, including image repositories, CI/CD pipelines, and orchestration configurations"
        else:
            requirements['Application Deployment'] = f"Current deployment procedures for {assessment_data.application_name}, including build processes, deployment scripts, and environment configuration"
        
        # Configuration management
        requirements['Configuration Management'] = f"Application configuration documentation for {env_name}, including environment variables, connection strings, and feature flags"
        
        # Data migration based on detected databases
        db_technologies = tech_stack.get('databases', [])
        if db_technologies:
            requirements['Data Migration Requirements'] = f"Database migration specifications for {', '.join(db_technologies[:2])}, including schema, data volume, and migration strategy"
        else:
            requirements['Data Migration Requirements'] = f"Data migration requirements for {assessment_data.application_name}, including database schema and data transfer specifications"
        
        return requirements
    
    def _extract_authentication_info(self, questions_answers: List[QuestionAnswer]) -> str:
        """Extract authentication information using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return None
        
        auth_prompt = f"""Analyze this application assessment and identify authentication mechanisms mentioned.

TRANSCRIPT DATA:
{qa_text[:1500]}

Based on the conversation, identify authentication systems mentioned such as:
- Active Directory
- LDAP
- OAuth
- SAML/SSO
- Database authentication
- Custom authentication

Return just the name of the primary authentication mechanism mentioned, or "None" if not specified."""

        # Get AI analysis
        result = self._llm_analyze(auth_prompt, None)
        
        if isinstance(result, str) and result.strip() and result.lower() != "none":
            return result.strip()
        
        return None
    
    def _generate_target_delivery_requirements(self, env_name: str, assessment_data: AssessmentReportData) -> Dict[str, str]:
        """Generate intelligent target migration delivery requirements using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return self._generate_default_target_requirements(env_name, assessment_data)
        
        context_data = {
            "application_name": assessment_data.application_name,
            "environment": env_name,
            "qa_content": qa_text[:2500],
            "total_qa_pairs": len(assessment_data.questions_answers)
        }
        
        target_prompt = f"""Analyze this application assessment and generate specific Azure target delivery requirements for the {env_name} environment.

Based on the conversation, generate detailed Azure target delivery requirements covering:

1. **Azure Compute Services**: Specific compute services based on application architecture
2. **Azure Database Services**: Database services matching current technology stack
3. **Networking and Security**: Azure networking and security configuration
4. **Monitoring and Management**: Azure monitoring and management services needed
5. **Backup and Recovery**: Azure backup and disaster recovery configuration
6. **Cost Management**: Azure cost optimization strategies and controls
7. **DevOps Integration**: Azure DevOps and deployment automation requirements

Provide specific, actionable Azure service recommendations based on what was discussed.

Return as JSON:
{{
    "target_requirements": {{
        "Azure Compute Services": "specific Azure compute services and configurations",
        "Azure Database Services": "appropriate Azure database services based on current technology",
        "Networking and Security": "Azure networking and security service requirements",
        "Monitoring and Management": "Azure Monitor and management service configuration", 
        "Backup and Recovery": "Azure backup and disaster recovery service setup",
        "Cost Management": "Azure cost management and optimization strategies",
        "DevOps Integration": "Azure DevOps and CI/CD automation requirements"
    }}
}}

Map actual technologies and requirements mentioned to appropriate Azure services with specific configurations."""

        ai_response = self._generate_ai_content(target_prompt, context_data)
        
        # Parse AI response
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                target_reqs = parsed_response.get('target_requirements', {})
                
                if target_reqs and len(target_reqs) > 0:
                    return target_reqs
                    
        except json.JSONDecodeError:
            pass
        
        # Fallback to context-based requirements if AI parsing fails
        return self._generate_default_target_requirements(env_name, assessment_data)
    
    def _generate_default_target_requirements(self, env_name: str, assessment_data: AssessmentReportData) -> Dict[str, str]:
        """Generate default target requirements based on application context and technology stack."""
        
        # Analyze technology stack for better Azure service mapping
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        
        requirements = {}
        
        # Azure Compute Services based on technology stack
        if tech_stack.get('containers'):
            requirements['Azure Compute Services'] = f"Azure Kubernetes Service (AKS) for {assessment_data.application_name}, including node pools, auto-scaling, and container orchestration"
        elif tech_stack.get('cloud_ready'):
            requirements['Azure Compute Services'] = f"Azure App Service Premium tier for {assessment_data.application_name}, including auto-scaling, deployment slots, and health monitoring"
        else:
            requirements['Azure Compute Services'] = f"Azure Virtual Machines for {assessment_data.application_name}, including VM sizing, availability sets, and load balancing configuration"
        
        # Azure Database Services based on detected technologies
        db_technologies = tech_stack.get('databases', [])
        if 'PostgreSQL' in str(db_technologies):
            requirements['Azure Database Services'] = "Azure Database for PostgreSQL with appropriate compute tier, backup retention, and high availability configuration"
        elif 'MySQL' in str(db_technologies):
            requirements['Azure Database Services'] = "Azure Database for MySQL with suitable performance tier, automated backups, and monitoring"
        elif any('SQL' in str(db) for db in db_technologies):
            requirements['Azure Database Services'] = "Azure SQL Database with appropriate service tier, elastic pools, and intelligent performance tuning"
        else:
            requirements['Azure Database Services'] = f"Azure database service selection based on {assessment_data.application_name} data requirements and performance needs"
        
        # Networking and Security - Environment specific
        if env_name.lower() == 'production':
            requirements['Networking and Security'] = "Azure Application Gateway, Network Security Groups, Azure Firewall, and private endpoints for production security and traffic management"
        else:
            requirements['Networking and Security'] = f"Network Security Groups, Azure Load Balancer, and secure network configuration for {env_name} environment"
        
        # Monitoring and Management based on technology stack
        if tech_stack.get('containers'):
            requirements['Monitoring and Management'] = "Azure Monitor for containers, Application Insights, Log Analytics workspace, and Kubernetes monitoring solutions"
        else:
            requirements['Monitoring and Management'] = f"Azure Monitor, Application Insights for {assessment_data.application_name}, Log Analytics, and automated alerting configuration"
        
        # Backup and Recovery - Environment specific
        if env_name.lower() == 'production':
            requirements['Backup and Recovery'] = "Azure Backup with policy-based backup, Azure Site Recovery for disaster recovery, and geo-redundant storage for critical data protection"
        else:
            requirements['Backup and Recovery'] = f"Azure Backup configuration for {env_name}, automated backup policies, and point-in-time recovery capabilities"
        
        # Cost Management - Environment specific
        if env_name.lower() == 'production':
            requirements['Cost Management'] = "Azure Cost Management with budget alerts, Reserved Instances for predictable workloads, and cost optimization recommendations"
        elif env_name.lower() == 'development':
            requirements['Cost Management'] = "Development cost controls including auto-shutdown policies, dev/test pricing, and resource lifecycle management"
        else:
            requirements['Cost Management'] = f"Cost monitoring and budget management for {env_name}, appropriate SKU selection, and resource optimization"
        
        # DevOps Integration based on technology and deployment patterns
        if tech_stack.get('containers'):
            requirements['DevOps Integration'] = "Azure DevOps with container-based CI/CD pipelines, Azure Container Registry, and Infrastructure as Code deployment"
        else:
            requirements['DevOps Integration'] = f"Azure DevOps pipelines for {assessment_data.application_name}, source control integration, and automated deployment processes"
        
        return requirements
    
    def _add_formatted_paragraph(self, doc, content: str):
        """Add a paragraph with proper formatting, converting markdown bold to Word bold."""
        if not content:
            return
        
        # Split content by lines to handle each line separately
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                doc.add_paragraph("")  # Add empty paragraph for spacing
                continue
                
            para = doc.add_paragraph()
            
            # Split line by ** markers to identify bold sections
            parts = line.split('**')
            
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Even index = regular text
                    if part:
                        para.add_run(part)
                else:
                    # Odd index = bold text
                    if part:
                        run = para.add_run(part)
                        run.bold = True
    
    def _add_decision_matrix_table(self, doc, assessment_data: AssessmentReportData):
        """Add decision matrix as a proper Word table with intelligent analysis."""
        
        # Add the heading and introduction
        para = doc.add_paragraph()
        title_run = para.add_run("Migration Decision Matrix")
        title_run.bold = True
        
        doc.add_paragraph("The following matrix outlines the key decisions made during the assessment:")
        
        # Create the table with headers
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set up headers
        header_cells = table.rows[0].cells
        headers = ['Decision Area', 'Options Considered', 'Selected Approach', 'Rationale']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header text bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Generate intelligent decision matrix based on Q&A analysis
        decisions = self._generate_decision_matrix(assessment_data)
        
        # Add each decision as a table row
        for decision in decisions:
            row_cells = table.add_row().cells
            row_cells[0].text = decision['area']
            row_cells[1].text = decision['options']
            row_cells[2].text = decision['selected']
            row_cells[3].text = decision['rationale']
        
        # Add key decisions rationale section
        doc.add_paragraph("")  # Add spacing
        
        rationale_para = doc.add_paragraph()
        rationale_title = rationale_para.add_run("Key Decisions Rationale:")
        rationale_title.bold = True
        
        # Generate intelligent rationale points based on decisions
        rationale_points = self._generate_decision_rationale(decisions, assessment_data)
        
        for point in rationale_points:
            bullet_para = doc.add_paragraph(f"• {point}")
            bullet_para.style = 'List Bullet'

    def _generate_decision_matrix(self, assessment_data: AssessmentReportData) -> List[Dict[str, str]]:
        """Generate intelligent decision matrix based on comprehensive AI analysis with consistent migration approach."""
        
        # Get centralized migration approach first
        migration_decision = self._determine_migration_approach(assessment_data.questions_answers)
        
        # Prepare comprehensive context for AI analysis
        decision_context = self._prepare_decision_analysis_context(assessment_data.questions_answers)
        
        context_data = {
            "application_name": assessment_data.application_name,
            "transcript_data": decision_context["qa_content"],
            "technology_indicators": decision_context["tech_indicators"],
            "architecture_context": decision_context["architecture_context"],
            "business_context": decision_context["business_context"],
            "security_context": decision_context["security_context"],
            "performance_context": decision_context["performance_context"],
            "total_qa_pairs": len(assessment_data.questions_answers),
            "answered_qa_pairs": len([qa for qa in assessment_data.questions_answers if qa.is_answered]),
            "chosen_migration_approach": migration_decision["approach"],
            "migration_justification": migration_decision["justification"]
        }
        
        prompt = f"""
Analyze this comprehensive application assessment and generate an intelligent migration decision matrix using the pre-determined migration approach for consistency.

MIGRATION APPROACH ALREADY DETERMINED: {migration_decision["approach"]}
JUSTIFICATION: {migration_decision["justification"]}

Generate decisions for these key areas, ensuring consistency with the chosen migration approach:

1. **Migration Strategy** (Use Pre-Determined)
   - Use the already determined approach: {migration_decision["approach"]}
   - Include the provided justification explaining why this approach was chosen

2. **Compute Platform Selection**
   - Analyze application architecture patterns discussed  
   - Evaluate containerization readiness mentioned
   - Assess scalability requirements identified
   - Select Azure compute service aligned with {migration_decision["approach"]} approach

3. **Database Platform Strategy**
   - Identify current database technologies mentioned
   - Analyze data complexity and volume discussed
   - Evaluate migration complexity factors
   - Recommend Azure database service consistent with {migration_decision["approach"]} strategy

4. **Network Security Approach**
   - Analyze security requirements mentioned
   - Evaluate compliance needs discussed  
   - Assess external connectivity requirements
   - Determine network security strategy appropriate for {migration_decision["approach"]}

5. **Identity Management Strategy**
   - Identify current authentication systems mentioned
   - Analyze user management requirements
   - Evaluate integration complexity
   - Recommend Azure identity approach aligned with {migration_decision["approach"]} migration

For each decision area, provide:
- Multiple options that were realistically considered
- The selected approach based on assessment findings and migration strategy consistency
- Detailed rationale referencing specific conversation content
- Alignment with the chosen {migration_decision["approach"]} migration approach

Return as JSON:
{{
    "decisions": [
        {{
            "area": "Migration Strategy",
            "options": "Rehost vs Replatform vs Refactor",
            "selected": "{migration_decision["approach"]}",
            "rationale": "{migration_decision["justification"]}"
        }},
        {{
            "area": "Compute Platform",
            "options": "platform options considered based on architecture",
            "selected": "optimal Azure compute service for this application",
            "rationale": "platform-specific reasoning based on application characteristics and {migration_decision["approach"]} approach"
        }},
        {{
            "area": "Database Platform",
            "options": "database options based on current technology",
            "selected": "appropriate Azure database service",
            "rationale": "database-specific reasoning based on technology stack and {migration_decision["approach"]} migration strategy"
        }},
        {{
            "area": "Network Security",
            "options": "security approaches considered",
            "selected": "security strategy based on requirements",
            "rationale": "security-specific reasoning based on compliance needs and {migration_decision["approach"]} approach"
        }},
        {{
            "area": "Identity Management",
            "options": "identity options evaluated",
            "selected": "identity solution based on current systems",
            "rationale": "identity-specific reasoning based on authentication requirements and {migration_decision["approach"]} migration"
        }}
    ]
}}

Base all decisions on actual application characteristics, technology stack, business requirements, and constraints mentioned in the assessment conversation. Ensure all decisions align with the {migration_decision["approach"]} migration approach. Avoid generic recommendations - make decisions specific to this application's context.
"""
        
        ai_response = self._generate_ai_content(prompt, context_data)
        
        # Parse AI response to extract decisions
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                if "decisions" in parsed_response:
                    return parsed_response["decisions"]
        except json.JSONDecodeError:
            pass
        
        # Fallback with minimal hardcoding - still better than the original
        return [
            {
                "area": "Migration Strategy",
                "options": "Rehost vs Replatform vs Refactor",
                "selected": "Assessment-based strategy (requires detailed analysis)",
                "rationale": f"Strategy selection for {assessment_data.application_name} requires analysis of technical complexity and business timeline from assessment"
            },
            {
                "area": "Compute Platform", 
                "options": "Azure VMs vs App Service vs Container Apps",
                "selected": "Platform selection based on architecture analysis",
                "rationale": "Compute platform recommendation depends on application architecture patterns identified in assessment"
            }
        ]
    
    def _prepare_decision_analysis_context(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Prepare comprehensive decision analysis context for AI processing."""
        
        qa_content = []
        tech_indicators = []
        architecture_context = []
        business_context = []
        security_context = []
        performance_context = []
        
        tech_keywords = [
            'technology', 'framework', 'language', 'database', 'server', 'container',
            'docker', 'kubernetes', 'microservices', 'api', 'rest', 'soap'
        ]
        
        architecture_keywords = [
            'architecture', 'design', 'pattern', 'monolith', 'distributed', 'layered',
            'component', 'service', 'tier', 'scale', 'load'
        ]
        
        business_keywords = [
            'business', 'timeline', 'deadline', 'budget', 'cost', 'priority', 'critical',
            'important', 'revenue', 'customer', 'user', 'availability'
        ]
        
        security_keywords = [
            'security', 'authentication', 'authorization', 'compliance', 'audit',
            'encrypt', 'certificate', 'access control', 'firewall'
        ]
        
        performance_keywords = [
            'performance', 'latency', 'throughput', 'response time', 'load', 'capacity',
            'scalability', 'availability', 'uptime'
        ]
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                qa_text = f"Q: {qa.question}\nA: {qa.answer}"
                qa_content.append(qa_text)
                
                content_lower = (qa.question + " " + qa.answer).lower()
                
                # Categorize Q&As by context type
                for keyword in tech_keywords:
                    if keyword in content_lower and keyword not in tech_indicators:
                        tech_indicators.append(keyword)
                
                for keyword in architecture_keywords:
                    if keyword in content_lower:
                        architecture_context.append(qa_text)
                        
                for keyword in business_keywords:
                    if keyword in content_lower:
                        business_context.append(qa_text)
                        
                for keyword in security_keywords:
                    if keyword in content_lower:
                        security_context.append(qa_text)
                        
                for keyword in performance_keywords:
                    if keyword in content_lower:
                        performance_context.append(qa_text)
        
        return {
            "qa_content": "\n\n".join(qa_content[:15]),  # More context for decisions
            "tech_indicators": tech_indicators,
            "architecture_context": architecture_context[:4],
            "business_context": business_context[:4],
            "security_context": security_context[:3],
            "performance_context": performance_context[:3]
        }

    def _extract_application_context(self, questions_answers: List[QuestionAnswer]) -> Dict[str, bool]:
        """Extract comprehensive application-specific context indicators using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return {
                'has_timeline_pressure': False,
                'has_compliance_needs': False,
                'has_large_database': False,
                'has_external_access': False,
                'has_multiple_users': False,
                'has_high_availability_needs': False,
                'has_performance_requirements': False,
                'has_integration_complexity': False
            }
        
        context_prompt = f"""Analyze this application assessment conversation and identify key application context indicators.

ASSESSMENT CONVERSATION:
{qa_text[:2000]}

Analyze the conversation for the following context indicators:

1. **Timeline Pressure**: Are there mentions of urgent deadlines, time constraints, or pressure to migrate quickly?
2. **Compliance Needs**: Are there mentions of regulatory requirements, compliance standards, auditing, or governance needs?
3. **Large Database**: Are there mentions of large databases, data volumes, complex data migration, or database size concerns?
4. **External Access**: Are there mentions of external users, public access, API endpoints, or internet-facing requirements?
5. **Multiple Users**: Are there mentions of many users, user management, authentication for multiple people, or user access control?
6. **High Availability**: Are there mentions of uptime requirements, 24/7 availability, disaster recovery, or business continuity?
7. **Performance Requirements**: Are there mentions of performance needs, latency concerns, speed requirements, or SLA targets?
8. **Integration Complexity**: Are there mentions of multiple systems, complex integrations, API dependencies, or system interconnections?

For each indicator, determine if it's clearly present in the conversation.

Return as JSON:
{{
    "context_analysis": {{
        "has_timeline_pressure": true/false,
        "has_compliance_needs": true/false, 
        "has_large_database": true/false,
        "has_external_access": true/false,
        "has_multiple_users": true/false,
        "has_high_availability_needs": true/false,
        "has_performance_requirements": true/false,
        "has_integration_complexity": true/false
    }},
    "evidence": {{
        "timeline_evidence": "specific quote or mention if found",
        "compliance_evidence": "specific quote or mention if found",
        "database_evidence": "specific quote or mention if found",
        "access_evidence": "specific quote or mention if found",
        "users_evidence": "specific quote or mention if found",
        "availability_evidence": "specific quote or mention if found",
        "performance_evidence": "specific quote or mention if found",
        "integration_evidence": "specific quote or mention if found"
    }}
}}

Base analysis only on what's actually mentioned in the conversation, not assumptions."""

        ai_response = self._generate_ai_content(context_prompt)
        
        # Parse AI response
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                context_analysis = parsed_response.get('context_analysis', {})
                
                # Ensure all required keys exist
                default_context = {
                    'has_timeline_pressure': False,
                    'has_compliance_needs': False,
                    'has_large_database': False,
                    'has_external_access': False,
                    'has_multiple_users': False,
                    'has_high_availability_needs': False,
                    'has_performance_requirements': False,
                    'has_integration_complexity': False
                }
                
                # Update with AI analysis results
                default_context.update(context_analysis)
                return default_context
                
        except json.JSONDecodeError:
            pass
        
        # Fallback to keyword-based analysis if AI parsing fails
        context = {
            'has_timeline_pressure': False,
            'has_compliance_needs': False,
            'has_large_database': False,
            'has_external_access': False,
            'has_multiple_users': False,
            'has_high_availability_needs': False,
            'has_performance_requirements': False,
            'has_integration_complexity': False
        }
        
        combined_text = qa_text.lower()
        
        # Timeline pressure indicators
        timeline_keywords = ['urgent', 'asap', 'deadline', 'quickly', 'fast', 'rush', 'immediate', 'soon']
        context['has_timeline_pressure'] = any(keyword in combined_text for keyword in timeline_keywords)
        
        # Compliance indicators
        compliance_keywords = ['compliance', 'regulatory', 'audit', 'governance', 'sox', 'hipaa', 'gdpr', 'pci']
        context['has_compliance_needs'] = any(keyword in combined_text for keyword in compliance_keywords)
        
        # Database size indicators
        database_keywords = ['large database', 'big data', 'terabyte', 'gigabyte', 'millions of records', 'data migration']
        context['has_large_database'] = any(keyword in combined_text for keyword in database_keywords)
        
        # External access indicators
        external_keywords = ['public', 'external', 'internet', 'customer', 'client access', 'api']
        context['has_external_access'] = any(keyword in combined_text for keyword in external_keywords)
        
        # Multi-user indicators
        user_keywords = ['users', 'people', 'team', 'multiple', 'authentication', 'login']
        context['has_multiple_users'] = any(keyword in combined_text for keyword in user_keywords)
        
        # High availability indicators
        availability_keywords = ['24/7', '99.9', 'uptime', 'availability', 'disaster recovery', 'failover']
        context['has_high_availability_needs'] = any(keyword in combined_text for keyword in availability_keywords)
        
        # Performance indicators
        performance_keywords = ['performance', 'fast', 'slow', 'latency', 'response time', 'speed']
        context['has_performance_requirements'] = any(keyword in combined_text for keyword in performance_keywords)
        
        # Integration complexity indicators
        integration_keywords = ['integrate', 'api', 'system', 'connect', 'interface', 'dependencies']
        context['has_integration_complexity'] = any(keyword in combined_text for keyword in integration_keywords)
        
        return context

    def _generate_decision_rationale(self, decisions: List[Dict[str, str]], assessment_data: AssessmentReportData) -> List[str]:
        """Generate high-level rationale points based on decisions made, using centralized migration approach."""
        
        rationale_points = []
        
        # Get the centralized migration approach for consistency
        migration_decision = self._determine_migration_approach(assessment_data.questions_answers)
        
        # Add migration strategy rationale (always first for consistency)
        rationale_points.append(f"Adopt {migration_decision['approach']} migration approach: {migration_decision['justification'][:100]}...")
        
        # Add technology-specific rationale
        compute_decision = next((d for d in decisions if d['area'] == 'Compute Platform'), None)
        if compute_decision and 'Container' in compute_decision.get('selected', ''):
            rationale_points.append(f"Leverage existing containerization for improved scalability and deployment aligned with {migration_decision['approach']} approach")
        elif compute_decision and 'App Service' in compute_decision.get('selected', ''):
            rationale_points.append(f"Utilize managed services to reduce operational overhead consistent with {migration_decision['approach']} strategy")
        
        # Add security rationale
        security_decision = next((d for d in decisions if d['area'] == 'Network Security'), None)
        if security_decision and 'Private' in security_decision.get('selected', ''):
            rationale_points.append(f"Implement enhanced security controls based on application requirements and {migration_decision['approach']} migration needs")
        
        # Add database-specific rationale
        db_decision = next((d for d in decisions if d['area'] == 'Database Platform'), None)
        if db_decision and 'Managed' in db_decision.get('selected', ''):
            rationale_points.append(f"Reduce database operational complexity through Azure managed services supporting {migration_decision['approach']} approach")
        
        # Add default rationale points if none generated
        if len(rationale_points) <= 1:  # Only migration approach added
            rationale_points.extend([
                f"Maintain compatibility with existing processes and tools during {migration_decision['approach']} migration",
                "Ensure minimal business disruption during transition",
                f"Enable future cloud optimization opportunities through {migration_decision['approach']} approach"
            ])
        
        return rationale_points[:4]  # Limit to 4 points

    def _enhance_decisions_with_llm(self, decisions: List[Dict[str, str]], assessment_data: AssessmentReportData) -> List[Dict[str, str]]:
        """Enhance all decision matrix entries with comprehensive LLM analysis."""
        
        if not self.llm_client:
            return decisions
        
        try:
            # Prepare comprehensive Q&A context
            qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                               if qa.is_answered and qa.answer != "Not addressed in transcript"])
            
            # Get technology and business context
            tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
            
            # Create comprehensive decision analysis prompt
            decisions_prompt = f"""Based on this detailed application assessment, enhance the migration decision matrix with intelligent, context-specific rationale:

APPLICATION CONTEXT:
Q&A Data: {qa_text[:2500]}

TECHNOLOGY ANALYSIS:
Technology Stack: {tech_stack}
Application Name: {assessment_data.application_name}

CURRENT DECISIONS:
{chr(10).join([f"• {d['area']}: {d['selected']} - {d['rationale']}" for d in decisions])}

Enhance each decision with application-specific rationale that references actual details from the Q&A assessment. 
Provide a JSON response with enhanced decisions:
{{
    "enhanced_decisions": [
        {{
            "area": "Migration Strategy",
            "selected": "enhanced selection based on analysis",
            "rationale": "detailed rationale that references specific Q&A findings (100 chars max)"
        }},
        {{
            "area": "Compute Platform", 
            "selected": "enhanced selection",
            "rationale": "application-specific reasoning (100 chars max)"
        }},
        {{
            "area": "Database Platform",
            "selected": "enhanced selection", 
            "rationale": "database-specific reasoning (100 chars max)"
        }},
        {{
            "area": "Network Security",
            "selected": "enhanced selection",
            "rationale": "security-specific reasoning (100 chars max)"
        }},
        {{
            "area": "Identity Management",
            "selected": "enhanced selection",
            "rationale": "identity-specific reasoning (100 chars max)"
        }}
    ]
}}

Make rationale highly specific to this application, referencing actual technologies, constraints, or requirements mentioned in the Q&A."""

            response = self.llm_client.invoke(decisions_prompt)
            
            import json
            result = json.loads(response.content.strip())
            enhanced_decisions = result.get('enhanced_decisions', [])
            
            # Apply enhancements to original decisions
            for i, decision in enumerate(decisions):
                if i < len(enhanced_decisions):
                    enhanced = enhanced_decisions[i]
                    decision['selected'] = enhanced.get('selected', decision['selected'])
                    decision['rationale'] = enhanced.get('rationale', decision['rationale'])
                    
                    # Add intelligence marker for verification
                    decision['rationale'] = f"LLM-Analysis: {decision['rationale']}"
            
            return decisions
            
        except Exception as e:
            print(f"LLM enhancement failed: {e}")
            return decisions
    
    def _add_key_contacts_table(self, doc, contacts: List[str]):
        """Add key contacts as a proper Word table with intelligent contact extraction."""
        
        if contacts:
            doc.add_paragraph("Key project contacts identified from the assessment:")
            for contact in contacts[:5]:  # Show up to 5 contacts
                bullet_para = doc.add_paragraph(f"• {contact}")
                bullet_para.style = 'List Bullet'
        else:
            # Generate intelligent contact roles based on application context
            doc.add_paragraph("The following table identifies the key contacts for this migration project:")
            
            # Create the table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set up headers
            header_cells = table.rows[0].cells
            headers = ['Role', 'Name', 'Email', 'Responsibilities']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Make header text bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Generate intelligent contact structure
            contacts_data = self._generate_project_contacts()
            
            # Add each contact as a table row
            for contact in contacts_data:
                row_cells = table.add_row().cells
                row_cells[0].text = contact['role']
                row_cells[1].text = contact['name']
                row_cells[2].text = contact['email']
                row_cells[3].text = contact['responsibilities']
            
            doc.add_paragraph("Note: Contact details to be finalized during project initiation phase.")

    def _generate_project_contacts(self) -> List[Dict[str, str]]:
        """Generate intelligent project contact structure based on migration complexity."""
        
        contacts_data = [
            {
                'role': 'Executive Sponsor',
                'name': 'To be assigned',
                'email': 'TBD',
                'responsibilities': 'Strategic oversight, budget approval, and executive decision-making'
            },
            {
                'role': 'Project Manager',
                'name': 'To be assigned',
                'email': 'TBD',
                'responsibilities': 'Project coordination, timeline management, and stakeholder communication'
            },
            {
                'role': 'Application Owner',
                'name': 'To be assigned',
                'email': 'TBD',
                'responsibilities': 'Business requirements validation and user acceptance testing'
            },
            {
                'role': 'Technical Lead',
                'name': 'To be assigned',
                'email': 'TBD',
                'responsibilities': 'Architecture design, technical decisions, and implementation oversight'
            },
            {
                'role': 'Infrastructure Team',
                'name': 'To be assigned',
                'email': 'TBD',
                'responsibilities': 'Azure environment setup, security configuration, and operational support'
            }
        ]
        
        return contacts_data
    
    def _add_cost_breakdown_table(self, doc, tech_stack: Dict[str, Any], total_min_cost: int, total_max_cost: int, cost_breakdown: List[str]):
        """Add cost breakdown as a proper Word table."""
        
        # Create the table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set up headers
        header_cells = table.rows[0].cells
        headers = ['Service Category', 'Estimated Monthly Cost']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header text bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add cost breakdown rows
        for cost_item in cost_breakdown:
            # Parse the markdown table format
            if "|" in cost_item and cost_item.strip().startswith("|"):
                # Remove leading/trailing pipes and split
                parts = [part.strip() for part in cost_item.split("|")[1:-1]]
                if len(parts) >= 2:
                    row_cells = table.add_row().cells
                    row_cells[0].text = parts[0]
                    row_cells[1].text = parts[1]
        
        # Add total row
        total_row = table.add_row().cells
        total_run = total_row[0].paragraphs[0].add_run("Total Estimated")
        total_run.bold = True
        cost_run = total_row[1].paragraphs[0].add_run(f"${total_min_cost:,} - ${total_max_cost:,}")
        cost_run.bold = True
        
        # Add technology-specific cost considerations
        doc.add_paragraph("")  # Add spacing
        
        considerations_para = doc.add_paragraph()
        considerations_title = considerations_para.add_run("Cost Analysis Based on Current Technology Stack:")
        considerations_title.bold = True
        
        if tech_stack.get('containers'):
            bullet_para = doc.add_paragraph("• Containerization Advantage: Existing containers reduce migration costs and enable efficient resource utilization")
            bullet_para.style = 'List Bullet'
        
        if tech_stack.get('cloud_ready'):
            bullet_para = doc.add_paragraph("• Cloud-Ready Architecture: Modern technology stack reduces migration costs and speeds up migration timeline")
            bullet_para.style = 'List Bullet'
        
        if tech_stack.get('databases'):
            bullet_para = doc.add_paragraph("• Database Migration: Leverage Azure Database Migration Service for cost-effective database transitions")
            bullet_para.style = 'List Bullet'
    
    def _extract_key_contacts(self, questions_answers: List[QuestionAnswer]) -> List[str]:
        """Extract key contacts from Q&A data."""
        contacts = []
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in ['contact', 'owner', 'responsible', 'team', 'manager']):
                    contacts.append(qa.answer)
        return contacts
    
    def _calculate_azure_costs(self, azure_services: Dict[str, Any], tech_stack: Dict[str, Any], deployment_method: str) -> tuple:
        """Calculate Azure costs using intelligent analysis based on application requirements."""
        
        # Prepare context for AI-driven cost analysis
        context_data = {
            "azure_services": azure_services,
            "technology_stack": tech_stack,
            "deployment_method": deployment_method,
            "databases": tech_stack.get('databases', []),
            "languages": tech_stack.get('languages', []),
            "frameworks": tech_stack.get('frameworks', []),
            "containers": tech_stack.get('containers', False),
            "cloud_ready": tech_stack.get('cloud_ready', False),
            "architecture_type": tech_stack.get('architecture_type', 'unknown'),
            "modernization_level": tech_stack.get('modernization_level', 'unknown')
        }
        
        prompt = """
Analyze this application technology stack and provide intelligent Azure cost estimates based on the specific technologies and architecture identified.

Base your cost analysis on:
- Technology stack complexity and requirements
- Architecture type and scalability needs  
- Database technologies and expected usage
- Compute requirements based on application type
- Networking complexity based on architecture
- Security requirements based on application context
- Monitoring needs based on technology stack

Provide realistic monthly cost estimates in USD for a production environment.

Return as JSON:
{
    "cost_analysis": {
        "total_min_cost": <number>,
        "total_max_cost": <number>,
        "cost_breakdown": [
            "| Service Category | $min - $max |",
            "| Another Service | $min - $max |"
        ],
        "cost_factors": [
            "Factor 1 affecting costs",
            "Factor 2 affecting costs"
        ],
        "optimization_opportunities": [
            "Cost optimization opportunity 1",
            "Cost optimization opportunity 2"
        ]
    }
}

Consider these Azure services and their typical cost ranges:
- Azure VMs: $200-$1200/month depending on size and performance needs
- Azure App Service: $150-$800/month based on tier and scaling requirements  
- Azure Kubernetes Service: $300-$1500/month including node pool costs
- Azure SQL Database: $100-$800/month based on DTU/vCore requirements
- Azure Database for PostgreSQL/MySQL: $100-$600/month based on compute/storage
- Azure Cache for Redis: $50-$400/month based on cache size and tier
- Azure Application Gateway: $100-$300/month including bandwidth
- Azure Monitor + App Insights: $50-$300/month based on telemetry volume
- Azure Key Vault: $5-$50/month based on operations
- Azure Storage: $20-$200/month based on volume and access patterns
- Azure Virtual Network: $20-$100/month based on complexity

Adjust estimates based on the specific technology stack and architecture complexity identified.
"""
        
        ai_response = self._generate_ai_content(prompt, context_data)
        
        # Parse AI response for cost data
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                cost_data = parsed_response.get('cost_analysis', {})
                
                total_min = cost_data.get('total_min_cost', 0)
                total_max = cost_data.get('total_max_cost', 0)
                breakdown = cost_data.get('cost_breakdown', [])
                
                if total_min > 0 and total_max > 0 and breakdown:
                    return total_min, total_max, breakdown
                    
        except json.JSONDecodeError:
            pass
        
        # Fallback to technology-based cost calculation if AI parsing fails
        return self._fallback_cost_calculation(azure_services, tech_stack, deployment_method)
    
    def _fallback_cost_calculation(self, azure_services: Dict[str, Any], tech_stack: Dict[str, Any], deployment_method: str) -> tuple:
        """Fallback cost calculation based on technology stack analysis."""
        total_min_cost = 0
        total_max_cost = 0
        cost_breakdown = []
        
        # Compute service costs based on deployment method and technology
        if deployment_method == "kubernetes" or tech_stack.get('containers'):
            cost_breakdown.append("| Azure Kubernetes Service (AKS) | $300 - $800 |")
            total_min_cost += 300
            total_max_cost += 800
        elif tech_stack.get('cloud_ready') and tech_stack.get('frameworks'):
            cost_breakdown.append("| Azure App Service (Premium) | $200 - $500 |")
            total_min_cost += 200
            total_max_cost += 500
        else:
            cost_breakdown.append("| Azure Virtual Machines | $400 - $800 |")
            total_min_cost += 400
            total_max_cost += 800
        
        # Database costs based on detected technologies
        db_technologies = tech_stack.get('databases', [])
        if 'PostgreSQL' in str(db_technologies):
            cost_breakdown.append("| Azure Database for PostgreSQL | $150 - $400 |")
            total_min_cost += 150
            total_max_cost += 400
        elif 'MySQL' in str(db_technologies):
            cost_breakdown.append("| Azure Database for MySQL | $150 - $350 |")
            total_min_cost += 150
            total_max_cost += 350
        elif any('SQL' in str(db) for db in db_technologies):
            cost_breakdown.append("| Azure SQL Database | $200 - $500 |")
            total_min_cost += 200
            total_max_cost += 500
        else:
            cost_breakdown.append("| Azure SQL Database (Default) | $200 - $500 |")
            total_min_cost += 200
            total_max_cost += 500
        
        # Add caching if Redis detected
        if 'Redis' in str(db_technologies):
            cost_breakdown.append("| Azure Cache for Redis | $100 - $250 |")
            total_min_cost += 100
            total_max_cost += 250
        
        # Standard infrastructure costs
        cost_breakdown.append("| Application Gateway + VNet | $150 - $300 |")
        total_min_cost += 150
        total_max_cost += 300
        
        cost_breakdown.append("| Key Vault + Security Center | $50 - $150 |")
        total_min_cost += 50
        total_max_cost += 150
        
        cost_breakdown.append("| Azure Monitor + App Insights | $100 - $250 |")
        total_min_cost += 100
        total_max_cost += 250
        
        cost_breakdown.append("| Blob Storage + File Storage | $50 - $150 |")
        total_min_cost += 50
        total_max_cost += 150
        
        return total_min_cost, total_max_cost, cost_breakdown
    
    def _add_cost_optimization_section(self, doc, migration_pattern: Dict[str, Any], deployment_method: str):
        """Add cost optimization recommendations section."""
        
        # Add migration pattern specific considerations
        if 'Replatform' in migration_pattern['pattern']:
            bullet_para = doc.add_paragraph("• Replatform Strategy: Moderate migration costs with significant long-term operational savings")
            bullet_para.style = 'List Bullet'
        elif 'Rehost' in migration_pattern['pattern']:
            bullet_para = doc.add_paragraph("• Rehost Strategy: Lower initial migration costs, gradual modernization approach")
            bullet_para.style = 'List Bullet'
        
        doc.add_paragraph("")  # Add spacing
        
        # Cost optimization opportunities
        optimization_para = doc.add_paragraph()
        optimization_title = optimization_para.add_run("Cost Optimization Opportunities:")
        optimization_title.bold = True
        
        optimization_points = [
            "Reserved Instances: 30-50% savings for predictable workloads",
            "Azure Hybrid Benefit: Leverage existing licenses for Windows/SQL Server",
            "Auto-scaling: Optimize resource utilization based on demand",
            "Spot Instances: Up to 90% savings for development/testing environments",
            "Azure Cost Management: Continuous monitoring and optimization"
        ]
        
        if deployment_method == "kubernetes":
            optimization_points.append("AKS Cost Optimization: Node auto-scaling and resource quotas")
        
        for point in optimization_points:
            bullet_para = doc.add_paragraph(f"• {point}")
            bullet_para.style = 'List Bullet'
        
        # Add note
        doc.add_paragraph("")  # Add spacing
        note_text = ("Note: Costs are indicative and based on recommended Azure services from technology analysis. "
                    "Actual costs may vary based on usage patterns, data transfer, and specific service configurations. "
                    "A detailed Azure Pricing Calculator assessment will be performed during planning phase.")
        
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(note_text)
        note_run.italic = True
    
    def _generate_environment_specific_content(self, env_name: str, section_type: str, assessment_data: AssessmentReportData) -> str:
        """Generate environment-specific content based on section type."""
        
        env_lower = env_name.lower()
        
        if section_type == "logical_architecture":
            return self._generate_logical_architecture_content(env_name, assessment_data)
        elif section_type == "network_flow":
            return self._generate_network_flow_content(env_name, assessment_data)
        elif section_type == "proposed_architecture":
            return self._generate_proposed_architecture_content(env_name, assessment_data)
        else:
            return f"Content for {env_name} environment - {section_type}"
    
    def _generate_logical_architecture_content(self, env_name: str, assessment_data: AssessmentReportData) -> str:
        """Generate logical architecture content for specific environment using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return f"The {env_name} environment architecture for {assessment_data.application_name} requires detailed analysis based on application requirements."
        
        architecture_prompt = f"""Analyze this application assessment and generate logical architecture content for the {env_name} environment.

APPLICATION: {assessment_data.application_name}
ENVIRONMENT: {env_name}

TRANSCRIPT DATA:
{qa_text[:2500]}

Based on the conversation, generate logical architecture content that:
1. Describes the current architecture mentioned in the discussion
2. Identifies key components and technologies discussed
3. Explains environment-specific requirements mentioned
4. References actual technologies, frameworks, or systems identified
5. Addresses scalability, availability, and performance considerations discussed

Generate a detailed architectural description (200-400 words) that references actual details from the conversation, not generic content.

Focus on what was actually discussed about the application architecture, technology stack, and environment requirements."""

        # Get AI analysis
        result = self._llm_analyze(architecture_prompt, 
            f"The {env_name} environment for {assessment_data.application_name} requires architecture analysis based on the application's specific requirements and technology stack.")
        
        if isinstance(result, str) and len(result) > 50:
            return result
        
        return f"The {env_name} environment architecture for {assessment_data.application_name} requires detailed analysis based on application requirements identified in the assessment."

    def _generate_network_flow_content(self, env_name: str, assessment_data: AssessmentReportData) -> str:
        """Generate network flow content for specific environment using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return f"Network flow analysis for {env_name} environment requires detailed assessment of {assessment_data.application_name} connectivity requirements."
        
        network_flow_prompt = f"""Analyze this application assessment and generate network flow content for the {env_name} environment.

APPLICATION: {assessment_data.application_name}
ENVIRONMENT: {env_name}

TRANSCRIPT DATA:
{qa_text[:2500]}

Based on the conversation, generate network flow content that:
1. Describes current network connectivity patterns mentioned
2. Identifies external systems or integrations discussed
3. Explains data flow and communication patterns from the discussion
4. References actual network requirements, ports, or protocols mentioned
5. Addresses security and connectivity concerns discussed

Generate a detailed network flow description (150-300 words) based on actual requirements discussed, not generic content.

Focus on real network patterns, integrations, and connectivity requirements mentioned in the conversation."""

        # Get AI analysis
        result = self._llm_analyze(network_flow_prompt, 
            f"Network flow for {env_name} environment requires analysis of {assessment_data.application_name} connectivity and integration requirements.")
        
        if isinstance(result, str) and len(result) > 50:
            return result
        
        return f"Network flow analysis for {env_name} environment requires detailed assessment of {assessment_data.application_name} connectivity requirements based on the application architecture."

    def _generate_proposed_architecture_content(self, env_name: str, assessment_data: AssessmentReportData) -> str:
        """Generate proposed Azure architecture content for specific environment using AI analysis and network traffic data."""
        
        # Start with target architecture if available
        content = ""
        if hasattr(assessment_data, 'target_architecture') and assessment_data.target_architecture:
            content += self._format_target_architecture_content(env_name, assessment_data.target_architecture)
            content += "\n\n"
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in assessment_data.questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            if content:
                return content + f"Additional architecture analysis for {env_name} environment requires detailed review of {assessment_data.application_name} specific requirements."
            return f"Proposed Azure architecture for {env_name} environment requires detailed analysis of {assessment_data.application_name} requirements."
        
        # Enhanced architecture prompt that incorporates network traffic analysis
        architecture_prompt = f"""Analyze this application assessment and generate proposed Azure architecture content for the {env_name} environment.

APPLICATION: {assessment_data.application_name}
ENVIRONMENT: {env_name}

TRANSCRIPT DATA:
{qa_text[:2000]}

NETWORK TRAFFIC ANALYSIS:
{self._format_network_analysis_summary(assessment_data)}

Based on the conversation and network traffic analysis, generate proposed Azure architecture that:
1. Maps current technologies mentioned to appropriate Azure services
2. Addresses specific requirements and constraints discussed
3. Incorporates network traffic patterns for optimal service placement
4. Proposes Azure services that match the discovered architecture patterns
5. Addresses scalability, security, and operational needs based on actual traffic

Generate detailed Azure architecture recommendations (300-500 words) with specific Azure services based on the actual application requirements and network traffic patterns.

Structure the response with sections like:
- Compute Resources (based on discovered applications and Azure Migrate data)
- Database (based on database technologies and port analysis)  
- Networking (based on traffic patterns and connectivity analysis)
- Security (based on discovered ports and security requirements)
- Monitoring (based on operational needs and traffic monitoring)

Use actual Azure service names and reference real requirements from the conversation and network analysis."""

        # Get AI analysis
        ai_result = self._llm_analyze(architecture_prompt, 
            f"Proposed Azure architecture for {env_name} environment requires analysis of {assessment_data.application_name} specific requirements and technology stack.")
        
        if isinstance(ai_result, str) and len(ai_result) > 100:
            content += ai_result
        else:
            content += f"Proposed Azure architecture for {env_name} environment requires detailed analysis of {assessment_data.application_name} requirements and current technology stack to recommend appropriate Azure services."
        
        return content

    def _extract_security_considerations(self, questions_answers: List[QuestionAnswer]) -> List[Dict[str, str]]:
        """Extract security-related information from Q&A data using comprehensive AI analysis following MigrationPlanGenerator pattern."""
        
        # Prepare comprehensive context for AI analysis
        security_context = self._prepare_security_analysis_context(questions_answers)
        
        context_data = {
            "transcript_data": security_context["qa_content"],
            "security_keywords_mentioned": security_context["security_indicators"],
            "compliance_mentions": security_context["compliance_references"],
            "authentication_context": security_context["auth_context"],
            "data_sensitivity_indicators": security_context["data_sensitivity"],
            "total_qa_pairs": len(questions_answers),
            "answered_qa_pairs": len([qa for qa in questions_answers if qa.is_answered])
        }
        
        prompt = """
Analyze this application assessment and extract comprehensive security considerations for Azure migration.

Focus on identifying and analyzing:

1. **Authentication & Authorization Requirements**
   - Current identity systems mentioned
   - User management approaches discussed
   - Access control requirements identified
   - Multi-factor authentication needs

2. **Data Protection & Encryption**
   - Sensitive data types mentioned in conversations
   - Current encryption practices discussed
   - Data classification requirements
   - Backup and recovery security needs

3. **Compliance & Regulatory Requirements**
   - Specific compliance frameworks mentioned (GDPR, HIPAA, SOX, etc.)
   - Audit requirements discussed
   - Data residency concerns
   - Retention policies mentioned

4. **Network Security Considerations**
   - Current network security practices
   - Firewall and access control requirements
   - VPN and connectivity security needs
   - Micro-segmentation requirements

5. **Application Security Requirements**
   - Code security practices mentioned
   - API security requirements
   - Certificate management needs
   - Secrets and configuration management

6. **Operational Security Considerations**
   - Monitoring and logging requirements
   - Incident response procedures
   - Security operations workflows
   - Vulnerability management approaches

Generate specific, actionable security considerations based on the actual conversation content. Each consideration should include:
- The specific requirement extracted from the discussion
- Detailed context from the conversation
- Business priority (High/Medium/Low) based on criticality mentioned
- Recommended Azure security service or approach
- Implementation complexity assessment

Return as JSON:
{
    "security_considerations": [
        {
            "requirement": "specific security requirement name",
            "details": "detailed explanation based on conversation content",
            "conversation_reference": "direct quote or reference from discussion",
            "priority": "High/Medium/Low",
            "azure_recommendation": "specific Azure security service",
            "implementation_complexity": "Low/Medium/High",
            "rationale": "why this is important for the migration"
        }
    ]
}

If specific security requirements are not clearly mentioned, make intelligent inferences based on:
- Application type and technology stack discussed
- Business context and industry indicators
- Integration and connectivity requirements mentioned
- Data handling patterns identified in conversation
"""
        
        ai_response = self._generate_ai_content(prompt, context_data)
        
        # Parse AI response to extract security considerations
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                if "security_considerations" in parsed_response:
                    return parsed_response["security_considerations"]
        except json.JSONDecodeError:
            pass
        
        # Fallback to default security considerations if AI parsing fails
        return [
            {
                "requirement": "Identity and Access Management",
                "details": "Implement comprehensive identity management for Azure migration",
                "conversation_reference": "Security requirements need detailed assessment",
                "priority": "High",
                "azure_recommendation": "Azure Active Directory with conditional access policies",
                "implementation_complexity": "Medium",
                "rationale": "Essential for secure cloud migration and user management"
            },
            {
                "requirement": "Data Encryption and Protection",
                "details": "Ensure data encryption in transit and at rest",
                "conversation_reference": "Data protection requirements to be determined",
                "priority": "High", 
                "azure_recommendation": "Azure Key Vault and Azure Disk Encryption",
                "implementation_complexity": "Medium",
                "rationale": "Critical for maintaining data security during and after migration"
            }
        ]
    
    def _prepare_security_analysis_context(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Prepare comprehensive security analysis context for AI processing."""
        
        qa_content = []
        security_indicators = []
        compliance_references = []
        auth_context = []
        data_sensitivity = []
        
        security_keywords = [
            'security', 'authentication', 'authorization', 'encryption', 'compliance',
            'audit', 'gdpr', 'hipaa', 'sox', 'pci', 'access control', 'firewall',
            'certificate', 'ssl', 'tls', 'identity', 'password', 'mfa', 'sso'
        ]
        
        compliance_keywords = [
            'gdpr', 'hipaa', 'sox', 'pci', 'compliance', 'regulation', 'audit',
            'retention', 'privacy', 'data protection', 'regulatory'
        ]
        
        auth_keywords = [
            'active directory', 'ldap', 'oauth', 'saml', 'sso', 'identity',
            'authentication', 'login', 'user management', 'access'
        ]
        
        data_keywords = [
            'sensitive', 'confidential', 'personal', 'customer data', 'payment',
            'financial', 'health', 'medical', 'private', 'classified'
        ]
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                qa_text = f"Q: {qa.question}\nA: {qa.answer}"
                qa_content.append(qa_text)
                
                # Check for security indicators
                content_lower = (qa.question + " " + qa.answer).lower()
                
                for keyword in security_keywords:
                    if keyword in content_lower and keyword not in security_indicators:
                        security_indicators.append(keyword)
                
                for keyword in compliance_keywords:
                    if keyword in content_lower and keyword not in compliance_references:
                        compliance_references.append(keyword)
                        
                for keyword in auth_keywords:
                    if keyword in content_lower:
                        auth_context.append(qa_text)
                        
                for keyword in data_keywords:
                    if keyword in content_lower:
                        data_sensitivity.append(qa_text)
        
        return {
            "qa_content": "\n\n".join(qa_content[:10]),  # Limit to prevent token overflow
            "security_indicators": security_indicators,
            "compliance_references": compliance_references, 
            "auth_context": auth_context[:3],  # Top 3 authentication-related Q&As
            "data_sensitivity": data_sensitivity[:3]  # Top 3 data sensitivity Q&As
        }
    
    def _extract_network_requirements(self, questions_answers: List[QuestionAnswer]) -> List[Dict[str, str]]:
        """Extract network-related requirements from Q&A data using comprehensive AI analysis following MigrationPlanGenerator pattern."""
        
        # Prepare comprehensive context for AI analysis
        network_context = self._prepare_network_analysis_context(questions_answers)
        
        context_data = {
            "transcript_data": network_context["qa_content"],
            "network_technologies_mentioned": network_context["network_indicators"],
            "connectivity_requirements": network_context["connectivity_context"],
            "performance_requirements": network_context["performance_context"],
            "integration_requirements": network_context["integration_context"],
            "total_qa_pairs": len(questions_answers),
            "answered_qa_pairs": len([qa for qa in questions_answers if qa.is_answered])
        }
        
        prompt = """
Analyze this application assessment and extract comprehensive network requirements for Azure migration.

Focus on identifying and analyzing:

1. **Connectivity Requirements**
   - Internet connectivity needs mentioned
   - On-premises connectivity requirements
   - VPN or ExpressRoute needs discussed
   - Network bandwidth requirements
   - Latency sensitivity mentioned

2. **Load Balancing and Traffic Management**
   - Current load balancing setup discussed
   - Traffic distribution requirements
   - Session affinity needs
   - Geographic distribution requirements

3. **DNS and Name Resolution**
   - DNS management requirements
   - Custom domain needs mentioned
   - Name resolution patterns discussed
   - External DNS dependencies

4. **Network Security and Segmentation**
   - Network isolation requirements
   - Firewall rules and access controls
   - Network segmentation needs
   - DMZ or perimeter security requirements

5. **Application Integration Requirements**
   - External service integrations mentioned
   - API connectivity requirements
   - Database connectivity patterns
   - File sharing and storage access needs

6. **Performance and Monitoring Requirements**
   - Network performance expectations
   - Monitoring and alerting needs
   - Capacity planning considerations
   - Disaster recovery network requirements

Generate specific, actionable network requirements based on the actual conversation content. Each requirement should include:
- The specific network requirement extracted from discussion
- Technical details mentioned in conversation
- Business impact of the requirement
- Recommended Azure networking service
- Configuration complexity assessment

Return as JSON:
{
    "network_requirements": [
        {
            "requirement": "specific network requirement name",
            "details": "detailed technical explanation from conversation",
            "conversation_reference": "relevant quote or context from discussion",
            "business_impact": "why this is important for the business",
            "azure_solution": "recommended Azure networking service",
            "configuration_complexity": "Low/Medium/High",
            "dependencies": "other systems or services this impacts"
        }
    ]
}

If specific network requirements are not clearly mentioned, make intelligent inferences based on:
- Application architecture patterns discussed
- Integration requirements identified
- Performance expectations mentioned
- Current infrastructure setup described
"""
        
        ai_response = self._generate_ai_content(prompt, context_data)
        
        # Parse AI response to extract network requirements
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                if "network_requirements" in parsed_response:
                    return parsed_response["network_requirements"]
        except json.JSONDecodeError:
            pass
        
        # Fallback to default network requirements if AI parsing fails
        return [
            {
                "requirement": "Load Balancing and High Availability",
                "details": "Implement load balancing for application availability and performance",
                "conversation_reference": "Network requirements need detailed assessment",
                "business_impact": "Ensures application availability and optimal user experience",
                "azure_solution": "Azure Load Balancer or Application Gateway",
                "configuration_complexity": "Medium",
                "dependencies": "Application servers and health monitoring"
            },
            {
                "requirement": "Secure Network Connectivity",
                "details": "Establish secure communication between application components",
                "conversation_reference": "Security and connectivity requirements to be determined",
                "business_impact": "Protects data in transit and ensures secure communications",
                "azure_solution": "Azure Virtual Network with Network Security Groups",
                "configuration_complexity": "Low",
                "dependencies": "Application architecture and security policies"
            }
        ]
    
    def _prepare_network_analysis_context(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Prepare comprehensive network analysis context for AI processing."""
        
        qa_content = []
        network_indicators = []
        connectivity_context = []
        performance_context = []
        integration_context = []
        
        network_keywords = [
            'network', 'connectivity', 'bandwidth', 'latency', 'vpn', 'load balancer',
            'dns', 'domain', 'firewall', 'port', 'protocol', 'subnet', 'vlan',
            'routing', 'gateway', 'proxy', 'cdn', 'ssl', 'certificate'
        ]
        
        connectivity_keywords = [
            'vpn', 'expressroute', 'internet', 'intranet', 'on-premises', 'hybrid',
            'connectivity', 'connection', 'remote access', 'site-to-site'
        ]
        
        performance_keywords = [
            'performance', 'latency', 'bandwidth', 'throughput', 'speed', 'capacity',
            'load', 'traffic', 'bottleneck', 'optimization'
        ]
        
        integration_keywords = [
            'integration', 'api', 'service', 'external', 'third-party', 'partner',
            'interface', 'endpoint', 'webhook', 'callback'
        ]
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                qa_text = f"Q: {qa.question}\nA: {qa.answer}"
                qa_content.append(qa_text)
                
                # Check for network indicators
                content_lower = (qa.question + " " + qa.answer).lower()
                
                for keyword in network_keywords:
                    if keyword in content_lower and keyword not in network_indicators:
                        network_indicators.append(keyword)
                
                for keyword in connectivity_keywords:
                    if keyword in content_lower:
                        connectivity_context.append(qa_text)
                        
                for keyword in performance_keywords:
                    if keyword in content_lower:
                        performance_context.append(qa_text)
                        
                for keyword in integration_keywords:
                    if keyword in content_lower:
                        integration_context.append(qa_text)
        
        return {
            "qa_content": "\n\n".join(qa_content[:10]),  # Limit to prevent token overflow
            "network_indicators": network_indicators,
            "connectivity_context": connectivity_context[:3],  # Top 3 connectivity-related Q&As
            "performance_context": performance_context[:3],  # Top 3 performance-related Q&As
            "integration_context": integration_context[:3]  # Top 3 integration-related Q&As
        }
    
    def _extract_network_requirements_enhanced(self, questions_answers: List[QuestionAnswer], dependency_analysis: Any = None) -> List[Dict[str, str]]:
        """Extract network-related requirements using transcript Q&A and dependency analysis data."""
        
        # Start with transcript-based analysis
        network_requirements = self._extract_network_requirements(questions_answers)
        
        # Enhanced analysis with dependency data if available
        if dependency_analysis:
            dependency_insights = self._analyze_dependency_network_requirements(dependency_analysis)
            
            # Merge dependency insights with transcript analysis
            for insight in dependency_insights:
                # Check if similar requirement exists from transcript
                existing_requirement = None
                for req in network_requirements:
                    if (req.get('category', '').lower() == insight.get('category', '').lower() or
                        any(keyword in req.get('requirement', '').lower() 
                            for keyword in insight.get('requirement', '').lower().split()[:3])):
                        existing_requirement = req
                        break
                
                if existing_requirement:
                    # Enhance existing requirement with dependency data
                    existing_requirement['requirement'] = f"{existing_requirement['requirement']} {insight.get('dependency_detail', '')}"
                    existing_requirement['source'] = f"{existing_requirement.get('source', 'Transcript analysis')} + Dependency analysis"
                else:
                    # Add new requirement from dependency analysis
                    network_requirements.append({
                        'category': insight.get('category', 'Network Connectivity'),
                        'requirement': insight.get('requirement', ''),
                        'source': 'Dependency analysis',
                        'priority': insight.get('priority', 'Medium')
                    })
        
        return network_requirements
    
    def _analyze_dependency_network_requirements(self, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Analyze dependency data to extract network requirements."""
        requirements = []
        
        try:
            # Analyze network segments for subnet/VLAN requirements
            if hasattr(dependency_analysis, 'network_segments') and dependency_analysis.network_segments:
                subnet_count = len(dependency_analysis.network_segments)
                if subnet_count > 1:
                    requirements.append({
                        'category': 'Network Segmentation',
                        'requirement': f'Multi-subnet architecture required with {subnet_count} network segments for proper isolation',
                        'dependency_detail': f'Based on {subnet_count} identified network segments in dependency analysis',
                        'priority': 'High'
                    })
            
            # Analyze connections for port and protocol requirements
            if hasattr(dependency_analysis, 'connections') and dependency_analysis.connections:
                protocols = set()
                ports = set()
                external_connections = 0
                
                for conn in dependency_analysis.connections:
                    if hasattr(conn, 'protocol') and conn.protocol:
                        protocols.add(conn.protocol.upper())
                    if hasattr(conn, 'port') and conn.port:
                        ports.add(str(conn.port))
                    if hasattr(conn, 'source_server') and hasattr(conn, 'target_server'):
                        if 'external' in (conn.source_server + conn.target_server).lower():
                            external_connections += 1
                
                if protocols:
                    requirements.append({
                        'category': 'Protocol Support',
                        'requirement': f'Support required for protocols: {", ".join(sorted(protocols))}',
                        'dependency_detail': f'Identified from {len(dependency_analysis.connections)} dependency connections',
                        'priority': 'High'
                    })
                
                if ports:
                    port_list = sorted([int(p) for p in ports if p.isdigit()])[:10]  # Top 10 ports
                    requirements.append({
                        'category': 'Port Configuration',
                        'requirement': f'Network Security Group rules needed for ports: {", ".join(map(str, port_list))}',
                        'dependency_detail': f'Based on analysis of application dependencies',
                        'priority': 'Medium'
                    })
                
                if external_connections > 0:
                    requirements.append({
                        'category': 'External Connectivity',
                        'requirement': f'External connectivity required for {external_connections} identified external dependencies',
                        'dependency_detail': 'May require Azure Firewall or Application Gateway configuration',
                        'priority': 'High'
                    })
            
            # Analyze external dependencies
            if hasattr(dependency_analysis, 'external_dependencies') and dependency_analysis.external_dependencies:
                ext_count = len(dependency_analysis.external_dependencies)
                requirements.append({
                    'category': 'Internet Connectivity',
                    'requirement': f'Internet access required for {ext_count} external services',
                    'dependency_detail': 'Identified external dependencies requiring outbound connectivity',
                    'priority': 'Medium'
                })
        
        except Exception as e:
            # Fallback in case of dependency analysis parsing issues
            requirements.append({
                'category': 'Network Analysis',
                'requirement': 'Standard Azure networking configuration recommended',
                'dependency_detail': f'Dependency analysis available but could not be fully parsed: {str(e)}',
                'priority': 'Medium'
            })
        
        return requirements
    
    def _extract_identity_providers(self, questions_answers: List[QuestionAnswer]) -> List[Dict[str, str]]:
        """Extract identity and authentication information using comprehensive AI analysis following MigrationPlanGenerator pattern."""
        
        # Prepare comprehensive context for AI analysis
        identity_context = self._prepare_identity_analysis_context(questions_answers)
        
        context_data = {
            "transcript_data": identity_context["qa_content"],
            "identity_technologies_mentioned": identity_context["identity_indicators"],
            "authentication_context": identity_context["auth_context"],
            "user_management_context": identity_context["user_mgmt_context"],
            "sso_requirements": identity_context["sso_context"],
            "total_qa_pairs": len(questions_answers),
            "answered_qa_pairs": len([qa for qa in questions_answers if qa.is_answered])
        }
        
        prompt = """
Analyze this application assessment and extract comprehensive identity management and authentication requirements for Azure migration.

Focus on identifying and analyzing:

1. **Current Authentication Systems**
   - Active Directory implementations mentioned
   - LDAP systems discussed
   - Database authentication patterns
   - Custom authentication solutions
   - Third-party identity providers

2. **User Management Requirements**
   - User provisioning processes discussed
   - Role-based access control requirements
   - User lifecycle management needs
   - Group management approaches

3. **Single Sign-On (SSO) Requirements**
   - SSO implementations mentioned
   - SAML configurations discussed
   - OAuth/OpenID Connect requirements
   - Federation requirements

4. **Identity Integration Needs**
   - On-premises identity integration
   - Multi-domain requirements
   - Trust relationships mentioned
   - Directory synchronization needs

5. **Access Control Requirements**
   - Privileged access management
   - Conditional access requirements
   - Multi-factor authentication needs
   - Risk-based authentication

6. **Identity Security Considerations**
   - Password policies mentioned
   - Identity governance requirements
   - Compliance requirements for identity
   - Audit and monitoring needs

Generate specific identity provider recommendations based on the actual conversation content. Each provider should include:
- The identity system or requirement identified
- Current implementation details mentioned
- Migration complexity assessment
- Recommended Azure identity service
- Integration requirements and dependencies

Return as JSON:
{
    "identity_providers": [
        {
            "provider": "identity system name or requirement",
            "current_implementation": "details from conversation about current setup",
            "conversation_reference": "relevant quote or context from discussion",
            "migration_complexity": "Low/Medium/High",
            "azure_recommendation": "specific Azure identity service",
            "integration_requirements": "what needs to be integrated or configured",
            "business_impact": "why this is important for the organization"
        }
    ]
}

If specific identity requirements are not clearly mentioned, make intelligent inferences based on:
- Application architecture and user access patterns discussed
- Security requirements mentioned
- Compliance needs identified
- Current technology stack and enterprise patterns
"""
        
        ai_response = self._generate_ai_content(prompt, context_data)
        
        # Parse AI response to extract identity providers
        try:
            if isinstance(ai_response, str):
                import json
                parsed_response = json.loads(ai_response)
                if "identity_providers" in parsed_response:
                    return parsed_response["identity_providers"]
        except json.JSONDecodeError:
            pass
        
        # Fallback to default identity providers if AI parsing fails
        return [
            {
                "provider": "Azure Active Directory",
                "current_implementation": "Identity management system to be determined from detailed assessment",
                "conversation_reference": "Identity requirements need detailed analysis",
                "migration_complexity": "Medium",
                "azure_recommendation": "Azure AD with conditional access policies",
                "integration_requirements": "User provisioning, role mapping, and SSO configuration",
                "business_impact": "Provides centralized identity management and secure access control"
            },
            {
                "provider": "Multi-Factor Authentication",
                "current_implementation": "Enhanced security authentication to be implemented",
                "conversation_reference": "Security enhancement requirement",
                "migration_complexity": "Low",
                "azure_recommendation": "Azure MFA integrated with Azure AD",
                "integration_requirements": "User enrollment and authentication flow integration",
                "business_impact": "Significantly improves security posture and reduces breach risk"
            }
        ]
    
    def _prepare_identity_analysis_context(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Prepare comprehensive identity analysis context for AI processing."""
        
        qa_content = []
        identity_indicators = []
        auth_context = []
        user_mgmt_context = []
        sso_context = []
        
        identity_keywords = [
            'active directory', 'ldap', 'authentication', 'identity', 'sso', 'saml',
            'oauth', 'openid', 'federation', 'user management', 'access control',
            'rbac', 'role', 'permission', 'login', 'password', 'mfa', 'multi-factor'
        ]
        
        auth_keywords = [
            'authentication', 'login', 'password', 'credential', 'token', 'session',
            'certificate', 'biometric', 'smart card'
        ]
        
        user_mgmt_keywords = [
            'user management', 'provisioning', 'deprovisioning', 'user lifecycle',
            'role assignment', 'group membership', 'access request'
        ]
        
        sso_keywords = [
            'single sign-on', 'sso', 'saml', 'federation', 'claims', 'trust',
            'identity provider', 'relying party'
        ]
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                qa_text = f"Q: {qa.question}\nA: {qa.answer}"
                qa_content.append(qa_text)
                
                # Check for identity indicators
                content_lower = (qa.question + " " + qa.answer).lower()
                
                for keyword in identity_keywords:
                    if keyword in content_lower and keyword not in identity_indicators:
                        identity_indicators.append(keyword)
                
                for keyword in auth_keywords:
                    if keyword in content_lower:
                        auth_context.append(qa_text)
                        
                for keyword in user_mgmt_keywords:
                    if keyword in content_lower:
                        user_mgmt_context.append(qa_text)
                        
                for keyword in sso_keywords:
                    if keyword in content_lower:
                        sso_context.append(qa_text)
        
        return {
            "qa_content": "\n\n".join(qa_content[:10]),  # Limit to prevent token overflow
            "identity_indicators": identity_indicators,
            "auth_context": auth_context[:3],  # Top 3 authentication-related Q&As
            "user_mgmt_context": user_mgmt_context[:3],  # Top 3 user management Q&As
            "sso_context": sso_context[:3]  # Top 3 SSO-related Q&As
        }
    
    def _extract_automation_details(self, questions_answers: List[QuestionAnswer]) -> List[Dict[str, str]]:
        """Extract automation and CI/CD information using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return [{"process": "Automation Assessment Required", "current_state": "No automation information available", "azure_approach": "Azure DevOps implementation recommended"}]
        
        automation_prompt = f"""Analyze this application assessment transcript and extract automation and deployment information.

TRANSCRIPT DATA:
{qa_text[:3000]}

Based on the conversation, identify:
1. Current deployment processes
2. Build and CI/CD pipelines
3. Testing automation
4. Infrastructure as Code usage
5. DevOps tools and practices
6. Release management processes
7. Monitoring and alerting automation

Provide a JSON response:
{{
    "automation_details": [
        {{
            "process": "automation process mentioned",
            "current_state": "how it's currently implemented",
            "tools_used": "tools or platforms mentioned",
            "azure_approach": "recommended Azure DevOps or automation solution"
        }}
    ]
}}

If no specific automation is mentioned, recommend appropriate Azure automation solutions based on the application architecture discussed."""

        # Get AI analysis
        result = self._llm_analyze(automation_prompt, {
            "automation_details": [
                {"process": "Continuous Integration", "current_state": "Manual or basic automation", "tools_used": "To be determined", "azure_approach": "Azure DevOps Pipelines with automated build and test"},
                {"process": "Deployment Automation", "current_state": "Manual deployment process", "tools_used": "Traditional tools", "azure_approach": "Azure DevOps Release Pipelines with Infrastructure as Code"},
                {"process": "Monitoring Automation", "current_state": "Basic monitoring", "tools_used": "Legacy monitoring tools", "azure_approach": "Azure Monitor with automated alerting and Log Analytics"}
            ]
        })
        
        if isinstance(result, dict) and "automation_details" in result:
            return result["automation_details"]
        
        return [{"process": "Automation Assessment Required", "current_state": "Unable to extract automation details from transcript", "azure_approach": "Azure DevOps implementation recommended"}]
    
    def _extract_customer_impact(self, questions_answers: List[QuestionAnswer]) -> List[str]:
        """Extract customer impact information using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return [{"impact_area": "Customer Impact Assessment Required", "description": "No customer impact information available", "mitigation": "Customer impact analysis needed during planning"}]
        
        impact_prompt = f"""Analyze this application assessment transcript and extract customer impact considerations for the migration.

TRANSCRIPT DATA:
{qa_text[:3000]}

Based on the conversation, identify:
1. Customer-facing services and features
2. Service availability requirements 
3. Performance expectations
4. Business continuity needs
5. Downtime tolerance
6. User experience considerations
7. SLA commitments mentioned

Provide a JSON response:
{{
    "customer_impact": [
        {{
            "impact_area": "specific area affecting customers",
            "description": "detailed impact based on discussion",
            "severity": "High/Medium/Low based on business criticality",
            "mitigation": "specific Azure strategy to minimize impact"
        }}
    ]
}}

Focus on real customer impacts mentioned in the conversation, not generic impacts."""

        # Get AI analysis
        result = self._llm_analyze(impact_prompt, {
            "customer_impact": [
                {"impact_area": "Service Availability", "description": "Application downtime affects customer access", "severity": "High", "mitigation": "Blue-green deployment strategy with Azure Traffic Manager"},
                {"impact_area": "Performance", "description": "Response time critical for user experience", "severity": "Medium", "mitigation": "Azure CDN and Application Gateway for optimized performance"},
                {"impact_area": "Data Integrity", "description": "Customer data must remain secure and accessible", "severity": "High", "mitigation": "Azure Backup and point-in-time recovery"}
            ]
        })
        
        if isinstance(result, dict) and "customer_impact" in result:
            return result["customer_impact"]
        
        return [{"impact_area": "Customer Impact Assessment Required", "description": "Unable to extract customer impact from transcript", "mitigation": "Customer impact analysis needed during planning"}]
    
    def _extract_operational_concerns(self, questions_answers: List[QuestionAnswer]) -> List[str]:
        """Extract operational challenges and concerns using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return [{"concern": "Operational Assessment Required", "current_challenge": "No operational information available", "azure_solution": "Comprehensive operational assessment needed"}]
        
        operational_prompt = f"""Analyze this application assessment transcript and extract operational challenges and concerns.

TRANSCRIPT DATA:
{qa_text[:3000]}

Based on the conversation, identify:
1. Current operational challenges mentioned
2. Maintenance and support difficulties
3. Monitoring and troubleshooting pain points
4. Backup and recovery concerns  
5. Performance management issues
6. Resource management challenges
7. Operational team concerns

Provide a JSON response:
{{
    "operational_concerns": [
        {{
            "concern": "specific operational challenge mentioned",
            "current_challenge": "detailed description of the problem",
            "impact": "how this affects operations",
            "azure_solution": "specific Azure service or approach to address this"
        }}
    ]
}}

Focus on real operational challenges discussed, not generic concerns."""

        # Get AI analysis
        result = self._llm_analyze(operational_prompt, {
            "operational_concerns": [
                {"concern": "Monitoring and Alerting", "current_challenge": "Limited visibility into application performance", "impact": "Reactive troubleshooting", "azure_solution": "Azure Monitor with Application Insights for comprehensive observability"},
                {"concern": "Backup and Recovery", "current_challenge": "Manual backup processes", "impact": "Risk of data loss and extended recovery times", "azure_solution": "Azure Backup with automated policies and point-in-time recovery"},
                {"concern": "Scalability Management", "current_challenge": "Manual scaling processes", "impact": "Poor resource utilization and performance", "azure_solution": "Azure auto-scaling and load balancing"}
            ]
        })
        
        if isinstance(result, dict) and "operational_concerns" in result:
            return result["operational_concerns"]
        
        return [{"concern": "Operational Assessment Required", "current_challenge": "Unable to extract operational concerns from transcript", "azure_solution": "Comprehensive operational assessment needed"}]
    
    def _extract_observability_info(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Extract monitoring and observability information using AI analysis."""
        
        # Prepare Q&A context for AI analysis
        qa_text = "\n".join([f"Q: {qa.question}\nA: {qa.answer}" for qa in questions_answers 
                           if qa.is_answered and qa.answer != "Not addressed in transcript"])
        
        if not qa_text.strip():
            return {
                'monitoring': ['Azure Monitor setup required - no observability information in transcript'],
                'alerts': ['Azure Monitor alerting configuration needed'],
                'events': ['Azure Event tracking to be configured'],
                'logging': ['Azure Log Analytics setup required']
            }
        
        observability_prompt = f"""Analyze this application assessment transcript and extract observability and monitoring requirements.

TRANSCRIPT DATA:
{qa_text[:3000]}

Based on the conversation, identify:
1. Current monitoring and observability practices
2. Alerting requirements and thresholds
3. Event tracking and audit needs
4. Logging requirements and retention
5. Performance monitoring needs
6. Health check requirements
7. Dashboard and reporting needs

Provide a JSON response:
{{
    "observability": {{
        "monitoring": ["specific monitoring requirements based on discussion"],
        "alerts": ["specific alerting needs mentioned in conversation"], 
        "events": ["event tracking requirements discussed"],
        "logging": ["logging needs and retention requirements mentioned"]
    }}
}}

Focus on real observability requirements mentioned, not generic recommendations."""

        # Get AI analysis
        result = self._llm_analyze(observability_prompt, {
            "observability": {
                'monitoring': ['Azure Monitor for application performance and infrastructure monitoring based on application requirements'],
                'alerts': ['Performance thresholds and availability alerts configured for business-critical metrics'],
                'events': ['Security events and operational monitoring for audit and compliance'],
                'logging': ['Application and system logs centralized in Azure Log Analytics with appropriate retention']
            }
        })
        
        if isinstance(result, dict) and "observability" in result:
            return result["observability"]
        
        return {
            'monitoring': ['Azure Monitor setup required - unable to extract specific requirements from transcript'],
            'alerts': ['Azure Monitor alerting configuration needed'],
            'events': ['Azure Event tracking to be configured'],
            'logging': ['Azure Log Analytics setup required']
        }
    
    def _generate_architecture_heatmap(self, azure_migrate_data: Any, questions_answers: List[QuestionAnswer]) -> List[Dict[str, str]]:
        """Generate architecture complexity heatmap based on actual Q&A responses."""
        
        # Start with base heatmap
        heatmap = [
            {'area': 'Overall Complexity', 'notes': 'Based on Azure Migrate assessment and Q&A analysis', 'ranking': 'Medium'},
            {'area': 'App Remediation', 'notes': 'Application modifications required', 'ranking': 'Low'},
            {'area': 'Data Migration', 'notes': 'Database and data store migration', 'ranking': 'Medium'},
            {'area': 'Network Configuration', 'notes': 'Network and security setup', 'ranking': 'Medium'},
            {'area': 'Integration Complexity', 'notes': 'External system integrations', 'ranking': 'Low'}
        ]
        
        # Analyze Q&A responses for complexity indicators
        complexity_indicators = {
            'high': ['complex', 'difficult', 'challenging', 'multiple databases', 'many integrations', 'custom', 'legacy'],
            'medium': ['some', 'moderate', 'standard', 'typical', 'several'],
            'low': ['simple', 'basic', 'minimal', 'straightforward', 'standard']
        }
        
        overall_complexity_score = 0
        integration_complexity_score = 0
        data_complexity_score = 0
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer and qa.answer != "Not addressed in transcript":
                answer_lower = qa.answer.lower()
                question_lower = qa.question.lower()
                
                # Check for high complexity indicators
                for indicator in complexity_indicators['high']:
                    if indicator in answer_lower:
                        overall_complexity_score += 2
                        
                        # Specific complexity areas
                        if any(keyword in question_lower for keyword in ['database', 'data', 'storage']):
                            data_complexity_score += 2
                        if any(keyword in question_lower for keyword in ['integration', 'api', 'service', 'external']):
                            integration_complexity_score += 2
                
                # Check for medium complexity indicators
                for indicator in complexity_indicators['medium']:
                    if indicator in answer_lower:
                        overall_complexity_score += 1
                        if any(keyword in question_lower for keyword in ['database', 'data', 'storage']):
                            data_complexity_score += 1
                        if any(keyword in question_lower for keyword in ['integration', 'api', 'service', 'external']):
                            integration_complexity_score += 1
        
        # Update rankings based on analysis
        if overall_complexity_score > 8:
            heatmap[0]['ranking'] = 'High'
            heatmap[0]['notes'] = f'High complexity identified from Q&A analysis (score: {overall_complexity_score})'
        elif overall_complexity_score > 4:
            heatmap[0]['ranking'] = 'Medium'
            heatmap[0]['notes'] = f'Medium complexity identified from Q&A analysis (score: {overall_complexity_score})'
        else:
            heatmap[0]['ranking'] = 'Low'
            heatmap[0]['notes'] = f'Low complexity identified from Q&A analysis (score: {overall_complexity_score})'
        
        if data_complexity_score > 4:
            heatmap[2]['ranking'] = 'High'
            heatmap[2]['notes'] = 'Complex data migration requirements identified'
        elif data_complexity_score > 2:
            heatmap[2]['ranking'] = 'Medium'
            heatmap[2]['notes'] = 'Moderate data migration complexity'
        else:
            heatmap[2]['ranking'] = 'Low'
            heatmap[2]['notes'] = 'Standard data migration approach'
        
        if integration_complexity_score > 4:
            heatmap[4]['ranking'] = 'High'
            heatmap[4]['notes'] = 'Complex integration requirements identified'
        elif integration_complexity_score > 2:
            heatmap[4]['ranking'] = 'Medium' 
            heatmap[4]['notes'] = 'Moderate integration complexity'
        else:
            heatmap[4]['ranking'] = 'Low'
            heatmap[4]['notes'] = 'Standard integration approach'
        
        return heatmap
    
    def _generate_architecture_heatmap_enhanced(self, azure_migrate_data: Any, questions_answers: List[QuestionAnswer], dependency_analysis: Any = None) -> List[Dict[str, str]]:
        """Generate enhanced architecture complexity heatmap using transcript Q&A, Azure Migrate data, and dependency analysis."""
        
        # Start with base analysis from transcript and Azure Migrate data
        heatmap = self._generate_architecture_heatmap(azure_migrate_data, questions_answers)
        
        # Enhance with dependency analysis if available
        if dependency_analysis:
            dependency_insights = self._analyze_dependency_complexity(dependency_analysis)
            
            # Update heatmap based on dependency analysis
            for insight in dependency_insights:
                # Find corresponding heatmap area
                matching_area = None
                for item in heatmap:
                    if insight['area'].lower() in item['area'].lower() or item['area'].lower() in insight['area'].lower():
                        matching_area = item
                        break
                
                if matching_area:
                    # Combine rankings (take higher complexity)
                    current_rank = matching_area['ranking']
                    new_rank = insight['ranking']
                    
                    rank_values = {'Low': 1, 'Medium': 2, 'High': 3}
                    if rank_values.get(new_rank, 1) > rank_values.get(current_rank, 1):
                        matching_area['ranking'] = new_rank
                        matching_area['notes'] = f"{matching_area['notes']} + {insight['notes']}"
                else:
                    # Add new complexity area from dependency analysis
                    heatmap.append(insight)
        
        return heatmap
    
    def _analyze_dependency_complexity(self, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Analyze dependency data to determine architecture complexity factors."""
        complexity_insights = []
        
        try:
            # Analyze network complexity
            if hasattr(dependency_analysis, 'network_segments') and dependency_analysis.network_segments:
                segment_count = len(dependency_analysis.network_segments)
                if segment_count > 3:
                    complexity_insights.append({
                        'area': 'Network Segmentation',
                        'ranking': 'High',
                        'notes': f'Complex network with {segment_count} segments requiring careful Azure VNET planning'
                    })
                elif segment_count > 1:
                    complexity_insights.append({
                        'area': 'Network Segmentation', 
                        'ranking': 'Medium',
                        'notes': f'Multi-segment network ({segment_count} segments) requiring subnet planning'
                    })
            
            # Analyze connection complexity
            if hasattr(dependency_analysis, 'connections') and dependency_analysis.connections:
                connection_count = len(dependency_analysis.connections)
                unique_protocols = set()
                unique_ports = set()
                
                for conn in dependency_analysis.connections:
                    if hasattr(conn, 'protocol') and conn.protocol:
                        unique_protocols.add(conn.protocol.upper())
                    if hasattr(conn, 'port') and conn.port:
                        unique_ports.add(conn.port)
                
                if connection_count > 50 or len(unique_protocols) > 5:
                    complexity_insights.append({
                        'area': 'Dependency Complexity',
                        'ranking': 'High', 
                        'notes': f'High dependency complexity: {connection_count} connections, {len(unique_protocols)} protocols'
                    })
                elif connection_count > 20 or len(unique_protocols) > 2:
                    complexity_insights.append({
                        'area': 'Dependency Complexity',
                        'ranking': 'Medium',
                        'notes': f'Moderate dependencies: {connection_count} connections, {len(unique_protocols)} protocols'
                    })
            
            # Analyze external dependency complexity
            if hasattr(dependency_analysis, 'external_dependencies') and dependency_analysis.external_dependencies:
                ext_count = len(dependency_analysis.external_dependencies)
                if ext_count > 10:
                    complexity_insights.append({
                        'area': 'External Integration',
                        'ranking': 'High',
                        'notes': f'High external dependency complexity: {ext_count} external systems'
                    })
                elif ext_count > 3:
                    complexity_insights.append({
                        'area': 'External Integration',
                        'ranking': 'Medium', 
                        'notes': f'Moderate external dependencies: {ext_count} external systems'
                    })
        
        except Exception as e:
            # Fallback in case of parsing issues
            complexity_insights.append({
                'area': 'Dependency Analysis',
                'ranking': 'Medium',
                'notes': f'Dependency data available but requires manual review: {str(e)}'
            })
        
        return complexity_insights
    
    def _generate_application_allocation(self, azure_migrate_data: Any) -> Dict[str, Any]:
        """Generate application allocation and scheduling information."""
        
        allocation = {
            'move_group': 'Wave 1 - Core Applications',
            'wave_allocation': 'Wave 1',
            'scheduling': 'Month 2-3',
            'migration_factory': 'Azure Migrate Service',
            'decisions': [
                {'area': 'Migration Tooling', 'decision': 'Azure Migrate + Azure Site Recovery'},
                {'area': 'Planning Dependencies', 'decision': 'Network and security configuration first'},
                {'area': 'Resource Allocation', 'decision': 'Dedicated migration team'},
                {'area': 'Testing Strategy', 'decision': 'Parallel testing environment'}
            ]
        }
        
        return allocation
    
    def _generate_supporting_documents(self) -> List[Dict[str, str]]:
        """Generate list of supporting documents."""
        
        documents = [
            {'artifact': 'Application Information Form', 'location': 'Generated from Q&A analysis'},
            {'artifact': 'Azure Migrate Assessment', 'location': 'Azure Migrate Portal'},
            {'artifact': 'Network Architecture Diagram', 'location': 'To be created'},
            {'artifact': 'Security Requirements Document', 'location': 'To be created'},
            {'artifact': 'Migration Plan Document', 'location': 'To be generated'}
        ]
        
        return documents
    
    def _replace_placeholders_in_template(self, doc: Document, assessment_data: AssessmentReportData):
        """Replace placeholders and template content in the document with actual findings."""
        
        # Define content mappings for different sections
        content_mappings = {
            # Application name replacement
            '<Application Name>': assessment_data.application_name,
            
            # External services and integrations
            '<External services, databases, or integrations with other systems per environment>': 
                self._format_external_services(assessment_data),
            
            # Security content
            '<Secrets Management, Certificates, store local password, Cryptographic Algorithm>':
                self._format_security_content(assessment_data.security_considerations),
            
            # Network content - multiple variations in template  
            '<VCI/SDWAN/Market, latency, bandwidth, protocols, middleware, encryption, load balancing>':
                self._format_network_content(assessment_data.network_requirements),
            '<VCI/SDWAN/Market, latency, bandwidth, protocols, middleware, encryption, load balancing, Iron Cloud>':
                self._format_network_content(assessment_data.network_requirements),
            
            # Identity provider content
            '< authentication/authorization for application / database / servers, AAA / SSO / AD /AAD / LDAP / SAML':
                self._format_identity_content(assessment_data.identity_providers),
            
            # Automation/CI-CD content
            '<Include details on how application and infrastructure are built and deployed including CI/CD':
                self._format_automation_content(assessment_data.automation_details),
            '<Include details on how application and infrastructure are built and deployed including CI/CD config':
                self._format_automation_content(assessment_data.automation_details),
            
            # Customer impact content
            '<Include details on what potential impact migration may have on internal as well as external interfaces>':
                self._format_customer_impact_content(assessment_data.customer_impact),
            '<Include details on what potential impact migration may have on internal as well as external interfa':
                self._format_customer_impact_content(assessment_data.customer_impact),
            
            # Operational concerns content
            '<Include details on current challenges, pain points application team is facing in operations/maintenance>':
                self._format_operational_concerns_content(assessment_data.operational_concerns),
            '<Include details on current challenges, pain points application team is facing in operations/mainten':
                self._format_operational_concerns_content(assessment_data.operational_concerns),
            '<Include details on current challenges, pain points application team is facing in opera':
                self._format_operational_concerns_content(assessment_data.operational_concerns),
            
            # Observability content
            'Monitoring <App/Infra/Security and solutions>':
                self._format_monitoring_content(assessment_data.observability),
            'Alerts: <metrics and thresholds, notification destination>':
                self._format_alerts_content(assessment_data.observability),
            'Events: <Event types and destination>':
                self._format_events_content(assessment_data.observability),
            
            # BCDR content
            'Regions / Locations / BCDR :<Primary site, DR site, RPO, RTO, SLA>':
                self._format_bcdr_content(assessment_data),
            
            # Migration acceptance tests
            '<Yes/No, if it needs to be created and type of tests required>':
                self._format_migration_tests_content(assessment_data),
            
            # Environment placeholders
            '<Environment>': 'Production',
            '< Production, Development, Pre-Production >': 'Production, Development, Pre-Production',
            '<Production, Development, Pre-Production>': 'Production, Development, Pre-Production',
            '< Production, Development, Pre-Production': 'Production, Development, Pre-Production',
            '< Environment>': 'Production',
            '<Environment> Logical Architecture': 'Production Logical Architecture',
            '< Environment> Logical Architecture': 'Production Logical Architecture',
            '<Environment> Application Network Flow': 'Production Application Network Flow',
            '< Environment> Application Network Flow': 'Production Application Network Flow',
            '<Environment> Proposed Architecture': 'Production Proposed Architecture',
            '< Environment> Proposed Architecture': 'Production Proposed Architecture',
            '<Environment> Application and Infrastructure RBAC': 'Production Application and Infrastructure RBAC',
            '<Environment> Azure Services RBAC': 'Production Azure Services RBAC',
            '<Environment> Azure Tagging': 'Production Azure Tagging',
            '<Environment> Source Delivery Information': 'Production Source Delivery Information',
            '<Environment> Target Delivery Information': 'Production Target Delivery Information',
            
            # Generic placeholders
            '<tag-name>': 'environment',
            '<value>': 'production',
            '<1 x IP Addresses for Application Tier, 2 x IP Addresses for Database Tier>': 
                '1 x IP Address for Application Tier, 1 x IP Address for Database Tier'
        }
        
        # Replace content in paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                new_text = paragraph.text
                
                # Apply all content mappings
                for placeholder, replacement in content_mappings.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, replacement)
                
                # Update paragraph text if changed
                if new_text != paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
    
    def _format_external_services(self, assessment_data: AssessmentReportData) -> str:
        """Format external services and integrations content."""
        external_items = []
        
        # Look for database information
        if hasattr(assessment_data, 'questions_answers'):
            for qa in assessment_data.questions_answers:
                if qa.is_answered and qa.answer != "Not addressed in transcript":
                    question_lower = qa.question.lower()
                    if any(keyword in question_lower for keyword in ['database', 'external', 'integration', 'service', 'api']):
                        external_items.append(f"• {qa.answer}")
        
        return '\n'.join(external_items) if external_items else "NA - No external services or integrations identified from current assessment"
    
    def _format_security_content(self, security_considerations: List[Dict]) -> str:
        """Format security considerations content."""
        if not security_considerations:
            return "NA - Security requirements to be determined during detailed assessment"
        
        security_items = []
        for item in security_considerations:
            security_items.append(f"• {item.get('requirement', 'Security requirement')}: {item.get('details', 'Details to be determined')}")
        
        return '\n'.join(security_items)
    
    def _format_network_content(self, network_requirements: List[Dict]) -> str:
        """Format network requirements content."""
        if not network_requirements:
            return "NA - Network requirements to be determined during detailed assessment"
        
        network_items = []
        for item in network_requirements:
            network_items.append(f"• {item.get('requirement', 'Network requirement')}: {item.get('details', 'Details to be determined')}")
        
        return '\n'.join(network_items)
    
    def _format_identity_content(self, identity_providers: List[Dict]) -> str:
        """Format identity provider content."""
        if not identity_providers:
            return "NA - Identity and authentication requirements to be determined during detailed assessment"
        
        identity_items = []
        for item in identity_providers:
            identity_items.append(f"• {item.get('provider', 'Identity provider')}: {item.get('details', 'Configuration to be determined')}")
        
        return '\n'.join(identity_items)
    
    def _format_automation_content(self, automation_details: List[Dict]) -> str:
        """Format automation and CI/CD content."""
        if not automation_details:
            return "NA - CI/CD and automation processes to be determined during detailed assessment"
        
        automation_items = []
        for item in automation_details:
            automation_items.append(f"• {item.get('process', 'Automation process')}: {item.get('current_state', 'Current state to be assessed')}")
        
        return '\n'.join(automation_items)
    
    def _format_customer_impact_content(self, customer_impact: List[Dict]) -> str:
        """Format customer impact content."""
        if not customer_impact:
            return "NA - Customer impact analysis to be conducted during migration planning phase"
        
        impact_items = []
        for item in customer_impact:
            impact_items.append(f"• {item.get('impact_area', 'Impact area')}: {item.get('description', 'Impact to be assessed')}")
        
        return '\n'.join(impact_items)
    
    def _format_migration_tests_content(self, assessment_data: AssessmentReportData) -> str:
        """Format migration acceptance tests content."""
        # Look for testing-related Q&A
        if hasattr(assessment_data, 'questions_answers'):
            for qa in assessment_data.questions_answers:
                if qa.is_answered and qa.answer != "Not addressed in transcript":
                    question_lower = qa.question.lower()
                    if any(keyword in question_lower for keyword in ['test', 'testing', 'validation', 'acceptance']):
                        return f"Yes - {qa.answer}"
        
        return "NA - Migration acceptance testing strategy to be defined during planning phase"
    
    def _format_operational_concerns_content(self, operational_concerns: List[Dict]) -> str:
        """Format operational concerns content."""
        if not operational_concerns:
            return "NA - Operational challenges and pain points to be identified during detailed assessment"
        
        concerns_items = []
        for item in operational_concerns:
            concerns_items.append(f"• {item.get('concern', 'Operational concern')}: {item.get('current_challenge', 'Challenge to be assessed')}")
        
        return '\n'.join(concerns_items)
    
    def _format_monitoring_content(self, observability: Dict[str, Any]) -> str:
        """Format monitoring content."""
        if observability.get('monitoring'):
            monitoring_info = '; '.join(observability['monitoring'][:2])
            return f"Monitoring: {monitoring_info}"
        else:
            return "Monitoring: Comprehensive observability strategy to be implemented with Azure Monitor"
    
    def _format_alerts_content(self, observability: Dict[str, Any]) -> str:
        """Format alerts content."""
        if observability.get('alerts'):
            alerts_info = '; '.join(observability['alerts'][:2])
            return f"Alerts: {alerts_info}"
        else:
            return "Alerts: Proactive alerting framework to be configured based on application requirements"
    
    def _format_events_content(self, observability: Dict[str, Any]) -> str:
        """Format events content."""
        if observability.get('events'):
            events_info = '; '.join(observability['events'][:2])
            return f"Events: {events_info}"
        else:
            return "Events: Event tracking and audit logging to be implemented per compliance requirements"
    
    def _format_bcdr_content(self, assessment_data: AssessmentReportData) -> str:
        """Format BCDR (Business Continuity Disaster Recovery) content with intelligent architecture recommendations."""
        
        # Analyze technology stack for BCDR strategy recommendations
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        deployment_method = self._analyze_deployment_method(assessment_data.questions_answers)
        
        # Look for BCDR-related Q&A first
        bcdr_info = []
        rpo_rto_info = None
        
        if hasattr(assessment_data, 'questions_answers'):
            for qa in assessment_data.questions_answers:
                if qa.is_answered and qa.answer != "Not addressed in transcript":
                    question_lower = qa.question.lower()
                    if any(keyword in question_lower for keyword in ['disaster', 'recovery', 'bcdr', 'backup', 'rpo', 'rto', 'availability']):
                        bcdr_info.append(f"• {qa.question}: {qa.answer}")
                        if any(term in question_lower for term in ['rpo', 'rto', 'availability']):
                            rpo_rto_info = qa.answer
        
        content = "**Business Continuity and Disaster Recovery Strategy:**\n\n"
        
        if bcdr_info:
            content += "From the assessment, the following BCDR requirements were identified:\n\n"
            content += "\n".join(bcdr_info[:3])
            content += "\n\n"
        else:
            content += "**BCDR Requirements**: N/A - Specific BCDR requirements not detailed in the transcript.\n\n"
        
        content += "**Recommended Azure BCDR Architecture:**\n\n"
        
        # Technology-specific BCDR recommendations
        if tech_stack['containers'] and deployment_method == 'kubernetes':
            content += """**Container-Based BCDR Strategy:**
• **Multi-Region AKS**: Deploy AKS clusters across multiple Azure regions
• **Container Registry Replication**: Geo-replicate container images for disaster recovery
• **Persistent Volume Backup**: Azure Backup for Kubernetes persistent volumes
• **Application State**: Implement stateless design with external state management
• **Traffic Routing**: Azure Traffic Manager for automatic failover between regions

"""
        elif tech_stack['containers']:
            content += """**Container Application BCDR Strategy:**
• **Multi-Region Deployment**: Deploy containers across multiple Azure regions
• **Container Image Backup**: Store container images in geo-replicated Azure Container Registry
• **Application Data Backup**: Azure Backup for application data and configuration
• **Load Balancer Failover**: Azure Application Gateway with backend pool failover
• **Automated Recovery**: Azure Site Recovery for container hosts

"""
        else:
            content += """**Virtual Machine BCDR Strategy:**
• **Azure Site Recovery**: Replicate VMs to secondary Azure region
• **VM Backup**: Azure Backup for virtual machine and application data protection
• **Network Replication**: Replicate virtual networks and security configurations
• **Load Balancer Configuration**: Multi-region load balancing with health probes
• **Manual/Automated Failover**: Configure based on RTO requirements

"""
        
        # Database-specific BCDR
        if tech_stack['databases']:
            content += "**Database BCDR Strategy:**\n"
            
            for db in tech_stack['databases']:
                if db == 'PostgreSQL':
                    content += """• **Azure Database for PostgreSQL**: 
  - Geo-redundant backup with point-in-time recovery
  - Read replicas in secondary regions for disaster recovery
  - Automatic failover with minimal downtime
"""
                elif db == 'Redis':
                    content += """• **Azure Cache for Redis**:
  - Zone redundancy for high availability within region
  - Geo-replication for cross-region disaster recovery
  - Data persistence options for recovery scenarios
"""
                elif db == 'MySQL':
                    content += """• **Azure Database for MySQL**:
  - Cross-region backup replication
  - Read replicas for disaster recovery
  - Automated backup with configurable retention
"""
            content += "\n"
        else:
            content += """**Database BCDR Strategy:**
• **Azure SQL Database**: Geo-replication with automatic failover groups
• **Point-in-Time Recovery**: Configurable backup retention (7-35 days)
• **Cross-Region Backup**: Geo-redundant backup storage
• **Read Replicas**: Secondary databases for disaster recovery scenarios

"""
        
        # Define recommended RTO/RPO based on application type
        if rpo_rto_info:
            content += f"**Recovery Objectives (from assessment):**\n• {rpo_rto_info}\n\n"
        else:
            content += """**Recommended Recovery Objectives:**
• **Recovery Time Objective (RTO)**: 4 hours - Maximum acceptable downtime
• **Recovery Point Objective (RPO)**: 1 hour - Maximum acceptable data loss
• **Availability Target**: 99.9% uptime - Approximately 8.76 hours downtime per year

"""
        
        content += "**Azure Native BCDR Services:**\n"
        content += """• **Azure Site Recovery**: Automated disaster recovery orchestration
• **Azure Backup**: Centralized backup management and monitoring
• **Azure Traffic Manager**: DNS-based traffic routing with health monitoring
• **Azure Monitor**: Continuous monitoring and alerting for BCDR events
• **Azure Resource Manager**: Infrastructure as Code for rapid environment recreation

**BCDR Implementation Phases:**
1. **Assessment**: Define RTO/RPO requirements and document dependencies
2. **Design**: Create multi-region architecture with appropriate Azure services
3. **Implementation**: Deploy BCDR infrastructure and configure replication
4. **Testing**: Regular disaster recovery drills and failover testing
5. **Documentation**: Maintain runbooks and escalation procedures
6. **Monitoring**: Continuous monitoring of backup and replication health"""
        
        return content
        
        # Update tables with actual data
        for table in doc.tables:
            self._update_table_with_data(table, assessment_data)
    
    def _create_embedded_template(self, assessment_data: AssessmentReportData) -> Document:
        """Create document with embedded template structure and populate with findings."""
        
        doc = Document()
        
        # Add title - clean application name without quotes
        clean_app_name = assessment_data.application_name.strip('"').strip("'").strip()
        title = doc.add_heading(clean_app_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('Application Assessment Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', 0)
        toc_content = self._generate_dynamic_toc(assessment_data.environments)
        
        toc_para = doc.add_paragraph(toc_content)
        toc_para.style = 'Normal'
        
        doc.add_page_break()
        
        # Introduction
        doc.add_heading('Introduction', 0)
        intro_content = self._format_introduction_content(assessment_data)
        self._add_formatted_paragraph(doc, intro_content)
        
        doc.add_page_break()
        
        # 1. Application Overview
        doc.add_heading('1	Application Overview', 0)
        
        # 1.1 Key Business Drivers
        doc.add_heading('1.1	Key Business Drivers', 1)
        business_drivers_content = self._format_business_drivers_content(assessment_data)
        self._add_formatted_paragraph(doc, business_drivers_content)
        
        # 1.2 Key Contacts
        doc.add_heading('1.2	Key Contacts', 1)
        contacts = self._extract_key_contacts(assessment_data.questions_answers)
        self._add_key_contacts_table(doc, contacts)
        
        # 1.3 Migration Strategy
        doc.add_heading('1.3	Migration Strategy', 1)
        
        # 1.3.1 Migration Pattern and Complexity
        doc.add_heading('1.3.1	Migration Pattern and Complexity', 2)
        pattern_content = self._format_migration_pattern_content(assessment_data)
        self._add_formatted_paragraph(doc, pattern_content)
        
        # 1.3.2 Technology Selection
        doc.add_heading('1.3.2	Technology Selection', 2)
        tech_content = self._format_technology_selection_content(assessment_data)
        self._add_formatted_paragraph(doc, tech_content)
        
        # 1.3.3 Indicative Azure Cost
        doc.add_heading('1.3.3	Indicative Azure Cost', 2)
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        architecture_type = self._analyze_architecture_type(assessment_data.questions_answers)
        deployment_method = self._analyze_deployment_method(assessment_data.questions_answers)
        migration_pattern = self._recommend_migration_pattern(tech_stack, architecture_type, deployment_method)
        azure_services = self._recommend_azure_services(tech_stack, migration_pattern['pattern'])
        
        # Extract cost info and build breakdown
        cost_info = []
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in ['cost', 'budget', 'price', 'estimate']):
                    cost_info.append(qa.answer)
        
        # Add introduction
        cost_para = doc.add_paragraph()
        cost_title = cost_para.add_run("Indicative Monthly Azure Costs:")
        cost_title.bold = True
        
        if cost_info:
            doc.add_paragraph("Based on the assessment discussion:")
            bullet_para = doc.add_paragraph(f"• {cost_info[0]}")
            bullet_para.style = 'List Bullet'
        
        estimated_para = doc.add_paragraph()
        estimated_title = estimated_para.add_run("Estimated costs based on recommended Azure services:")
        estimated_title.bold = True
        
        # Calculate costs and build cost breakdown
        total_min_cost, total_max_cost, cost_breakdown = self._calculate_azure_costs(azure_services, tech_stack, deployment_method)
        
        # Add the cost table
        self._add_cost_breakdown_table(doc, tech_stack, total_min_cost, total_max_cost, cost_breakdown)
        
        # Add cost optimization section
        self._add_cost_optimization_section(doc, migration_pattern, deployment_method)
        
        # 1.4 Database Information
        doc.add_heading('1.4	Database Information', 1)
        db_content = self._format_database_information_content(assessment_data)
        self._add_formatted_paragraph(doc, db_content)
        
        # 1.5 Macro Dependencies
        doc.add_heading('1.5	Macro Dependencies', 1)
        dependencies_content = self._format_macro_dependencies_content(assessment_data)
        self._add_formatted_paragraph(doc, dependencies_content)
        
        # 1.6 Security Considerations
        doc.add_heading('1.6	Security Considerations', 1)
        security_content = self._format_security_content(assessment_data.security_considerations)
        self._add_formatted_paragraph(doc, security_content)
        
        # 1.7 Resiliency Configuration  
        doc.add_heading('1.7	Resiliency Configuration', 1)
        bcdr_content = self._format_bcdr_content(assessment_data)
        self._add_formatted_paragraph(doc, bcdr_content)
        
        # 1.8 Network Access Requirements
        doc.add_heading('1.8	Network Access Requirements', 1)
        network_content = self._format_network_content(assessment_data.network_requirements)
        self._add_formatted_paragraph(doc, network_content)
        
        # 1.9 Identity Providers
        doc.add_heading('1.9	Identity Providers', 1)
        identity_content = self._format_identity_content(assessment_data.identity_providers)
        self._add_formatted_paragraph(doc, identity_content)
        
        # 1.10 Automation
        doc.add_heading('1.10	Automation', 1)
        automation_content = self._format_automation_content(assessment_data.automation_details)
        self._add_formatted_paragraph(doc, automation_content)
        
        # 1.11 Customer Impact
        doc.add_heading('1.11	Customer Impact', 1)
        impact_content = self._format_customer_impact_content(assessment_data.customer_impact)
        self._add_formatted_paragraph(doc, impact_content)
        
        # 1.12 Operational Concerns
        doc.add_heading('1.12	Operational Concerns', 1)
        operational_content = self._format_operational_concerns_content(assessment_data.operational_concerns)
        self._add_formatted_paragraph(doc, operational_content)
        
        # 1.13 Migration Acceptance Tests
        doc.add_heading('1.13	Migration Acceptance Tests', 1)
        tests_content = self._format_migration_tests_content(assessment_data)
        self._add_formatted_paragraph(doc, tests_content)
        
        # 1.14 Observability
        doc.add_heading('1.14	Observability', 1)
        monitoring_content = self._format_monitoring_content(assessment_data.observability)
        alerts_content = self._format_alerts_content(assessment_data.observability)
        events_content = self._format_events_content(assessment_data.observability)
        self._add_formatted_paragraph(doc, monitoring_content)
        self._add_formatted_paragraph(doc, alerts_content)
        self._add_formatted_paragraph(doc, events_content)
        
        # 2. Supporting Documents
        doc.add_heading('2	Supporting Documents', 0)
        doc.add_paragraph('The following table provides a summary of the supporting documents to support the planning and migration of the application.')
        
        # Supporting Documents Table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Artefact'
        header_cells[1].text = 'Information Location'
        
        for doc_item in assessment_data.supporting_documents:
            row_cells = table.add_row().cells
            row_cells[0].text = doc_item['artifact']
            row_cells[1].text = doc_item['location']
        
        # 3. Current Logical Architecture
        doc.add_heading('3	Current Logical Architecture', 0)
        doc.add_paragraph('The following section provides a view of the logical architecture of the application per environment.')
        
        # Add dynamic environment sections with specific content
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'3.{i+1}	{env} Logical Architecture', 1)
            doc.add_paragraph(f'The following provides the logical architecture view of the {env} environment.')
            
            # Add environment-specific content
            env_content = self._generate_environment_specific_content(env, "logical_architecture", assessment_data)
            self._add_formatted_paragraph(doc, env_content)
            
            doc.add_paragraph(f'Figure: {env} Current Logical View')
            doc.add_paragraph("")  # Add spacing
        
        # 4. Application Network Flow
        doc.add_heading('4	Application Network Flow', 0)
        doc.add_paragraph('The following section provides the details for the application network flow required by the application.')
        
        # Add network flow for each environment with detailed content
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'4.{i+1}	{env} Application Network Flow', 1)
            doc.add_paragraph(f'The following diagram provides the application network flow for the {env} environment.')
            
            # Add environment-specific network flow content
            network_content = self._generate_environment_specific_content(env, "network_flow", assessment_data)
            self._add_formatted_paragraph(doc, network_content)
            
            doc.add_paragraph(f'Figure: {env} Application Network Flow Diagram')
            
            # Add network flow table for the first environment (most detailed)
            if i == 0:
                flow_table = doc.add_table(rows=1, cols=2)
                flow_table.style = 'Table Grid'
                flow_header = flow_table.rows[0].cells
                flow_header[0].text = 'Step'
                flow_header[1].text = 'Details'
                
                # Add network flow details from findings
                network_steps = self._extract_network_flow_steps(assessment_data)
                for j, step in enumerate(network_steps, 1):
                    row_cells = flow_table.add_row().cells
                    row_cells[0].text = str(j)
                    row_cells[1].text = step
            
            doc.add_paragraph("")  # Add spacing
        
        # 5. Proposed Architecture in Azure
        doc.add_heading('5	Proposed Architecture in Azure', 0)
        doc.add_paragraph('The following section details the proposed architecture per environment of the application when being migrated to Azure.')
        
        # Add dynamic environment sections for proposed architecture with detailed content
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'5.{i+1}	{env} Proposed Architecture', 1)
            doc.add_paragraph(f'The following diagram represents the proposed architecture for the {env} environment.')
            
            # Add environment-specific proposed architecture content
            proposed_content = self._generate_environment_specific_content(env, "proposed_architecture", assessment_data)
            self._add_formatted_paragraph(doc, proposed_content)
            
            doc.add_paragraph(f'Figure: {env} Proposed Architecture Diagram')
            doc.add_paragraph("")  # Add spacing
        
        # 5.4 Low Level Design - Network Traffic Analysis
        doc.add_heading('5.4	Low Level Design - Network Traffic Analysis', 1)
        doc.add_paragraph('Based on the comprehensive analysis of network dependency data, the following low-level design recommendations provide detailed insights into the proposed Azure architecture.')
        
        # Add detailed network analysis content
        print(f"DEBUG: Checking target_architecture - hasattr: {hasattr(assessment_data, 'target_architecture')}")
        if hasattr(assessment_data, 'target_architecture'):
            print(f"DEBUG: target_architecture value: {assessment_data.target_architecture}")
            print(f"DEBUG: target_architecture type: {type(assessment_data.target_architecture)}")
        
        if hasattr(assessment_data, 'target_architecture') and assessment_data.target_architecture:
            try:
                print("DEBUG: About to generate network analysis content...")
                network_analysis_content = self._generate_comprehensive_network_analysis(assessment_data.target_architecture)
                print(f"DEBUG: Generated content length: {len(network_analysis_content) if network_analysis_content else 0}")
                print(f"DEBUG: Content preview: {network_analysis_content[:200] if network_analysis_content else 'None'}")
                self._add_formatted_paragraph(doc, network_analysis_content)
                print("DEBUG: Successfully added network analysis content to document")
            except Exception as e:
                print(f"Warning: Error generating network analysis content: {e}")
                print(f"DEBUG: Exception type: {type(e)}")
                import traceback
                traceback.print_exc()
                doc.add_paragraph("Network traffic analysis encountered an error. Low-level design will be based on application requirements and Azure best practices.")
        else:
            print("DEBUG: No target_architecture found - adding fallback message")
            doc.add_paragraph("Network traffic analysis data is not available. Low-level design will be based on application requirements and Azure best practices.")
        
        doc.add_paragraph("")  # Add spacing
        
        # 6. Architecture Heatmap
        doc.add_heading('6	Architecture Heatmap', 0)
        doc.add_paragraph('Architectural heatmap is a high-level ranking of key concerns that are relevant to application migration to Azure.')
        
        heatmap_table = doc.add_table(rows=1, cols=3)
        heatmap_table.style = 'Table Grid'
        heatmap_header = heatmap_table.rows[0].cells
        heatmap_header[0].text = 'Area'
        heatmap_header[1].text = 'Notes'
        heatmap_header[2].text = 'Ranking'
        
        for item in assessment_data.architecture_heatmap:
            row_cells = heatmap_table.add_row().cells
            row_cells[0].text = item['area']
            row_cells[1].text = item['notes']
            row_cells[2].text = item['ranking']
        
        # 7. Decision Matrix
        doc.add_heading('7	Decision Matrix', 0)
        self._add_decision_matrix_table(doc, assessment_data)
        
        # 8. Application Allocation and Scheduling
        doc.add_heading('8	Application Allocation and Scheduling', 0)
        doc.add_paragraph('The application allocation and scheduling cover the final decisions regarding the application to be migrated.')
        
        allocation_table = doc.add_table(rows=1, cols=4)
        allocation_table.style = 'Table Grid'
        allocation_header = allocation_table.rows[0].cells
        allocation_header[0].text = 'Move Group'
        allocation_header[1].text = 'Wave Allocation'
        allocation_header[2].text = 'Scheduling'
        allocation_header[3].text = 'Migration Factory'
        
        if assessment_data.application_allocation:
            allocation = assessment_data.application_allocation
            row_cells = allocation_table.add_row().cells
            row_cells[0].text = allocation.get('move_group', 'Wave 1 - Core Applications')
            row_cells[1].text = allocation.get('wave_allocation', 'Wave 1')
            row_cells[2].text = allocation.get('scheduling', 'Month 2-3')
            row_cells[3].text = allocation.get('migration_factory', 'Azure Migrate Service')
        
        # 9. Appendix
        doc.add_heading('9	Appendix', 0)
        
        # 9.1 Additional Backlog Items
        doc.add_heading('9.1	Additional Backlog Items', 1)
        doc.add_paragraph('List any additional work items that needs to be included to complete the migration')
        
        backlog_table = doc.add_table(rows=1, cols=2)
        backlog_table.style = 'Table Grid'
        backlog_header = backlog_table.rows[0].cells
        backlog_header[0].text = 'Area'
        backlog_header[1].text = 'Final Decision'
        
        backlog_items = self._extract_backlog_items(assessment_data)
        for item in backlog_items:
            row_cells = backlog_table.add_row().cells
            row_cells[0].text = item['area']
            row_cells[1].text = item['decision']
        
        # 9.2 Application and Infrastructure RBAC Information
        doc.add_heading('9.2	Application and Infrastructure RBAC Information', 1)
        doc.add_paragraph('The following tables provides the RBAC information for the application and infrastructure it\'s hosted on.')
        
        # Add dynamic environment sections for RBAC
        rbac_items = self._extract_rbac_information(assessment_data)
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'9.2.{i+1}	{env} Application and Infrastructure RBAC', 2)
            rbac_table = doc.add_table(rows=1, cols=3)
            rbac_table.style = 'Table Grid'
            rbac_header = rbac_table.rows[0].cells
            rbac_header[0].text = 'Areas'
            rbac_header[1].text = 'Role'
            rbac_header[2].text = 'Access List'
            
            for item in rbac_items:
                row_cells = rbac_table.add_row().cells
                row_cells[0].text = item['area']
                row_cells[1].text = item['role']
                row_cells[2].text = item['access']
        
        # 9.3 Azure Services RBAC Information
        doc.add_heading('9.3	Azure Services RBAC Information', 1)
        doc.add_paragraph('The following tables provides the Azure RBAC information for the Azure services to be configured when hosting the application.')
        
        # Add dynamic environment sections for Azure Services RBAC
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'9.3.{i+1}	{env} Azure Services RBAC', 2)
            azure_rbac_table = doc.add_table(rows=1, cols=5)
            azure_rbac_table.style = 'Table Grid'
            azure_rbac_header = azure_rbac_table.rows[0].cells
            azure_rbac_header[0].text = 'Name'
            azure_rbac_header[1].text = 'User ID'
            azure_rbac_header[2].text = 'User Email address'
            azure_rbac_header[3].text = 'Access Type'
            azure_rbac_header[4].text = 'Roles'
            
            # Add placeholder RBAC entries for each environment
            for _ in range(3):
                row_cells = azure_rbac_table.add_row().cells
                row_cells[0].text = 'To be determined'
                row_cells[1].text = 'TBD'
                row_cells[2].text = 'TBD'
                row_cells[3].text = 'Reader Access'
                row_cells[4].text = 'Application / Infra/ Testing'
        
        # 9.4 Azure Tagging
        doc.add_heading('9.4	Azure Tagging', 1)
        doc.add_paragraph('The following tables provides the Azure tagging information to be used when applying the Azure Tags to the application components.')
        
        # Add dynamic environment sections for Azure Tagging
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'9.4.{i+1}	{env} Azure Tagging', 2)
            tagging_table = doc.add_table(rows=1, cols=4)
            tagging_table.style = 'Table Grid'
            tagging_header = tagging_table.rows[0].cells
            tagging_header[0].text = 'Tag Name'
            tagging_header[1].text = 'Type'
            tagging_header[2].text = 'Description'
            tagging_header[3].text = 'Value'
            
            # Add environment-specific tagging information
            tag_row = tagging_table.add_row().cells
            tag_row[0].text = 'environment'
            tag_row[1].text = 'Free text (3-15 char)'
            tag_row[2].text = 'Cost allocation and reporting.'
            tag_row[3].text = env.lower()
        
        # 9.5 Source Migration Delivery Information
        doc.add_heading('9.5	Source Migration Delivery Information', 1)
        doc.add_paragraph('The following tables provide the source migration delivery information to support the migration per environment.')
        
        # Add dynamic environment sections for Source Migration with intelligent content
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'9.5.{i+1}	{env} Source Delivery Information', 2)
            source_table = doc.add_table(rows=1, cols=2)
            source_table.style = 'Table Grid'
            source_header = source_table.rows[0].cells
            source_header[0].text = 'Requirements'
            source_header[1].text = 'Comments'
            
            # Generate intelligent source delivery requirements
            source_requirements = self._generate_source_delivery_requirements(env, assessment_data)
            for req_name, req_details in source_requirements.items():
                row_cells = source_table.add_row().cells
                row_cells[0].text = req_name
                row_cells[1].text = req_details
        
        # 9.6 Target Migration Delivery Information
        doc.add_heading('9.6	Target Migration Delivery Information', 1)
        doc.add_paragraph('The following tables provide the target migration delivery information to support the migration per environment.')
        
        # Add dynamic environment sections for Target Migration with intelligent content
        for i, env in enumerate(assessment_data.environments):
            doc.add_heading(f'9.6.{i+1}	{env} Target Delivery Information', 2)
            target_table = doc.add_table(rows=1, cols=2)
            target_table.style = 'Table Grid'
            target_header = target_table.rows[0].cells
            target_header[0].text = 'Requirements'
            target_header[1].text = 'Comments'
            
            # Generate intelligent target delivery requirements
            target_requirements = self._generate_target_delivery_requirements(env, assessment_data)
            for req_name, req_details in target_requirements.items():
                row_cells = target_table.add_row().cells
                row_cells[0].text = req_name
                row_cells[1].text = req_details
        
        return doc
    
    def _extract_network_flow_steps(self, assessment_data: AssessmentReportData) -> List[str]:
        """Extract network flow steps from assessment data."""
        steps = []
        
        # Look for network-related Q&A
        if hasattr(assessment_data, 'questions_answers'):
            for qa in assessment_data.questions_answers:
                if qa.is_answered and qa.answer != "Not addressed in transcript":
                    question_lower = qa.question.lower()
                    if any(keyword in question_lower for keyword in ['network', 'flow', 'connection', 'port', 'protocol']):
                        steps.append(qa.answer)
        
        # Add default steps if none found
        if not steps:
            steps = [
                'User accesses application through load balancer',
                'Load balancer distributes traffic to application tier',
                'Application tier connects to database tier',
                'Data flows between application and database'
            ]
        
        return steps[:10]  # Limit to 10 steps
    
    def _extract_backlog_items(self, assessment_data: AssessmentReportData) -> List[Dict[str, str]]:
        """Extract backlog items from assessment data."""
        backlog_items = [
            {'area': 'Migration Tooling', 'decision': 'Azure Migrate + Azure Site Recovery'},
            {'area': 'Planning Dependencies', 'decision': 'Network and security configuration first'},
            {'area': 'Resource Allocation', 'decision': 'Dedicated migration team'},
            {'area': 'Testing Strategy', 'decision': 'Parallel testing environment'}
        ]
        
        # Look for additional items from Q&A
        if hasattr(assessment_data, 'questions_answers'):
            for qa in assessment_data.questions_answers:
                if qa.is_answered and qa.answer != "Not addressed in transcript":
                    question_lower = qa.question.lower()
                    if any(keyword in question_lower for keyword in ['backlog', 'work item', 'todo', 'additional']):
                        backlog_items.append({
                            'area': 'Additional Work Item',
                            'decision': qa.answer[:100] + '...' if len(qa.answer) > 100 else qa.answer
                        })
        
        return backlog_items
    
    def _extract_rbac_information(self, assessment_data: AssessmentReportData) -> List[Dict[str, str]]:
        """Extract RBAC information from assessment data."""
        rbac_items = [
            {'area': 'Application', 'role': 'Administrator', 'access': f'{assessment_data.application_name} Admins'},
            {'area': 'Application', 'role': 'User', 'access': f'{assessment_data.application_name} Users'},
            {'area': 'Infrastructure', 'role': 'Administrator', 'access': 'Infrastructure Admins'},
            {'area': 'Database', 'role': 'Administrator', 'access': 'Database Admins'}
        ]
        
        return rbac_items
    
    def _update_table_with_data(self, table, assessment_data: AssessmentReportData):
        """Update table contents with assessment data."""
        
        try:
            if len(table.rows) == 0:
                return
                
            # Get header row to identify table type
            header_text = ' '.join([cell.text.strip().lower() for cell in table.rows[0].cells])
            
            # Supporting documents table
            if 'artefact' in header_text and 'information location' in header_text:
                # Clear existing rows except header
                for i in range(len(table.rows) - 1, 0, -1):
                    table._element.remove(table.rows[i]._element)
                
                # Add assessment data
                for doc in assessment_data.supporting_documents:
                    row = table.add_row()
                    row.cells[0].text = doc['artifact']
                    row.cells[1].text = doc['location']
            
            # Architecture heatmap table
            elif 'area' in header_text and 'ranking' in header_text:
                # Clear existing rows except header
                for i in range(len(table.rows) - 1, 0, -1):
                    table._element.remove(table.rows[i]._element)
                
                # Add heatmap data
                for item in assessment_data.architecture_heatmap:
                    row = table.add_row()
                    row.cells[0].text = item['area']
                    if len(row.cells) > 1:
                        row.cells[1].text = item['notes']
                    if len(row.cells) > 2:
                        row.cells[2].text = item['ranking']
            
            # Application allocation table
            elif 'move group' in header_text and 'wave allocation' in header_text:
                if len(table.rows) > 1 and assessment_data.application_allocation:
                    allocation = assessment_data.application_allocation
                    row = table.rows[1]  # First data row
                    if len(row.cells) > 0:
                        row.cells[0].text = allocation.get('move_group', 'Wave 1 - Core Applications')
                    if len(row.cells) > 1:
                        row.cells[1].text = allocation.get('wave_allocation', 'Wave 1')
                    if len(row.cells) > 2:
                        row.cells[2].text = allocation.get('scheduling', 'Month 2-3')
                    if len(row.cells) > 3:
                        row.cells[3].text = allocation.get('migration_factory', 'Azure Migrate Service')
            
            # Application allocation decisions table
            elif 'area' in header_text and 'final decision' in header_text:
                if assessment_data.application_allocation and 'decisions' in assessment_data.application_allocation:
                    # Clear existing rows except header
                    for i in range(len(table.rows) - 1, 0, -1):
                        table._element.remove(table.rows[i]._element)
                    
                    # Add decision data
                    for decision in assessment_data.application_allocation['decisions']:
                        row = table.add_row()
                        row.cells[0].text = decision['area']
                        row.cells[1].text = decision['decision']
            
            # Network flow details table (generic pattern)
            elif 'step' in header_text and 'details' in header_text:
                if assessment_data.network_requirements:
                    # Clear existing rows except header
                    for i in range(len(table.rows) - 1, 0, -1):
                        table._element.remove(table.rows[i]._element)
                    
                    # Add network requirement steps
                    for idx, req in enumerate(assessment_data.network_requirements[:5], 1):
                        row = table.add_row()
                        row.cells[0].text = str(idx)
                        row.cells[1].text = req['details'][:100] + "..." if len(req['details']) > 100 else req['details']
            
            # RBAC information table
            elif 'areas' in header_text and 'role' in header_text and 'access list' in header_text:
                # Clear existing rows except header
                for i in range(len(table.rows) - 1, 0, -1):
                    table._element.remove(table.rows[i]._element)
                
                # Add generic RBAC entries based on application
                rbac_entries = [
                    {'area': 'Application', 'role': 'Administrator', 'access': f'{assessment_data.application_name} Admins'},
                    {'area': 'Application', 'role': 'User', 'access': f'{assessment_data.application_name} Users'},
                    {'area': 'Infrastructure', 'role': 'Administrator', 'access': 'Infrastructure Admins'},
                    {'area': 'Database', 'role': 'Administrator', 'access': 'Database Admins'}
                ]
                
                for entry in rbac_entries:
                    row = table.add_row()
                    row.cells[0].text = entry['area']
                    row.cells[1].text = entry['role']
                    row.cells[2].text = entry['access']
                        
        except Exception as e:
            print(f"Warning: Could not update table: {e}")
            # Continue processing other tables
    
    def _create_assessment_document(self, doc: Document, assessment_data: AssessmentReportData):
        """Create assessment document from scratch if no template is available."""
        
        # Title
        title = doc.add_heading(f'{assessment_data.application_name} - Application Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            f'This document provides a comprehensive assessment of the {assessment_data.application_name} '
            'application for migration to Microsoft Azure. The assessment covers current architecture, '
            'security requirements, network dependencies, and migration recommendations.'
        )
        
        # Security Considerations
        if assessment_data.security_considerations:
            doc.add_heading('Security Considerations', level=1)
            for item in assessment_data.security_considerations:
                doc.add_paragraph(f"• {item['requirement']}: {item['details']}", style='List Bullet')
        
        # Network Requirements
        if assessment_data.network_requirements:
            doc.add_heading('Network Access Requirements', level=1)
            for item in assessment_data.network_requirements:
                doc.add_paragraph(f"• {item['requirement']}: {item['details']}", style='List Bullet')
        
        # Add more sections as needed...
        
        # Architecture Heatmap
        if assessment_data.architecture_heatmap:
            doc.add_heading('Architecture Heatmap', level=1)
            
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Area'
            header_cells[1].text = 'Notes'
            header_cells[2].text = 'Ranking'
            
            # Data rows
            for item in assessment_data.architecture_heatmap:
                row_cells = table.add_row().cells
                row_cells[0].text = item['area']
                row_cells[1].text = item['notes']
                row_cells[2].text = item['ranking']
    
    def _format_introduction_content(self, assessment_data: AssessmentReportData) -> str:
        """Format introduction content from assessment data."""
        return f"""This Application Assessment Report for {assessment_data.application_name} provides a comprehensive analysis of the current application architecture, requirements, and recommendations for migration to Microsoft Azure. 

The assessment has been conducted based on customer interviews, technical documentation review, and application analysis. This document serves as the foundation for migration planning and Azure architecture design.

The key areas covered in this assessment include:
• Application overview and business drivers
• Current architecture and dependencies
• Security and compliance requirements
• Network access patterns
• Migration strategy and Azure service recommendations
• Risk assessment and mitigation strategies"""
    
    def _format_business_drivers_content(self, assessment_data: AssessmentReportData) -> str:
        """Format business drivers content using LLM analysis of Q&A data."""
        
        if not self.llm_client:
            return self._basic_business_drivers_extraction(assessment_data)
        
        # Compile all answered questions for LLM analysis
        qa_context = []
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer not in ["Not addressed in transcript", "Error in analysis"]:
                qa_context.append(f"Q: {qa.question}\nA: {qa.answer}")
        
        if not qa_context:
            return self._basic_business_drivers_extraction(assessment_data)
        
        context_text = "\n\n".join(qa_context)
        
        prompt = f"""Analyze the following Q&A conversation to extract and identify key business drivers for an Azure cloud migration.

Q&A Context:
{context_text}

Based on this conversation, identify the primary business drivers, motivations, and benefits for migrating to Azure. Look for:
- Cost considerations and optimization goals
- Security and compliance requirements  
- Scalability and performance needs
- Modernization objectives
- Operational efficiency goals
- Business continuity and disaster recovery needs
- Strategic technology initiatives

Format your response as a numbered list of the top 5 business drivers, each with a brief explanation:

1. [Driver Name] - [Brief explanation based on the conversation]
2. [Driver Name] - [Brief explanation based on the conversation]  
3. [Driver Name] - [Brief explanation based on the conversation]
4. [Driver Name] - [Brief explanation based on the conversation]
5. [Driver Name] - [Brief explanation based on the conversation]

Focus on drivers that are explicitly mentioned or strongly implied in the conversation. If specific drivers aren't clear, infer logical drivers based on the technology and business context discussed."""

        try:
            response = self.llm_client.invoke(prompt)
            content = response.content.strip()
            
            if content and len(content) > 50:
                return f"The key business drivers for this migration include:\n\n{content}"
            else:
                print("Warning: LLM business drivers analysis returned insufficient content")
                return self._basic_business_drivers_extraction(assessment_data)
                
        except Exception as e:
            print(f"Warning: LLM business drivers analysis failed: {str(e)}")
            return self._basic_business_drivers_extraction(assessment_data)
    
    def _basic_business_drivers_extraction(self, assessment_data: AssessmentReportData) -> str:
        """Fallback basic business drivers extraction."""
        drivers = []
        
        # Extract business drivers from Q&A
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in ['business', 'driver', 'reason', 'motivation', 'benefit']):
                    drivers.append(qa.answer)
        
        if drivers:
            content = "The key business drivers for this migration include:\n\n"
            for i, driver in enumerate(drivers[:5], 1):  # Limit to 5 drivers
                content += f"{i}. {driver}\n"
        else:
            content = """The primary business drivers for this Azure migration include:

1. Cost Optimization - Reduce infrastructure and operational costs through Azure's pay-as-you-go model
2. Enhanced Security - Leverage Azure's advanced security capabilities and compliance certifications  
3. Improved Scalability - Enable automatic scaling to handle varying workloads
4. Modernization - Update legacy systems and adopt cloud-native technologies
5. Business Continuity - Implement robust disaster recovery and backup solutions"""
        
        return content
    
    def _format_key_contacts_content(self, assessment_data: AssessmentReportData) -> str:
        """Format key contacts content."""
        # Try to extract contact information from Q&A
        contacts = []
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in ['contact', 'owner', 'responsible', 'team', 'manager']):
                    contacts.append(qa.answer)
        
        if contacts:
            content = "Key project contacts identified:\n\n"
            for contact in contacts[:3]:  # Limit to 3 contacts
                content += f"• {contact}\n"
        else:
            content = """The following table identifies the key contacts for this migration project:

| Role | Name | Email | Responsibilities |
|------|------|-------|------------------|
| Project Sponsor | TBD | TBD | Executive oversight and approval |
| Technical Lead | TBD | TBD | Technical architecture and decisions |
| Application Owner | TBD | TBD | Business requirements and acceptance |
| Infrastructure Team | TBD | TBD | Environment setup and maintenance |

Note: Contact details to be updated during project initiation."""
        
        return content
    
    def _format_migration_pattern_content(self, assessment_data: AssessmentReportData) -> str:
        """Format migration pattern and complexity content with intelligent architectural recommendations."""
        
        # Analyze technology stack and architecture from Q&A
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        architecture_type = self._analyze_architecture_type(assessment_data.questions_answers)
        deployment_method = self._analyze_deployment_method(assessment_data.questions_answers)
        
        # Determine optimal migration pattern based on analysis
        migration_pattern = self._recommend_migration_pattern(tech_stack, architecture_type, deployment_method)
        complexity_assessment = self._assess_migration_complexity(tech_stack, architecture_type, deployment_method)
        
        content = f"""Based on the application assessment, the recommended migration approach is:

**{migration_pattern['pattern']}** - {migration_pattern['description']}

**Rationale:**
{migration_pattern['rationale']}

**Complexity Assessment: {complexity_assessment['level']}**
{complexity_assessment['description']}

**Key Migration Considerations:**
{chr(10).join([f"• {consideration}" for consideration in migration_pattern['considerations']])}

**Recommended Migration Phases:**
{chr(10).join([f"{i+1}. {phase}" for i, phase in enumerate(migration_pattern['phases'])])}"""
        
        return content
    
    def _analyze_technology_stack(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Analyze technology stack from Q&A responses using LLM analysis."""
        
        if not self.llm_client:
            # Fallback to basic extraction if no LLM available
            return self._basic_technology_extraction(questions_answers)
        
        # Compile all answered questions for LLM analysis
        qa_context = []
        for qa in questions_answers:
            if qa.is_answered and qa.answer not in ["Not addressed in transcript", "Error in analysis"]:
                qa_context.append(f"Q: {qa.question}\nA: {qa.answer}")
        
        if not qa_context:
            return {
                'languages': [],
                'frameworks': [],
                'databases': [],
                'infrastructure': [],
                'containers': False,
                'cloud_ready': False,
                'qa_context': questions_answers
            }
        
        context_text = "\n\n".join(qa_context)
        
        prompt = f"""Analyze the following Q&A conversation to extract detailed technology stack information for an Azure migration assessment.

Q&A Context:
{context_text}

Please analyze and provide a comprehensive technology stack assessment in the following JSON format:

{{
    "languages": ["list of programming languages mentioned"],
    "frameworks": ["list of frameworks and libraries mentioned"],
    "databases": ["list of database technologies mentioned"],
    "infrastructure": ["list of infrastructure components mentioned"],
    "containers": true/false,
    "cloud_ready": true/false,
    "architecture_type": "monolithic/n-tier/microservices",
    "deployment_method": "traditional/containerized/kubernetes",
    "modernization_level": "legacy/modern/cloud-native",
    "complexity_factors": ["list of factors that increase migration complexity"],
    "migration_readiness": "low/medium/high"
}}

Focus on:
- Programming languages and versions
- Web frameworks and application servers
- Database systems and versions
- Container technologies (Docker, Kubernetes)
- Load balancers, web servers, middleware
- Development and deployment tools
- Cloud readiness indicators
- Architecture patterns mentioned

Be specific and only include technologies that are clearly mentioned or strongly implied in the conversation."""

        try:
            response = self.llm_client.invoke(prompt)
            import json
            import re
            
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                tech_analysis = json.loads(json_match.group())
                
                # Add the Q&A context for later use
                tech_analysis['qa_context'] = questions_answers
                
                return tech_analysis
            else:
                print("Warning: Could not parse LLM technology analysis response")
                return self._basic_technology_extraction(questions_answers)
                
        except Exception as e:
            print(f"Warning: LLM technology analysis failed: {str(e)}")
            return self._basic_technology_extraction(questions_answers)
    
    def _basic_technology_extraction(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Fallback basic technology extraction when LLM is not available."""
        tech_stack = {
            'languages': [],
            'frameworks': [],
            'databases': [],
            'infrastructure': [],
            'containers': False,
            'cloud_ready': False,
            'qa_context': questions_answers
        }
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                answer_lower = qa.answer.lower()
                question_lower = qa.question.lower()
                combined_text = f"{question_lower} {answer_lower}"
                
                # Detect programming languages
                if any(lang in combined_text for lang in ['python', 'java', '.net', 'c#', 'node', 'javascript']):
                    if 'python' in combined_text:
                        tech_stack['languages'].append('Python')
                    if 'java' in combined_text:
                        tech_stack['languages'].append('Java')
                    if '.net' in combined_text or 'c#' in combined_text:
                        tech_stack['languages'].append('.NET')
                
                # Detect frameworks
                if any(fw in combined_text for fw in ['django', 'flask', 'spring', 'express', 'react', 'angular']):
                    if 'django' in combined_text:
                        tech_stack['frameworks'].append('Django')
                        tech_stack['cloud_ready'] = True  # Django is cloud-friendly
                    if 'flask' in combined_text:
                        tech_stack['frameworks'].append('Flask')
                    if 'spring' in combined_text:
                        tech_stack['frameworks'].append('Spring')
                
                # Detect databases
                if any(db in combined_text for db in ['postgres', 'mysql', 'oracle', 'sql server', 'redis', 'mongo']):
                    if 'postgres' in combined_text and 'PostgreSQL' not in tech_stack['databases']:
                        tech_stack['databases'].append('PostgreSQL')
                    if 'redis' in combined_text and 'Redis' not in tech_stack['databases']:
                        tech_stack['databases'].append('Redis')
                    if 'mysql' in combined_text and 'MySQL' not in tech_stack['databases']:
                        tech_stack['databases'].append('MySQL')
                
                # Detect containerization
                if any(container in combined_text for container in ['docker', 'kubernetes', 'k8s', 'pod', 'container']):
                    tech_stack['containers'] = True
                    tech_stack['cloud_ready'] = True
                
                # Detect infrastructure
                if any(infra in combined_text for infra in ['nginx', 'apache', 'load balancer', 'statefulset']):
                    if 'nginx' in combined_text and 'Nginx' not in tech_stack['infrastructure']:
                        tech_stack['infrastructure'].append('Nginx')
                    if 'load balancer' in combined_text and 'Load Balancer' not in tech_stack['infrastructure']:
                        tech_stack['infrastructure'].append('Load Balancer')
        
        return tech_stack
    
    def _analyze_architecture_type(self, questions_answers: List[QuestionAnswer]) -> str:
        """Analyze application architecture type using LLM."""
        
        if not self.llm_client:
            return self._basic_architecture_analysis(questions_answers)
        
        # Compile context for LLM
        qa_context = []
        for qa in questions_answers:
            if qa.is_answered and qa.answer not in ["Not addressed in transcript", "Error in analysis"]:
                qa_context.append(f"Q: {qa.question}\nA: {qa.answer}")
        
        if not qa_context:
            return 'n-tier'  # Default assumption
        
        context_text = "\n\n".join(qa_context)
        
        prompt = f"""Analyze the following Q&A conversation to determine the application architecture type for Azure migration planning.

Q&A Context:
{context_text}

Based on the conversation, determine the most likely application architecture type from these options:
- monolithic: Single deployable unit, all components tightly integrated
- n-tier: Traditional multi-layer architecture (presentation, business, data layers)
- microservices: Multiple independent services communicating via APIs
- distributed: Multiple components but not necessarily microservices
- legacy: Older architecture patterns, mainframe or monolithic systems

Respond with only one word from the options above, choosing the best match based on:
- How components are described
- Deployment patterns mentioned
- Service boundaries discussed
- Integration patterns
- Scalability requirements

Architecture type:"""

        try:
            response = self.llm_client.invoke(prompt)
            arch_type = response.content.strip().lower()
            
            # Validate response
            valid_types = ['monolithic', 'n-tier', 'microservices', 'distributed', 'legacy']
            if arch_type in valid_types:
                return arch_type
            else:
                print(f"Warning: Invalid architecture type from LLM: {arch_type}")
                return self._basic_architecture_analysis(questions_answers)
                
        except Exception as e:
            print(f"Warning: LLM architecture analysis failed: {str(e)}")
            return self._basic_architecture_analysis(questions_answers)
    
    def _basic_architecture_analysis(self, questions_answers: List[QuestionAnswer]) -> str:
        """Fallback basic architecture analysis."""
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                combined_text = f"{qa.question.lower()} {qa.answer.lower()}"
                
                if 'microservice' in combined_text:
                    return 'microservices'
                elif any(tier in combined_text for tier in ['3 tier', 'three tier', 'n-tier']):
                    return 'n-tier'
                elif 'monolith' in combined_text:
                    return 'monolithic'
        
        # Default assumption based on common patterns
        return 'n-tier'
    
    def _analyze_deployment_method(self, questions_answers: List[QuestionAnswer]) -> str:
        """Analyze current deployment method using LLM."""
        
        if not self.llm_client:
            return self._basic_deployment_analysis(questions_answers)
        
        # Compile context for LLM
        qa_context = []
        for qa in questions_answers:
            if qa.is_answered and qa.answer not in ["Not addressed in transcript", "Error in analysis"]:
                qa_context.append(f"Q: {qa.question}\nA: {qa.answer}")
        
        if not qa_context:
            return 'traditional'
        
        context_text = "\n\n".join(qa_context)
        
        prompt = f"""Analyze the following Q&A conversation to determine the current deployment method for Azure migration planning.

Q&A Context:
{context_text}

Based on the conversation, determine the current deployment method from these options:
- traditional: VM-based deployment, manual processes, traditional server setup
- containerized: Uses Docker containers but not necessarily orchestrated
- kubernetes: Uses Kubernetes or similar container orchestration
- serverless: Uses serverless functions or similar event-driven architecture
- hybrid: Mix of different deployment methods

Respond with only one word from the options above, choosing the best match based on:
- Infrastructure mentioned (VMs, containers, Kubernetes, etc.)
- Deployment processes described
- Orchestration tools mentioned
- Platform services used

Deployment method:"""

        try:
            response = self.llm_client.invoke(prompt)
            deployment_method = response.content.strip().lower()
            
            # Validate response
            valid_methods = ['traditional', 'containerized', 'kubernetes', 'serverless', 'hybrid']
            if deployment_method in valid_methods:
                return deployment_method
            else:
                print(f"Warning: Invalid deployment method from LLM: {deployment_method}")
                return self._basic_deployment_analysis(questions_answers)
                
        except Exception as e:
            print(f"Warning: LLM deployment analysis failed: {str(e)}")
            return self._basic_deployment_analysis(questions_answers)
    
    def _basic_deployment_analysis(self, questions_answers: List[QuestionAnswer]) -> str:
        """Fallback basic deployment analysis."""
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                combined_text = f"{qa.question.lower()} {qa.answer.lower()}"
                
                if any(k8s in combined_text for k8s in ['kubernetes', 'k8s', 'pod', 'namespace']):
                    return 'kubernetes'
                elif 'docker' in combined_text:
                    return 'containers'
                elif any(vm in combined_text for vm in ['virtual machine', 'vm', 'server']):
                    return 'virtual_machines'
        
        return 'traditional'
    
    def _recommend_migration_pattern(self, tech_stack: Dict[str, Any], architecture_type: str, deployment_method: str) -> Dict[str, Any]:
        """Recommend optimal migration pattern using LLM analysis."""
        
        if not self.llm_client:
            return self._basic_migration_pattern_recommendation(tech_stack, architecture_type, deployment_method)
        
        # Prepare comprehensive context for LLM
        tech_summary = {
            'languages': tech_stack.get('languages', []),
            'frameworks': tech_stack.get('frameworks', []),
            'databases': tech_stack.get('databases', []),
            'infrastructure': tech_stack.get('infrastructure', []),
            'containers': tech_stack.get('containers', False),
            'cloud_ready': tech_stack.get('cloud_ready', False),
            'architecture_type': architecture_type,
            'deployment_method': deployment_method,
            'modernization_level': tech_stack.get('modernization_level', 'unknown'),
            'migration_readiness': tech_stack.get('migration_readiness', 'unknown')
        }
        
        prompt = f"""Based on the following technology analysis, recommend the optimal Azure migration pattern and strategy.

Current Technology Stack Analysis:
{tech_summary}

Please provide a comprehensive migration recommendation in the following JSON format:

{{
    "pattern": "Migration pattern name (e.g., Rehost, Replatform, Refactor, etc.)",
    "description": "Brief description of the recommended approach",
    "rationale": "Detailed explanation of why this pattern is optimal for this technology stack",
    "considerations": [
        "Key technical consideration 1",
        "Key technical consideration 2",
        "Key technical consideration 3"
    ],
    "phases": [
        "Phase 1 description",
        "Phase 2 description", 
        "Phase 3 description"
    ],
    "complexity_level": "Low/Medium/High",
    "estimated_timeline": "timeframe estimate",
    "azure_services": [
        "Recommended Azure service 1",
        "Recommended Azure service 2"
    ],
    "risks": [
        "Migration risk 1",
        "Migration risk 2"
    ],
    "benefits": [
        "Migration benefit 1", 
        "Migration benefit 2"
    ]
}}

Consider these migration patterns:
- Rehost (Lift & Shift): Minimal changes, move as-is to Azure VMs
- Replatform: Some optimization for cloud, use managed services 
- Refactor: Significant changes to leverage cloud-native features
- Rebuild: Complete rewrite using cloud-native technologies
- Replace: Use SaaS alternatives

Choose the pattern that best balances speed, cost, and long-term benefits for this specific technology stack."""

        try:
            response = self.llm_client.invoke(prompt)
            import json
            import re
            
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                migration_recommendation = json.loads(json_match.group())
                return migration_recommendation
            else:
                print("Warning: Could not parse LLM migration pattern response")
                return self._basic_migration_pattern_recommendation(tech_stack, architecture_type, deployment_method)
                
        except Exception as e:
            print(f"Warning: LLM migration pattern analysis failed: {str(e)}")
            return self._basic_migration_pattern_recommendation(tech_stack, architecture_type, deployment_method)
    
    def _basic_migration_pattern_recommendation(self, tech_stack: Dict[str, Any], architecture_type: str, deployment_method: str) -> Dict[str, Any]:
        """Fallback basic migration pattern recommendation."""
        
        # If already containerized and cloud-ready
        if tech_stack['containers'] and tech_stack['cloud_ready']:
            return {
                'pattern': 'Replatform (Containerized)',
                'description': 'Migrate existing containerized application to Azure Container Services with minimal changes',
                'rationale': 'Application is already containerized and uses cloud-friendly technologies, making replatforming the optimal approach for faster migration with immediate cloud benefits.',
                'considerations': [
                    'Existing container images can be reused with minimal modifications',
                    'Kubernetes configuration needs Azure-specific adaptations',
                    'Network and storage configurations require Azure integration',
                    'CI/CD pipelines need to be updated for Azure deployments',
                    'Monitoring and logging integration with Azure services'
                ],
                'phases': [
                    'Assessment and Azure service mapping',
                    'Container registry setup and image migration',
                    'Azure Kubernetes Service (AKS) cluster preparation',
                    'Network and security configuration',
                    'Application deployment and testing',
                    'Production cutover and optimization'
                ]
            }
        
        # If using modern frameworks but not containerized
        elif tech_stack['cloud_ready'] and not tech_stack['containers']:
            return {
                'pattern': 'Rehost with Containerization',
                'description': 'Lift-and-shift to Azure VMs initially, then containerize for cloud-native benefits',
                'rationale': 'Modern technology stack supports containerization. A phased approach reduces risk while enabling cloud-native transformation.',
                'considerations': [
                    'Initial VM migration provides quick cloud benefits',
                    'Application architecture supports containerization',
                    'Database migration strategy for cloud-native services',
                    'Load balancer and networking reconfiguration',
                    'Gradual transition to managed Azure services'
                ],
                'phases': [
                    'Initial assessment and Azure VM sizing',
                    'Lift-and-shift to Azure Virtual Machines',
                    'Application containerization and testing',
                    'Migration to Azure Container Instances or AKS',
                    'Integration with Azure managed services',
                    'Performance optimization and cost management'
                ]
            }
        
        # Traditional applications
        else:
            return {
                'pattern': 'Rehost (Lift and Shift)',
                'description': 'Migrate existing application to Azure VMs with minimal changes, providing foundation for future modernization',
                'rationale': 'Traditional deployment approach requires minimal initial changes, reducing migration risk while establishing cloud presence for future optimization.',
                'considerations': [
                    'Preserve existing application architecture during migration',
                    'Minimal code changes required for initial migration',
                    'Network connectivity and security controls adaptation',
                    'Database migration with minimal downtime strategies',
                    'Monitoring and backup integration with Azure services'
                ],
                'phases': [
                    'Infrastructure assessment and Azure resource planning',
                    'Network and security architecture design',
                    'VM provisioning and configuration',
                    'Application and database migration',
                    'Testing and validation in Azure environment',
                    'Production cutover and performance tuning'
                ]
            }
    
    def _assess_migration_complexity(self, tech_stack: Dict[str, Any], architecture_type: str, deployment_method: str) -> Dict[str, str]:
        """Assess migration complexity based on technical factors."""
        
        complexity_score = 0
        factors = []
        
        # Technology stack complexity
        if tech_stack['containers']:
            complexity_score += 1
            factors.append("Containerized deployment reduces complexity")
        else:
            complexity_score += 2
            factors.append("Traditional deployment requires more migration effort")
        
        if tech_stack['cloud_ready']:
            complexity_score += 1
            factors.append("Cloud-friendly technology stack")
        else:
            complexity_score += 3
            factors.append("Legacy technology stack requires modernization")
        
        # Architecture complexity
        if architecture_type == 'microservices':
            complexity_score += 3
            factors.append("Microservices architecture increases coordination complexity")
        elif architecture_type == 'n-tier':
            complexity_score += 2
            factors.append("N-tier architecture requires careful tier migration planning")
        else:
            complexity_score += 1
            factors.append("Monolithic architecture simplifies migration coordination")
        
        # Database complexity
        if len(tech_stack['databases']) > 1:
            complexity_score += 2
            factors.append("Multiple database technologies increase migration complexity")
        
        # Determine final complexity level
        if complexity_score <= 3:
            level = "Low"
            description = "Straightforward migration with minimal technical challenges. Well-suited for cloud deployment with existing technology choices."
        elif complexity_score <= 6:
            level = "Medium"
            description = "Moderate complexity requiring careful planning and phased approach. Some technology adaptations needed for optimal cloud deployment."
        else:
            level = "High"
            description = "Complex migration requiring significant planning, potential refactoring, and multiple phases. Consider modernization opportunities during migration."
        
        return {
            'level': level,
            'description': f"{description}\n\nKey factors: {'; '.join(factors)}"
        }

    def _format_technology_selection_content(self, assessment_data: AssessmentReportData) -> str:
        """Format technology selection content using LLM analysis for comprehensive architectural recommendations."""
        
        if not self.llm_client:
            return self._basic_technology_selection_content(assessment_data)
        
        # Compile all answered questions for LLM analysis
        qa_context = []
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer not in ["Not addressed in transcript", "Error in analysis"]:
                qa_context.append(f"Q: {qa.question}\nA: {qa.answer}")
        
        if not qa_context:
            return self._basic_technology_selection_content(assessment_data)
        
        context_text = "\n\n".join(qa_context)
        
        prompt = f"""Analyze the following Q&A conversation to provide comprehensive Azure technology selection and architecture recommendations for migration.

Q&A Context:
{context_text}

Based on this conversation, provide a detailed technology selection analysis covering:

1. **Current Technology Stack Analysis**: Identify technologies, frameworks, databases, and infrastructure components mentioned
2. **Migration Strategy Recommendation**: Recommend the best Azure migration approach (Rehost, Replatform, Refactor, etc.)
3. **Azure Services Recommendations**: Specific Azure services mapped to current technology components
4. **Architecture Considerations**: Design patterns, scalability, security, and integration recommendations
5. **Modernization Opportunities**: Areas where cloud-native services could improve the solution

Format your response as a comprehensive technology selection document with clear sections and specific recommendations based on what was discussed in the conversation. Include:

- Specific technology versions and configurations mentioned
- Database migration strategies for the technologies identified
- Compute service recommendations based on current deployment
- Network and security service recommendations
- Cost optimization opportunities
- Implementation phases and priorities

Focus on making specific recommendations based on the actual technologies and requirements discussed rather than generic advice."""

        try:
            response = self.llm_client.invoke(prompt)
            content = response.content.strip()
            
            if content and len(content) > 200:
                return f"**Azure Technology Selection and Architecture Strategy:**\n\n{content}"
            else:
                print("Warning: LLM technology selection analysis returned insufficient content")
                return self._basic_technology_selection_content(assessment_data)
                
        except Exception as e:
            print(f"Warning: LLM technology selection analysis failed: {str(e)}")
            return self._basic_technology_selection_content(assessment_data)
    
    def _basic_technology_selection_content(self, assessment_data: AssessmentReportData) -> str:
        """Fallback basic technology selection content."""
        
        # Analyze current technology stack
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        architecture_type = self._analyze_architecture_type(assessment_data.questions_answers)
        deployment_method = self._analyze_deployment_method(assessment_data.questions_answers)
        migration_pattern = self._recommend_migration_pattern(tech_stack, architecture_type, deployment_method)
        
        # Generate Azure service recommendations based on analysis
        azure_services = self._recommend_azure_services(tech_stack, migration_pattern['pattern'])
        
        content = "**Azure Technology Selection and Architecture Strategy:**\n\n"
        
        # Show current technology stack analysis if detected
        if any([tech_stack['languages'], tech_stack['frameworks'], tech_stack['databases'], tech_stack['infrastructure']]):
            content += "**Current Technology Stack Analysis:**\n\n"
            
            if tech_stack['languages']:
                content += f"• **Programming Languages**: {', '.join(tech_stack['languages'])} - Modern, cloud-compatible languages\n"
            if tech_stack['frameworks']:
                content += f"• **Application Frameworks**: {', '.join(tech_stack['frameworks'])} - Proven frameworks with Azure integration support\n"
            if tech_stack['databases']:
                content += f"• **Database Technologies**: {', '.join(tech_stack['databases'])} - Well-supported with Azure managed services\n"
            if tech_stack['infrastructure']:
                content += f"• **Infrastructure Components**: {', '.join(tech_stack['infrastructure'])} - Can be replaced with Azure native services\n"
            
            content += f"• **Deployment Approach**: {deployment_method.title()} - "
            if tech_stack['containers']:
                content += "Containerized deployment enables cloud-native migration strategies\n"
            else:
                content += "Traditional deployment suitable for lift-and-shift approach\n"
            
            content += f"• **Cloud Readiness**: {'High' if tech_stack['cloud_ready'] else 'Medium'} - "
            content += "Technology stack is well-suited for cloud migration\n" if tech_stack['cloud_ready'] else "Some modernization opportunities available\n"
            
            content += "\n"
        else:
            content += "**Application Technology Details**: N/A - Specific technology stack not detailed in the transcript.\n\n"
        
        # Add migration strategy alignment
        content += f"**Selected Migration Strategy**: {migration_pattern['pattern']}\n\n"
        content += f"**Strategy Rationale**: {migration_pattern['rationale']}\n\n"
        
        # Recommended Azure Services with architectural justification
        content += "**Recommended Azure Services Architecture:**\n\n"
        
        # Group services by architectural layer
        for category, services in azure_services.items():
            if services:
                content += f"**{category}:**\n"
                for service in services:
                    content += f"• **{service['name']}**\n"
                    content += f"  - *Purpose*: {service['description']}\n"
                    content += f"  - *Architectural Fit*: {service['rationale']}\n"
                content += "\n"
        
        # Add architecture-specific recommendations
        content += "**Architectural Design Considerations:**\n\n"
        
        if architecture_type == 'microservices':
            content += "• **Microservices Architecture**: Design Azure service mesh with proper service discovery and inter-service communication\n"
            content += "• **Service Isolation**: Implement proper network segmentation and security boundaries\n"
            content += "• **API Management**: Use Azure API Management for service orchestration and external API exposure\n"
        elif architecture_type == 'n-tier':
            content += "• **N-Tier Architecture**: Implement clear separation between presentation, business, and data layers\n"
            content += "• **Load Balancing**: Design proper load distribution across application tiers\n"
            content += "• **Database Tier**: Implement high availability and disaster recovery for data layer\n"
        else:
            content += "• **Monolithic Architecture**: Plan for eventual decomposition into smaller, manageable services\n"
            content += "• **Modernization Path**: Identify opportunities for gradual migration to cloud-native patterns\n"
        
        if tech_stack['containers']:
            content += "• **Container Orchestration**: Leverage Kubernetes patterns for scalability and resilience\n"
            content += "• **Container Security**: Implement Azure security best practices for container workloads\n"
            content += "• **Image Management**: Establish secure container image lifecycle and vulnerability scanning\n"
        
        # Technology-specific considerations
        if 'Django' in tech_stack['frameworks']:
            content += "• **Django Framework**: Configure Azure App Service or AKS for Django applications with proper static file handling\n"
        if 'PostgreSQL' in tech_stack['databases']:
            content += "• **PostgreSQL Migration**: Plan for Azure Database for PostgreSQL with connection pooling and performance optimization\n"
        if 'Redis' in tech_stack['databases']:
            content += "• **Redis Caching**: Implement Azure Cache for Redis with proper data persistence and security\n"
        
        content += "\n**Service Integration and Dependencies:**\n\n"
        content += "• **Service Mesh**: Implement proper service-to-service communication patterns\n"
        content += "• **Data Flow**: Design secure data flow between application and database tiers\n"
        content += "• **Monitoring Integration**: Ensure comprehensive observability across all service layers\n"
        content += "• **Security Integration**: Implement Azure security services consistently across all components\n"
        content += "• **Backup and Recovery**: Design coordinated backup strategies across all service tiers\n"
        
        content += "\n**Future Modernization Opportunities:**\n\n"
        if not tech_stack['containers']:
            content += "• **Containerization**: Future migration to containerized deployment for improved portability\n"
        content += "• **Serverless Integration**: Identify components suitable for Azure Functions migration\n"
        content += "• **Managed Services**: Gradual adoption of additional Azure managed services for operational efficiency\n"
        content += "• **DevOps Integration**: Implementation of Azure DevOps for CI/CD and infrastructure automation\n"
        
        return content
    
    def _recommend_azure_services(self, tech_stack: Dict[str, Any], deployment_method: str) -> Dict[str, List[Dict[str, str]]]:
        """Recommend Azure services based on technology stack analysis using AI."""
        
        if not self.llm_client:
            # Fallback to basic recommendations if no LLM available
            return {
                'Compute Services': [{'name': 'Azure Virtual Machines', 'description': 'Scalable compute resources', 'rationale': 'General compute service for application hosting'}],
                'Database Services': [{'name': 'Azure SQL Database', 'description': 'Managed database service', 'rationale': 'Cloud database service'}],
                'Networking Services': [{'name': 'Azure Virtual Network', 'description': 'Private networking', 'rationale': 'Network isolation and security'}],
                'Security Services': [{'name': 'Azure Key Vault', 'description': 'Secrets management', 'rationale': 'Secure credential storage'}],
                'Monitoring & Management': [{'name': 'Azure Monitor', 'description': 'Monitoring and alerting', 'rationale': 'Application and infrastructure monitoring'}],
                'Storage Services': [{'name': 'Azure Storage', 'description': 'Cloud storage', 'rationale': 'General purpose storage'}]
            }
        
        azure_services_prompt = f"""Based on this technology stack analysis, recommend specific Azure services that best match the application requirements.

TECHNOLOGY STACK:
{tech_stack}

DEPLOYMENT METHOD: {deployment_method}

Provide Azure service recommendations in the following JSON format:
{{
    "recommendations": {{
        "Compute Services": [
            {{"name": "specific Azure compute service", "description": "what it provides", "rationale": "why this service fits the technology stack"}}
        ],
        "Database Services": [
            {{"name": "specific Azure database service", "description": "what it provides", "rationale": "why this matches the detected database technology"}}
        ],
        "Networking Services": [
            {{"name": "specific Azure networking service", "description": "what it provides", "rationale": "why this networking approach is recommended"}}
        ],
        "Security Services": [
            {{"name": "specific Azure security service", "description": "what it provides", "rationale": "why this security service is needed"}}
        ],
        "Monitoring & Management": [
            {{"name": "specific Azure monitoring service", "description": "what it provides", "rationale": "why this monitoring approach fits"}}
        ],
        "Storage Services": [
            {{"name": "specific Azure storage service", "description": "what it provides", "rationale": "why this storage solution is appropriate"}}
        ]
    }}
}}

Recommendations should:
1. Map detected technologies to appropriate Azure services
2. Consider the deployment method (traditional/containerized/kubernetes)
3. Match database technologies to Azure database services
4. Include security and monitoring services appropriate for the technology stack
5. Provide specific rationale based on the actual technologies detected"""

        # Get AI analysis
        result = self._llm_analyze(azure_services_prompt, {
            "recommendations": {
                'Compute Services': [{'name': 'Azure Virtual Machines', 'description': 'Scalable compute resources', 'rationale': 'General compute service based on technology stack'}],
                'Database Services': [{'name': 'Azure SQL Database', 'description': 'Managed database service', 'rationale': 'Managed database service appropriate for application'}],
                'Networking Services': [{'name': 'Azure Virtual Network', 'description': 'Private networking', 'rationale': 'Network isolation and security'}],
                'Security Services': [{'name': 'Azure Key Vault', 'description': 'Secrets management', 'rationale': 'Secure credential storage'}],
                'Monitoring & Management': [{'name': 'Azure Monitor', 'description': 'Monitoring and alerting', 'rationale': 'Application and infrastructure monitoring'}],
                'Storage Services': [{'name': 'Azure Storage', 'description': 'Cloud storage', 'rationale': 'General purpose storage'}]
            }
        })
        
        if isinstance(result, dict) and "recommendations" in result:
            return result["recommendations"]
        
        # Fallback structure
        return {
            'Compute Services': [{'name': 'Azure Virtual Machines', 'description': 'Scalable compute resources', 'rationale': 'General compute service for application hosting'}],
            'Database Services': [{'name': 'Azure SQL Database', 'description': 'Managed database service', 'rationale': 'Cloud database service'}],
            'Networking Services': [{'name': 'Azure Virtual Network', 'description': 'Private networking', 'rationale': 'Network isolation and security'}],
            'Security Services': [{'name': 'Azure Key Vault', 'description': 'Secrets management', 'rationale': 'Secure credential storage'}],
            'Monitoring & Management': [{'name': 'Azure Monitor', 'description': 'Monitoring and alerting', 'rationale': 'Application and infrastructure monitoring'}],
            'Storage Services': [{'name': 'Azure Storage', 'description': 'Cloud storage', 'rationale': 'General purpose storage'}]
        }

    def _format_azure_cost_content(self, assessment_data: AssessmentReportData) -> str:
        """Format Azure cost content based on intelligent technology analysis and recommended services."""
        
        # Analyze current technology stack to determine recommended services
        tech_stack = self._analyze_technology_stack(assessment_data.questions_answers)
        architecture_type = self._analyze_architecture_type(assessment_data.questions_answers)
        deployment_method = self._analyze_deployment_method(assessment_data.questions_answers)
        migration_pattern = self._recommend_migration_pattern(tech_stack, architecture_type, deployment_method)
        azure_services = self._recommend_azure_services(tech_stack, migration_pattern['pattern'])
        
        # Try to extract cost information from Q&A first
        cost_info = []
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                if any(keyword in question_lower for keyword in ['cost', 'budget', 'price', 'estimate']):
                    cost_info.append(qa.answer)
        
        content = "**Indicative Monthly Azure Costs:**\n\n"
        
        if cost_info:
            content += "Based on the assessment discussion:\n"
            content += f"• {cost_info[0]}\n\n"
        
        content += "**Estimated costs based on recommended Azure services:**\n\n"
        
        # Calculate costs based on recommended services
        total_min_cost = 0
        total_max_cost = 0
        cost_breakdown = []
        
        # Compute service costs
        for service_category, services in azure_services.items():
            category_min = 0
            category_max = 0
            
            if service_category == "Compute Services":
                if deployment_method == "kubernetes":
                    cost_breakdown.append("| Azure Kubernetes Service (AKS) | $300 - $800 |")
                    category_min, category_max = 300, 800
                elif any('App Service' in str(service) for service in services):
                    cost_breakdown.append("| Azure App Service (Premium) | $200 - $500 |")
                    category_min, category_max = 200, 500
                else:
                    cost_breakdown.append("| Azure Virtual Machines | $400 - $800 |")
                    category_min, category_max = 400, 800
            
            elif service_category == "Database Services":
                db_cost_min = 0
                db_cost_max = 0
                
                if 'PostgreSQL' in tech_stack['databases']:
                    cost_breakdown.append("| Azure Database for PostgreSQL | $150 - $400 |")
                    db_cost_min += 150
                    db_cost_max += 400
                
                if 'Redis' in tech_stack['databases']:
                    cost_breakdown.append("| Azure Cache for Redis | $100 - $250 |")
                    db_cost_min += 100
                    db_cost_max += 250
                
                if 'MySQL' in tech_stack['databases']:
                    cost_breakdown.append("| Azure Database for MySQL | $150 - $350 |")
                    db_cost_min += 150
                    db_cost_max += 350
                
                if not tech_stack['databases']:
                    cost_breakdown.append("| Azure SQL Database | $200 - $500 |")
                    db_cost_min, db_cost_max = 200, 500
                
                category_min, category_max = db_cost_min, db_cost_max
            
            elif service_category == "Networking Services":
                cost_breakdown.append("| Application Gateway + VNet | $150 - $300 |")
                category_min, category_max = 150, 300
            
            elif service_category == "Security Services":
                cost_breakdown.append("| Key Vault + Security Center | $50 - $150 |")
                category_min, category_max = 50, 150
            
            elif service_category == "Monitoring & Management":
                cost_breakdown.append("| Azure Monitor + App Insights | $100 - $250 |")
                category_min, category_max = 100, 250
            
            elif service_category == "Storage Services":
                cost_breakdown.append("| Blob Storage + File Storage | $50 - $150 |")
                category_min, category_max = 50, 150
            
            total_min_cost += category_min
            total_max_cost += category_max
        
        # Create cost table
        content += "| Service Category | Estimated Monthly Cost |\n"
        content += "|------------------|----------------------|\n"
        content += "\n".join(cost_breakdown)
        content += f"\n| **Total Estimated** | **${total_min_cost:,} - ${total_max_cost:,}** |\n\n"
        
        # Add technology-specific cost considerations
        content += "**Cost Analysis Based on Current Technology Stack:**\n\n"
        
        if tech_stack['containers']:
            content += "• **Containerization Advantage**: Existing containers reduce migration costs and enable efficient resource utilization\n"
        
        if tech_stack['cloud_ready']:
            content += "• **Cloud-Ready Stack**: Modern technology stack reduces migration complexity and operational costs\n"
        
        if len(tech_stack['databases']) > 1:
            content += f"• **Multi-Database Setup**: {len(tech_stack['databases'])} database technologies detected, requiring careful cost optimization\n"
        
        # Add migration pattern specific considerations
        if 'Replatform' in migration_pattern['pattern']:
            content += "• **Replatform Strategy**: Moderate migration costs with significant long-term operational savings\n"
        elif 'Rehost' in migration_pattern['pattern']:
            content += "• **Rehost Strategy**: Lower initial migration costs, gradual modernization approach\n"
        
        content += "\n**Cost Optimization Opportunities:**\n"
        content += "• **Reserved Instances**: 30-50% savings for predictable workloads\n"
        content += "• **Azure Hybrid Benefit**: Leverage existing licenses for Windows/SQL Server\n"
        content += "• **Auto-scaling**: Optimize resource utilization based on demand\n"
        content += "• **Spot Instances**: Up to 90% savings for development/testing environments\n"
        content += "• **Azure Cost Management**: Continuous monitoring and optimization\n"
        
        if deployment_method == "kubernetes":
            content += "• **AKS Cost Optimization**: Node auto-scaling and resource quotas\n"
        
        content += "\n*Note: Costs are indicative and based on recommended Azure services from technology analysis. "
        content += "Actual costs may vary based on usage patterns, data transfer, and specific service configurations. "
        content += "A detailed Azure Pricing Calculator assessment will be performed during planning phase.*"
        
        return content
    
    def _format_database_information_content(self, assessment_data: AssessmentReportData) -> str:
        """Format database information content with intelligent migration recommendations."""
        
        # Extract database info from Q&A
        db_info = []
        detected_databases = []
        
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                answer_lower = qa.answer.lower()
                
                if any(keyword in question_lower for keyword in ['database', 'sql', 'data', 'storage']):
                    db_info.append(f"• {qa.question}: {qa.answer}")
                
                # Detect specific database technologies
                if any(db in answer_lower for db in ['postgres', 'postgresql']):
                    detected_databases.append('PostgreSQL')
                if any(db in answer_lower for db in ['redis']):
                    detected_databases.append('Redis')
                if any(db in answer_lower for db in ['mysql']):
                    detected_databases.append('MySQL')
                if any(db in answer_lower for db in ['oracle']):
                    detected_databases.append('Oracle')
                if any(db in answer_lower for db in ['sql server', 'sqlserver']):
                    detected_databases.append('SQL Server')
        
        content = "**Database Configuration and Requirements:**\n\n"
        
        if db_info:
            content += "From the assessment, the following database information was identified:\n\n"
            content += "\n".join(db_info[:5])  # Limit to 5 items
            content += "\n\n"
        else:
            content += "**Application Database Details:** N/A - Not specifically addressed in the transcript.\n\n"
        
        # Provide intelligent migration strategy based on detected databases
        if detected_databases:
            content += "**Recommended Database Migration Strategy:**\n\n"
            
            for db in set(detected_databases):  # Remove duplicates
                if db == 'PostgreSQL':
                    content += f"""**{db} Migration:**
• **Target Platform**: Azure Database for PostgreSQL Flexible Server
• **Migration Method**: Azure Database Migration Service or pg_dump/pg_restore
• **High Availability**: Built-in high availability with zone redundancy
• **Backup Strategy**: Automated daily backups with point-in-time recovery (up to 35 days)
• **Security**: SSL/TLS encryption, Azure AD integration, Advanced Threat Protection

"""
                elif db == 'Redis':
                    content += f"""**{db} Migration:**
• **Target Platform**: Azure Cache for Redis Premium tier
• **Migration Method**: Redis data migration using MIGRATE command or backup/restore
• **High Availability**: Zone redundancy and geo-replication support
• **Performance**: In-memory performance with persistence options
• **Security**: SSL encryption, virtual network isolation, access policies

"""
                elif db == 'MySQL':
                    content += f"""**{db} Migration:**
• **Target Platform**: Azure Database for MySQL Flexible Server
• **Migration Method**: Azure Database Migration Service or mysqldump
• **High Availability**: Zone-redundant high availability
• **Backup Strategy**: Automated backups with configurable retention
• **Security**: SSL encryption, Azure AD authentication, threat protection

"""
                else:
                    content += f"""**{db} Migration:**
• **Assessment Required**: Detailed compatibility assessment needed
• **Migration Tools**: Azure Database Migration Service evaluation
• **Alternatives**: Consider Azure SQL Database for modernization opportunities
• **Support**: Review Azure support for {db} or migration to supported alternatives

"""
        else:
            # Provide general database migration recommendations
            content += """**Recommended Database Strategy:**

Since specific database technologies were not detailed in the assessment, the following general approach is recommended:

• **Assessment Phase**: Conduct detailed database inventory and dependency analysis
• **Platform Selection**: Evaluate Azure managed database services based on current setup
• **Migration Approach**: 
  - Phase 1: Migrate databases to Azure VMs (rehost) for minimal disruption
  - Phase 2: Migrate to Azure managed database services for operational benefits
• **High Availability**: Implement Azure-native high availability and disaster recovery
• **Security**: Enable encryption at rest and in transit, implement Azure AD integration
• **Monitoring**: Deploy Azure Monitor for Databases for performance insights"""
        
        content += "\n\n**Database Migration Best Practices:**\n"
        content += """• Perform thorough compatibility testing in non-production environments
• Implement robust backup and rollback procedures
• Plan for minimal downtime using online migration techniques
• Establish performance baselines before and after migration
• Configure monitoring and alerting for database health and performance
• Document connection string changes and application configuration updates"""
        
        return content

    def _format_macro_dependencies_content(self, assessment_data: AssessmentReportData) -> str:
        """Format macro dependencies content with intelligent architecture analysis."""
        
        # Extract dependency info from Q&A
        dep_info = []
        integration_patterns = []
        
        for qa in assessment_data.questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                answer_lower = qa.answer.lower()
                
                if any(keyword in question_lower for keyword in ['depend', 'integration', 'service', 'api', 'connection']):
                    dep_info.append(f"• {qa.question}: {qa.answer}")
                
                # Detect integration patterns
                if 'api' in answer_lower:
                    integration_patterns.append('REST APIs')
                if any(term in answer_lower for term in ['message', 'queue', 'event']):
                    integration_patterns.append('Message Queuing')
                if any(term in answer_lower for term in ['database', 'shared']):
                    integration_patterns.append('Shared Database')
                if 'file' in answer_lower:
                    integration_patterns.append('File-based Integration')
        
        content = "**System Dependencies and Integration Architecture:**\n\n"
        
        if dep_info:
            content += "The following dependencies and integrations were identified:\n\n"
            content += "\n".join(dep_info[:5])
            content += "\n\n"
        else:
            content += "**Application Dependencies:** N/A - Specific dependencies not detailed in the transcript.\n\n"
        
        # Provide intelligent architectural recommendations
        content += "**Recommended Integration Architecture for Azure:**\n\n"
        
        if integration_patterns:
            unique_patterns = list(set(integration_patterns))
            
            for pattern in unique_patterns:
                if pattern == 'REST APIs':
                    content += """**API Integration Strategy:**
• **API Management**: Azure API Management for centralized API governance
• **Service Discovery**: Azure Service Bus or Azure Event Grid for service coordination
• **Authentication**: Azure AD B2C or Azure AD for API security
• **Load Balancing**: Azure Application Gateway with SSL termination
• **Monitoring**: Azure Monitor and Application Insights for API performance tracking

"""
                elif pattern == 'Message Queuing':
                    content += """**Messaging Integration Strategy:**
• **Message Broker**: Azure Service Bus for enterprise messaging patterns
• **Event Streaming**: Azure Event Hubs for high-throughput event processing
• **Queue Management**: Azure Storage Queues for simple message queuing
• **Dead Letter Handling**: Built-in dead letter queue support
• **Monitoring**: Service Bus metrics and Azure Monitor integration

"""
                elif pattern == 'Shared Database':
                    content += """**Database Integration Strategy:**
• **Data Architecture**: Implement database per service pattern for microservices
• **Data Synchronization**: Azure Data Factory for ETL processes
• **Event Sourcing**: Azure Event Store or Cosmos DB for event-driven architecture
• **CQRS Pattern**: Separate read/write databases using Azure SQL and Cosmos DB
• **Data Security**: Row-level security and column encryption

"""
                elif pattern == 'File-based Integration':
                    content += """**File Integration Strategy:**
• **Storage**: Azure Blob Storage with hierarchical namespace
• **File Processing**: Azure Functions for serverless file processing
• **Workflow**: Azure Logic Apps for file-based workflow automation
• **Monitoring**: Azure Storage Analytics and file system events
• **Security**: Storage account access policies and encryption

"""
        else:
            # General integration recommendations
            content += """**General Integration Architecture:**

For cloud-native integration patterns, the following Azure services are recommended:

• **API Gateway**: Azure API Management for external API exposure and management
• **Internal Communication**: Azure Service Bus for reliable inter-service messaging
• **Event-Driven Architecture**: Azure Event Grid for reactive system integration
• **Data Integration**: Azure Data Factory for data pipeline orchestration
• **Workflow Automation**: Azure Logic Apps for business process automation
• **Security**: Azure Key Vault for secrets management across services

"""
        
        content += "**Migration Integration Strategy:**\n"
        content += """• **Phase 1**: Establish Azure backbone services (Service Bus, API Management)
• **Phase 2**: Migrate applications with maintained integration points
• **Phase 3**: Modernize integration patterns using cloud-native services
• **Phase 4**: Implement monitoring and observability across all integrations

**Best Practices:**
• Implement circuit breaker patterns for resilient integrations
• Use Azure Monitor for end-to-end distributed tracing
• Design for eventual consistency in distributed systems
• Implement proper retry policies with exponential backoff
• Use managed identities for secure service-to-service authentication"""
        
        return content
    
    def _format_decision_matrix_content(self, assessment_data: AssessmentReportData) -> str:
        """Format decision matrix content with intelligent AI-driven analysis."""
        
        # Generate intelligent decision matrix using AI
        decisions = self._generate_decision_matrix(assessment_data)
        
        if not decisions:
            # If AI analysis fails, provide minimal fallback
            return """**Migration Decision Matrix**

The following matrix outlines the key decisions to be made during the assessment:

| Decision Area | Options Considered | Selected Approach | Rationale |
|---------------|-------------------|-------------------|-----------|
| Migration Strategy | Assessment required | To be determined | Detailed analysis of application architecture and business requirements needed |
| Compute Platform | Analysis pending | To be determined | Platform selection based on application characteristics from assessment |
| Database | Technology review needed | To be determined | Database strategy based on current technology stack analysis |
| Networking | Security assessment required | To be determined | Network approach based on security and accessibility requirements |
| Authentication | Identity analysis needed | To be determined | Authentication strategy based on current systems and integration needs |

**Key Considerations:**
• Comprehensive application assessment required for informed decision-making
• Technology stack analysis needed for optimal platform selection
• Business requirements evaluation essential for strategy alignment
• Security and compliance assessment required for network and identity decisions"""
        
        # Format decisions into table
        content = """**Migration Decision Matrix**

The following matrix outlines the key decisions made during the assessment:

| Decision Area | Options Considered | Selected Approach | Rationale |
|---------------|-------------------|-------------------|-----------|"""
        
        for decision in decisions[:5]:  # Limit to 5 key decisions
            area = decision.get('area', 'Unknown')
            options = decision.get('options', 'Options analysis pending')
            selected = decision.get('selected', 'To be determined')
            rationale = decision.get('rationale', 'Detailed analysis required')
            
            # Truncate long text for table format
            if len(options) > 40:
                options = options[:37] + "..."
            if len(selected) > 30:
                selected = selected[:27] + "..."
            if len(rationale) > 50:
                rationale = rationale[:47] + "..."
            
            content += f"\n| {area} | {options} | {selected} | {rationale} |"
        
        # Generate rationale points
        rationale_points = self._generate_decision_rationale(decisions, assessment_data)
        
        content += "\n\n**Key Decisions Rationale:**"
        for point in rationale_points:
            content += f"\n• {point}"
        
        return content
    
    def _generate_target_architecture(self, questions_answers: List[QuestionAnswer], azure_migrate_data: Any, dependency_analysis: Any) -> TargetArchitecture:
        """Generate target architecture recommendations based on network traffic analysis and dependency data."""
        
        # Get network connections from dependency analysis
        network_connections = []
        if dependency_analysis and hasattr(dependency_analysis, 'connections'):
            network_connections = dependency_analysis.connections
        
        # Create target architecture object
        target_architecture = TargetArchitecture(
            network_connections=network_connections,
            subnet_recommendations=[],
            nsg_rules=[],
            load_balancer_config=[],
            recommendations=[]
        )
        
        try:
            if network_connections:
                # Generate subnet recommendations based on IP patterns
                target_architecture.subnet_recommendations = self._create_subnet_recommendations(network_connections)
                
                # Generate NSG rules based on discovered ports
                target_architecture.nsg_rules = self._create_nsg_rules(network_connections)
                
                # Generate load balancer recommendations based on multi-server patterns
                target_architecture.load_balancer_config = self._create_load_balancer_config(network_connections)
                
                # Add general recommendations
                target_architecture.recommendations = self._create_architecture_recommendations(network_connections, azure_migrate_data)
        
        except Exception as e:
            print(f"Warning: Error generating target architecture: {e}")
            # Add fallback recommendations
            target_architecture.recommendations = ["Target architecture requires detailed network analysis"]
        
        return target_architecture
    
    def _create_subnet_recommendations(self, network_connections: List) -> List[SubnetRecommendation]:
        """Create subnet recommendations based on network connections."""
        subnets = []
        
        # Analyze IP patterns
        ip_networks = {}
        for conn in network_connections:
            if hasattr(conn, 'source_ip') and conn.source_ip:
                network = '.'.join(conn.source_ip.split('.')[:3])
                if network not in ip_networks:
                    ip_networks[network] = set()
                ip_networks[network].add(conn.source_ip)
        
        # Create subnet recommendations
        for i, (network, ips) in enumerate(ip_networks.items(), 1):
            subnets.append(SubnetRecommendation(
                name=f"app-subnet-{i}",
                address_range=f"{network}.0/24",
                purpose=f"Application tier - hosts {len(ips)} discovered services",
                service_endpoints=["Microsoft.Storage", "Microsoft.KeyVault"]
            ))
        
        # Add standard subnets if none discovered
        if not subnets:
            subnets.append(SubnetRecommendation(
                name="app-subnet-1",
                address_range="10.0.1.0/24",
                purpose="Application tier subnet",
                service_endpoints=["Microsoft.Storage", "Microsoft.KeyVault"]
            ))
        
        return subnets
    
    def _create_nsg_rules(self, network_connections: List) -> List[NSGRule]:
        """Create NSG rules based on discovered network traffic."""
        rules = []
        ports_discovered = set()
        
        # Extract unique ports from connections
        for conn in network_connections:
            if hasattr(conn, 'destination_port') and conn.destination_port:
                try:
                    port = str(conn.destination_port)
                    if port.isdigit():
                        ports_discovered.add(port)
                except:
                    pass
        
        # Create rules for discovered ports
        priority = 1000
        for port in sorted(ports_discovered):
            rule_name = f"Allow_Port_{port}"
            description = self._get_port_description(port)
            
            rules.append(NSGRule(
                name=rule_name,
                direction="inbound",
                protocol="TCP",
                source_address_prefix="10.0.0.0/16",
                destination_address_prefix="*",
                destination_port=port,
                priority=priority,
                description=description
            ))
            priority += 10
        
        # Add standard rules if none discovered
        if not rules:
            rules.append(NSGRule(
                name="Allow_HTTP",
                direction="inbound",
                protocol="TCP",
                source_address_prefix="*",
                destination_address_prefix="*",
                destination_port="80",
                priority=1000,
                description="Allow HTTP traffic"
            ))
        
        return rules
    
    def _create_load_balancer_config(self, network_connections: List) -> List[LoadBalancerConfig]:
        """Create load balancer configuration based on network patterns."""
        load_balancers = []
        
        # Analyze for load balancer patterns (multiple servers on same ports)
        port_servers = {}
        for conn in network_connections:
            if hasattr(conn, 'destination_port') and hasattr(conn, 'destination_ip'):
                port = str(conn.destination_port)
                if port not in port_servers:
                    port_servers[port] = set()
                port_servers[port].add(conn.destination_ip)
        
        # Create load balancers for ports with multiple servers
        for port, servers in port_servers.items():
            if len(servers) > 1:
                load_balancers.append(LoadBalancerConfig(
                    name=f"lb-app-{port}",
                    type="internal",
                    frontend_ip_config="Dynamic",
                    backend_pools=[f"backend-pool-{port}"],
                    load_balancing_rules=[LoadBalancingRule(
                        frontend_port=port,
                        backend_port=port,
                        protocol="TCP"
                    )]
                ))
        
        return load_balancers
    
    def _create_architecture_recommendations(self, network_connections: List, azure_migrate_data: Any) -> List[str]:
        """Create general architecture recommendations."""
        recommendations = []
        
        connection_count = len(network_connections)
        recommendations.append(f"Analyzed {connection_count} network connections for architecture planning")
        
        # Port analysis
        ports = set()
        for conn in network_connections:
            if hasattr(conn, 'destination_port') and conn.destination_port:
                ports.add(str(conn.destination_port))
        
        if ports:
            recommendations.append(f"Configure NSG rules for {len(ports)} discovered ports: {', '.join(sorted(ports)[:5])}")
        
        # IP analysis
        unique_ips = set()
        for conn in network_connections:
            if hasattr(conn, 'source_ip') and conn.source_ip:
                unique_ips.add(conn.source_ip)
            if hasattr(conn, 'destination_ip') and conn.destination_ip:
                unique_ips.add(conn.destination_ip)
        
        if unique_ips:
            recommendations.append(f"Plan VNet addressing for {len(unique_ips)} discovered IP addresses")
        
        # Add Azure Migrate insights
        if azure_migrate_data and hasattr(azure_migrate_data, 'servers'):
            recommendations.append(f"Consider Azure Migrate recommendations for {len(azure_migrate_data.servers)} servers")
        
        return recommendations
    
    def _get_port_description(self, port: str) -> str:
        """Get description for common ports."""
        port_descriptions = {
            "80": "HTTP web traffic",
            "443": "HTTPS secure web traffic",
            "3389": "RDP remote desktop",
            "22": "SSH secure shell",
            "1433": "SQL Server database",
            "3306": "MySQL database",
            "5432": "PostgreSQL database",
            "6379": "Redis cache",
            "8080": "HTTP alternative port",
            "9000": "Application server"
        }
        return port_descriptions.get(port, f"Application traffic on port {port}")
    
    def _generate_subnet_recommendations(self, source_ips: set, destination_ips: set, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Generate subnet recommendations based on IP traffic analysis."""
        recommendations = []
        
        # Analyze IP patterns to suggest subnet structure
        all_ips = source_ips.union(destination_ips)
        ip_networks = {}
        
        for ip in all_ips:
            if ip and '.' in ip:
                try:
                    # Extract network prefix (first 3 octets)
                    network_prefix = '.'.join(ip.split('.')[:3])
                    if network_prefix not in ip_networks:
                        ip_networks[network_prefix] = []
                    ip_networks[network_prefix].append(ip)
                except:
                    continue
        
        # Generate subnet recommendations
        for i, (network, ips) in enumerate(ip_networks.items(), 1):
            recommendations.append({
                'subnet_name': f'app-subnet-{i}',
                'address_space': f'{network}.0/24',
                'purpose': f'Application tier - hosts {len(ips)} identified services',
                'security_requirements': 'NSG required for traffic filtering'
            })
        
        # Add standard Azure subnets
        recommendations.extend([
            {
                'subnet_name': 'gateway-subnet',
                'address_space': 'To be determined based on VNet planning',
                'purpose': 'VPN Gateway for hybrid connectivity',
                'security_requirements': 'Gateway-specific NSG rules'
            },
            {
                'subnet_name': 'bastion-subnet',
                'address_space': 'To be determined based on VNet planning', 
                'purpose': 'Azure Bastion for secure management access',
                'security_requirements': 'Bastion-specific NSG rules'
            }
        ])
        
        return recommendations
    
    def _generate_nsg_rules(self, ports_used: set, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Generate Network Security Group rules based on identified traffic patterns."""
        nsg_rules = []
        
        # Create rules for discovered ports
        for port in sorted(ports_used):
            if port and port.isdigit():
                port_int = int(port)
                service_name = self._identify_service_by_port(port_int)
                
                nsg_rules.append({
                    'rule_name': f'Allow-{service_name}-{port}',
                    'direction': 'Inbound',
                    'priority': str(1000 + len(nsg_rules)),
                    'source': 'Application subnet',
                    'destination': 'Application subnet',
                    'port': port,
                    'protocol': 'TCP',
                    'action': 'Allow',
                    'description': f'Allow {service_name} traffic on port {port}'
                })
        
        # Add standard security rules
        nsg_rules.extend([
            {
                'rule_name': 'Deny-All-Inbound',
                'direction': 'Inbound',
                'priority': '4096',
                'source': 'Any',
                'destination': 'Any',
                'port': '*',
                'protocol': '*',
                'action': 'Deny',
                'description': 'Deny all other inbound traffic'
            },
            {
                'rule_name': 'Allow-HTTPS-Internet',
                'direction': 'Inbound',
                'priority': '100',
                'source': 'Internet',
                'destination': 'Any',
                'port': '443',
                'protocol': 'TCP',
                'action': 'Allow',
                'description': 'Allow HTTPS from Internet'
            }
        ])
        
        return nsg_rules
    
    def _identify_service_by_port(self, port: int) -> str:
        """Identify service type by port number."""
        common_ports = {
            80: 'HTTP',
            443: 'HTTPS',
            22: 'SSH',
            3389: 'RDP',
            1433: 'SQL-Server',
            3306: 'MySQL',
            5432: 'PostgreSQL',
            1521: 'Oracle',
            25: 'SMTP',
            587: 'SMTP-Submission',
            110: 'POP3',
            143: 'IMAP',
            993: 'IMAPS',
            995: 'POP3S',
            53: 'DNS',
            123: 'NTP',
            161: 'SNMP',
            389: 'LDAP',
            636: 'LDAPS',
            8080: 'HTTP-Alt',
            8443: 'HTTPS-Alt'
        }
        
        return common_ports.get(port, f'Port-{port}')
    
    def _generate_load_balancer_recommendations(self, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Generate load balancer recommendations based on traffic patterns."""
        recommendations = []
        
        # Analyze traffic to determine if load balancing is needed
        if hasattr(dependency_analysis, 'connections') and dependency_analysis.connections:
            # Count connections to each destination to identify high-traffic services
            destination_traffic = {}
            for conn in dependency_analysis.connections:
                if hasattr(conn, 'destination_ip') and conn.destination_ip:
                    dest = conn.destination_ip
                    destination_traffic[dest] = destination_traffic.get(dest, 0) + 1
            
            # Recommend load balancing for high-traffic destinations
            for dest, count in destination_traffic.items():
                if count > 5:  # Threshold for load balancing consideration
                    recommendations.append({
                        'type': 'Azure Application Gateway',
                        'target': dest,
                        'reason': f'High traffic volume ({count} connections identified)',
                        'configuration': 'Web Application Firewall enabled, SSL termination'
                    })
        
        # Always recommend standard load balancing patterns
        recommendations.extend([
            {
                'type': 'Azure Load Balancer (Standard)',
                'target': 'Application tier',
                'reason': 'High availability and scalability',
                'configuration': 'Internal load balancer for backend services'
            },
            {
                'type': 'Azure Application Gateway',
                'target': 'Web tier',
                'reason': 'Layer 7 load balancing and WAF protection',
                'configuration': 'Public-facing with SSL/TLS termination'
            }
        ])
        
        return recommendations
    
    def _generate_compute_recommendations(self, applications: set, dependency_analysis: Any, azure_migrate_data: Any) -> List[Dict[str, str]]:
        """Generate compute service recommendations based on identified applications and traffic."""
        recommendations = []
        
        # Analyze applications for compute recommendations
        for app in applications:
            if app and app.lower() not in ['system', 'unknown', 'n/a']:
                recommendations.append({
                    'application': app,
                    'recommended_service': self._recommend_azure_service(app),
                    'justification': f'Based on application type and traffic patterns',
                    'scaling_approach': 'Auto-scaling based on demand'
                })
        
        # Add Azure Migrate recommendations if available
        if azure_migrate_data and hasattr(azure_migrate_data, 'servers'):
            for server in azure_migrate_data.servers[:5]:  # Limit to first 5 servers
                recommendations.append({
                    'application': f'{server.server_name} ({server.server_type})',
                    'recommended_service': f'Azure VM {getattr(server, "azure_vm_size", "Standard_D2s_v3")}',
                    'justification': f'Based on Azure Migrate assessment - {getattr(server, "readiness", "Ready")}',
                    'scaling_approach': 'VM Scale Sets for horizontal scaling'
                })
        
        return recommendations
    
    def _recommend_azure_service(self, application: str) -> str:
        """Recommend appropriate Azure service based on application type."""
        app_lower = application.lower()
        
        if any(keyword in app_lower for keyword in ['web', 'http', 'iis', 'apache', 'nginx']):
            return 'Azure App Service or Azure Container Apps'
        elif any(keyword in app_lower for keyword in ['database', 'sql', 'mysql', 'postgres', 'oracle']):
            return 'Azure Database Service (PaaS) or SQL Server on VM'
        elif any(keyword in app_lower for keyword in ['cache', 'redis', 'memcache']):
            return 'Azure Cache for Redis'
        elif any(keyword in app_lower for keyword in ['file', 'share', 'storage']):
            return 'Azure Files or Azure Blob Storage'
        else:
            return 'Azure Virtual Machine or Azure Container Apps'
    
    def _generate_integration_recommendations(self, dependency_analysis: Any) -> List[Dict[str, str]]:
        """Generate integration recommendations based on discovered connections."""
        integrations = []
        
        if hasattr(dependency_analysis, 'connections') and dependency_analysis.connections:
            # Analyze external connections
            external_connections = []
            internal_connections = []
            
            for conn in dependency_analysis.connections:
                if hasattr(conn, 'source_ip') and hasattr(conn, 'destination_ip'):
                    # Simple heuristic: different network prefixes suggest external connectivity
                    source_net = '.'.join(conn.source_ip.split('.')[:3]) if conn.source_ip else ''
                    dest_net = '.'.join(conn.destination_ip.split('.')[:3]) if conn.destination_ip else ''
                    
                    if source_net != dest_net:
                        external_connections.append(conn)
                    else:
                        internal_connections.append(conn)
            
            # Recommend integration patterns
            if external_connections:
                integrations.append({
                    'type': 'Hybrid Connectivity',
                    'requirement': f'{len(external_connections)} external connections identified',
                    'recommendation': 'Azure ExpressRoute or Site-to-Site VPN',
                    'implementation': 'Virtual Network Gateway with appropriate bandwidth'
                })
            
            if internal_connections:
                integrations.append({
                    'type': 'Service-to-Service Communication',
                    'requirement': f'{len(internal_connections)} internal connections identified',
                    'recommendation': 'Azure Private Endpoints and Service Endpoints',
                    'implementation': 'Private networking for secure service communication'
                })
        
        return integrations
    
    def _extract_business_requirements_for_architecture(self, questions_answers: List[QuestionAnswer]) -> Dict[str, Any]:
        """Extract business requirements that impact architecture decisions."""
        
        requirements = {
            'scalability_considerations': [],
            'cost_optimization': [],
            'security_requirements': []
        }
        
        for qa in questions_answers:
            if qa.is_answered and qa.answer != "Not addressed in transcript":
                question_lower = qa.question.lower()
                answer_lower = qa.answer.lower()
                
                # Scalability requirements
                if any(keyword in question_lower for keyword in ['scale', 'performance', 'user', 'load']):
                    requirements['scalability_considerations'].append({
                        'requirement': qa.question,
                        'consideration': qa.answer,
                        'architecture_impact': 'Consider auto-scaling and load balancing'
                    })
                
                # Cost requirements
                if any(keyword in question_lower for keyword in ['cost', 'budget', 'price', 'expense']):
                    requirements['cost_optimization'].append({
                        'requirement': qa.question,
                        'consideration': qa.answer,
                        'architecture_impact': 'Optimize for cost-effectiveness'
                    })
                
                # Security requirements
                if any(keyword in question_lower for keyword in ['security', 'compliance', 'data protection', 'encryption']):
                    requirements['security_requirements'].append({
                        'requirement': qa.question,
                        'consideration': qa.answer,
                        'architecture_impact': 'Implement zero-trust security model'
                    })
        
        return requirements
    
    def _format_azure_migrate_recommendations(self, azure_migrate_data: Any) -> Dict[str, Any]:
        """Format Azure Migrate recommendations for target architecture."""
        
        recommendations = {
            'servers_assessed': 0,
            'ready_for_migration': 0,
            'recommended_vm_sizes': [],
            'estimated_monthly_cost': 0
        }
        
        if hasattr(azure_migrate_data, 'servers'):
            recommendations['servers_assessed'] = len(azure_migrate_data.servers)
            
            for server in azure_migrate_data.servers:
                if hasattr(server, 'readiness') and 'ready' in str(server.readiness).lower():
                    recommendations['ready_for_migration'] += 1
                
                if hasattr(server, 'azure_vm_size'):
                    recommendations['recommended_vm_sizes'].append({
                        'server': server.server_name,
                        'vm_size': server.azure_vm_size,
                        'current_specs': f'{getattr(server, "cpu_cores", "Unknown")} cores, {getattr(server, "memory_gb", "Unknown")} GB RAM'
                    })
                
                if hasattr(server, 'estimated_cost'):
                    try:
                        cost = float(str(server.estimated_cost).replace('$', '').replace(',', ''))
                        recommendations['estimated_monthly_cost'] += cost
                    except:
                        pass
        
        return recommendations
    
    def _generate_fallback_architecture(self) -> Dict[str, Any]:
        """Generate fallback architecture when dependency analysis is not available."""
        
        return {
            'compute_recommendations': [
                {
                    'application': 'Web Application',
                    'recommended_service': 'Azure App Service',
                    'justification': 'Standard recommendation for web applications',
                    'scaling_approach': 'Built-in auto-scaling'
                }
            ],
            'network_architecture': {
                'subnet_recommendations': [
                    {
                        'subnet_name': 'web-subnet',
                        'address_space': '10.0.1.0/24',
                        'purpose': 'Web tier hosting',
                        'security_requirements': 'NSG with web traffic rules'
                    },
                    {
                        'subnet_name': 'app-subnet', 
                        'address_space': '10.0.2.0/24',
                        'purpose': 'Application tier hosting',
                        'security_requirements': 'NSG with app-specific rules'
                    },
                    {
                        'subnet_name': 'data-subnet',
                        'address_space': '10.0.3.0/24',
                        'purpose': 'Database tier hosting',
                        'security_requirements': 'NSG with database access rules'
                    }
                ],
                'network_diagram_description': 'Standard 3-tier architecture with web, application, and data layers'
            },
            'integration_points': [
                {
                    'type': 'Standard Integration',
                    'requirement': 'Basic application integration needs',
                    'recommendation': 'Azure API Management for API governance',
                    'implementation': 'Standard API management patterns'
                }
            ]
        }
