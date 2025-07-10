from typing import Dict, Any, List
from datetime import datetime
from dataclasses import asdict
from .StateBase import AzureMigrationPlan, MigrationWave, MigrationRisk, CostEstimate

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class MigrationPlanDocumentExporter:
    """Exports Azure Migration Plan to Word document format."""
    
    def __init__(self):
        self.docx_available = DOCX_AVAILABLE
    
    def export_to_word(self, migration_plan, output_path: str) -> bool:
        """
        Export migration plan to Word document.
        
        Args:
            migration_plan: Complete migration plan data
            output_path: Path to save the Word document
            
        Returns:
            True if successful, False otherwise
        """
        if not self.docx_available:
            print("Warning: python-docx not available. Install with: pip install python-docx")
            return False
        
        try:
            doc = Document()
            
            # Set document styles
            self._setup_document_styles(doc)
            
            # Title page
            self._add_title_page(doc, migration_plan)
            
            # Table of contents placeholder
            self._add_table_of_contents(doc)
            
            # Executive summary
            self._add_executive_summary(doc, migration_plan)
            
            # Current state assessment
            self._add_current_state_assessment(doc, migration_plan)
            
            # Target architecture
            self._add_target_architecture(doc, migration_plan)
            
            # Migration strategy
            self._add_migration_strategy(doc, migration_plan)
            
            # Migration timeline and waves
            self._add_migration_timeline(doc, migration_plan)
            
            # Risk assessment
            self._add_risk_assessment(doc, migration_plan)
            
            # Cost analysis
            self._add_cost_analysis(doc, migration_plan)
            
            # Implementation plan
            self._add_implementation_plan(doc, migration_plan)
            
            # Governance and compliance
            self._add_governance_compliance(doc, migration_plan)
            
            # Success metrics
            self._add_success_metrics(doc, migration_plan)
            
            # Appendices
            self._add_appendices(doc, migration_plan)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating Word document: {str(e)}")
            return False
    
    def _setup_document_styles(self, doc):
        """Setup document styles for professional appearance."""
        if not self.docx_available:
            return
            
        # Configure normal style
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        
        # Configure heading styles
        for i in range(1, 6):
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.color.rgb = None  # Use default color
    
    def _add_title_page(self, doc, migration_plan):
        """Add title page to document."""
        if not self.docx_available:
            return
            
        # Title
        title = doc.add_paragraph()
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = migration_plan.project_name
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run("Azure Migration Plan")
        subtitle_run.font.size = Pt(18)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add space
        doc.add_paragraph()
        
        # Document info
        info_table = doc.add_table(rows=4, cols=2)
        info_table.cell(0, 0).text = "Document Version:"
        info_table.cell(0, 1).text = migration_plan.document_version
        info_table.cell(1, 0).text = "Created Date:"
        info_table.cell(1, 1).text = migration_plan.created_date
        info_table.cell(2, 0).text = "Created By:"
        info_table.cell(2, 1).text = migration_plan.created_by
        info_table.cell(3, 0).text = "Last Updated:"
        info_table.cell(3, 1).text = migration_plan.last_updated
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents placeholder."""
        doc.add_heading('Table of Contents', level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run("1. Executive Summary\n")
        toc_para.add_run("2. Current State Assessment\n")
        toc_para.add_run("3. Target Architecture\n")
        toc_para.add_run("4. Migration Strategy\n")
        toc_para.add_run("5. Migration Timeline and Waves\n")
        toc_para.add_run("6. Risk Assessment\n")
        toc_para.add_run("7. Cost Analysis\n")
        toc_para.add_run("8. Implementation Plan\n")
        toc_para.add_run("9. Governance and Compliance\n")
        toc_para.add_run("10. Success Metrics\n")
        toc_para.add_run("11. Appendices\n")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, migration_plan):
        """Add executive summary section."""
        doc.add_heading('1. Executive Summary', level=1)
        
        # Executive summary content
        doc.add_paragraph(migration_plan.executive_summary)
        
        # Business case
        doc.add_heading('1.1 Business Case', level=2)
        doc.add_paragraph(migration_plan.business_case)
        
        # Key metrics
        doc.add_heading('1.2 Key Metrics', level=2)
        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Table Grid'
        
        metrics_table.cell(0, 0).text = "Total Servers"
        metrics_table.cell(0, 1).text = str(len(migration_plan.azure_migrate_data.servers))
        metrics_table.cell(1, 0).text = "Migration Waves"
        metrics_table.cell(1, 1).text = str(len(migration_plan.migration_waves))
        metrics_table.cell(2, 0).text = "Total Investment"
        metrics_table.cell(2, 1).text = f"${migration_plan.total_investment:,.2f}"
        metrics_table.cell(3, 0).text = "Expected Annual Savings"
        metrics_table.cell(3, 1).text = f"${migration_plan.expected_savings:,.2f}"
        metrics_table.cell(4, 0).text = "Project Duration"
        metrics_table.cell(4, 1).text = f"{migration_plan.migration_timeline.total_duration_months} months"
    
    def _add_current_state_assessment(self, doc, migration_plan):
        """Add current state assessment section."""
        doc.add_heading('2. Current State Assessment', level=1)
        
        current_infra = migration_plan.current_infrastructure
        
        # Infrastructure overview
        doc.add_heading('2.1 Infrastructure Overview', level=2)
        overview_para = doc.add_paragraph()
        overview_para.add_run(f"The current infrastructure consists of {current_infra['total_servers']} servers ")
        overview_para.add_run(f"with a total of {current_infra['total_cpu_cores']} CPU cores, ")
        overview_para.add_run(f"{current_infra['total_memory_gb']:.0f} GB of memory, ")
        overview_para.add_run(f"and {current_infra['total_storage_gb']:.0f} GB of storage.")
        
        # OS Distribution
        doc.add_heading('2.2 Operating System Distribution', level=2)
        if 'os_distribution' in current_infra:
            os_table = doc.add_table(rows=len(current_infra['os_distribution']) + 1, cols=2)
            os_table.style = 'Table Grid'
            os_table.cell(0, 0).text = "Operating System"
            os_table.cell(0, 1).text = "Count"
            
            for i, (os_name, count) in enumerate(current_infra['os_distribution'].items(), 1):
                os_table.cell(i, 0).text = os_name
                os_table.cell(i, 1).text = str(count)
        
        # Server specifications
        doc.add_heading('2.3 Server Specifications Summary', level=2)
        specs_para = doc.add_paragraph()
        specs_para.add_run(f"Average CPU cores per server: {current_infra.get('average_cpu_per_server', 0):.1f}\n")
        specs_para.add_run(f"Average memory per server: {current_infra.get('average_memory_per_server', 0):.1f} GB\n")
        specs_para.add_run(f"Average storage per server: {current_infra.get('average_storage_per_server', 0):.1f} GB")
    
    def _add_target_architecture(self, doc, migration_plan):
        """Add target architecture section."""
        doc.add_heading('3. Target Architecture', level=1)
        
        # Architecture overview
        doc.add_heading('3.1 Architecture Overview', level=2)
        arch_para = doc.add_paragraph()
        arch_para.add_run(f"The target Azure architecture includes {len(migration_plan.target_services)} ")
        arch_para.add_run("Azure services designed to provide scalability, security, and high availability. ")
        arch_para.add_run("The architecture follows Azure Well-Architected Framework principles.")
        
        # Target services
        doc.add_heading('3.2 Target Azure Services', level=2)
        if migration_plan.target_services:
            services_table = doc.add_table(rows=len(migration_plan.target_services) + 1, cols=4)
            services_table.style = 'Table Grid'
            services_table.cell(0, 0).text = "Service"
            services_table.cell(0, 1).text = "Component Count"
            services_table.cell(0, 2).text = "Migration Strategy"
            services_table.cell(0, 3).text = "Estimated Effort"
            
            for i, service in enumerate(migration_plan.target_services, 1):
                services_table.cell(i, 0).text = service.get('service_name', '')
                services_table.cell(i, 1).text = str(service.get('component_count', 0))
                services_table.cell(i, 2).text = service.get('migration_strategy', '')
                services_table.cell(i, 3).text = service.get('estimated_effort', '')
        
        # Architecture diagram reference
        doc.add_heading('3.3 Architecture Diagram', level=2)
        doc.add_paragraph("The detailed architecture diagram is available in the following formats:")
        diagram_para = doc.add_paragraph()
        diagram_para.add_run("• VSDX format for Microsoft Visio editing\n")
        diagram_para.add_run("• SVG format for web viewing")
    
    def _add_migration_strategy(self, doc, migration_plan):
        """Add migration strategy section."""
        doc.add_heading('4. Migration Strategy', level=1)
        
        # Overall approach
        doc.add_heading('4.1 Migration Approach', level=2)
        doc.add_paragraph(migration_plan.migration_approach)
        
        # Migration waves overview
        doc.add_heading('4.2 Migration Waves Overview', level=2)
        waves_para = doc.add_paragraph()
        waves_para.add_run(f"The migration is structured into {len(migration_plan.migration_waves)} waves ")
        waves_para.add_run("to minimize risk and ensure systematic progression from pilot to production systems.")
        
        for wave in migration_plan.migration_waves:
            doc.add_heading(f'Wave {wave.wave_number}: {wave.name}', level=3)
            wave_para = doc.add_paragraph()
            wave_para.add_run(f"Description: {wave.description}\n")
            wave_para.add_run(f"Duration: {wave.duration_weeks} weeks\n")
            wave_para.add_run(f"Servers: {len(wave.servers)}\n")
            wave_para.add_run(f"Risk Level: {wave.risk_level}\n")
            wave_para.add_run(f"Estimated Cost: ${wave.estimated_cost:,.2f}")
    
    def _add_migration_timeline(self, doc, migration_plan):
        """Add migration timeline section."""
        doc.add_heading('5. Migration Timeline and Waves', level=1)
        
        # Timeline overview
        doc.add_heading('5.1 Timeline Overview', level=2)
        timeline_para = doc.add_paragraph()
        timeline_para.add_run(f"Total project duration: {migration_plan.migration_timeline.total_duration_months} months\n")
        timeline_para.add_run(f"Number of migration waves: {len(migration_plan.migration_timeline.waves)}")
        
        # Key milestones
        doc.add_heading('5.2 Key Milestones', level=2)
        if migration_plan.migration_timeline.key_milestones:
            milestones_table = doc.add_table(rows=len(migration_plan.migration_timeline.key_milestones) + 1, cols=3)
            milestones_table.style = 'Table Grid'
            milestones_table.cell(0, 0).text = "Milestone"
            milestones_table.cell(0, 1).text = "Date"
            milestones_table.cell(0, 2).text = "Description"
            
            for i, milestone in enumerate(migration_plan.migration_timeline.key_milestones, 1):
                milestones_table.cell(i, 0).text = milestone.get('milestone', '')
                milestones_table.cell(i, 1).text = milestone.get('date', '')
                milestones_table.cell(i, 2).text = milestone.get('description', '')
        
        # Resource requirements
        doc.add_heading('5.3 Resource Requirements', level=2)
        if migration_plan.migration_timeline.resource_requirements:
            for role, requirement in migration_plan.migration_timeline.resource_requirements.items():
                doc.add_paragraph(f"• {role.replace('_', ' ').title()}: {requirement}")
    
    def _add_risk_assessment(self, doc, migration_plan):
        """Add risk assessment section."""
        doc.add_heading('6. Risk Assessment', level=1)
        
        # Risk overview
        doc.add_heading('6.1 Risk Overview', level=2)
        risk_para = doc.add_paragraph()
        risk_para.add_run(f"A total of {len(migration_plan.risks)} risks have been identified ")
        risk_para.add_run("and assessed for this migration project. Each risk includes mitigation strategies ")
        risk_para.add_run("and assigned ownership for tracking and resolution.")
        
        # Risk matrix
        doc.add_heading('6.2 Risk Register', level=2)
        if migration_plan.risks:
            risk_table = doc.add_table(rows=len(migration_plan.risks) + 1, cols=6)
            risk_table.style = 'Table Grid'
            risk_table.cell(0, 0).text = "Risk ID"
            risk_table.cell(0, 1).text = "Description"
            risk_table.cell(0, 2).text = "Impact"
            risk_table.cell(0, 3).text = "Probability"
            risk_table.cell(0, 4).text = "Mitigation"
            risk_table.cell(0, 5).text = "Owner"
            
            for i, risk in enumerate(migration_plan.risks, 1):
                risk_table.cell(i, 0).text = risk.risk_id
                risk_table.cell(i, 1).text = risk.description[:100] + "..." if len(risk.description) > 100 else risk.description
                risk_table.cell(i, 2).text = risk.impact
                risk_table.cell(i, 3).text = risk.probability
                risk_table.cell(i, 4).text = risk.mitigation[:100] + "..." if len(risk.mitigation) > 100 else risk.mitigation
                risk_table.cell(i, 5).text = risk.owner
        
        # Assumptions and constraints
        doc.add_heading('6.3 Assumptions', level=2)
        for assumption in migration_plan.assumptions[:10]:  # Limit to first 10
            doc.add_paragraph(f"• {assumption}")
        
        doc.add_heading('6.4 Constraints', level=2)
        for constraint in migration_plan.constraints[:10]:  # Limit to first 10
            doc.add_paragraph(f"• {constraint}")
    
    def _add_cost_analysis(self, doc, migration_plan):
        """Add cost analysis section."""
        doc.add_heading('7. Cost Analysis', level=1)
        
        # Cost overview
        doc.add_heading('7.1 Cost Overview', level=2)
        cost_para = doc.add_paragraph()
        cost_para.add_run(f"Total migration investment: ${migration_plan.total_investment:,.2f}\n")
        cost_para.add_run(f"Expected annual savings: ${migration_plan.expected_savings:,.2f}\n")
        cost_para.add_run(f"Return on investment: {migration_plan.total_investment / migration_plan.expected_savings * 12:.1f} months")
        
        # Cost breakdown
        doc.add_heading('7.2 Cost Breakdown', level=2)
        if migration_plan.cost_estimates:
            cost_table = doc.add_table(rows=len(migration_plan.cost_estimates) + 1, cols=5)
            cost_table.style = 'Table Grid'
            cost_table.cell(0, 0).text = "Category"
            cost_table.cell(0, 1).text = "Current Monthly"
            cost_table.cell(0, 2).text = "Azure Monthly"
            cost_table.cell(0, 3).text = "Migration Cost"
            cost_table.cell(0, 4).text = "Annual Savings"
            
            for i, cost in enumerate(migration_plan.cost_estimates, 1):
                cost_table.cell(i, 0).text = cost.category
                cost_table.cell(i, 1).text = f"${cost.current_monthly_cost:,.2f}"
                cost_table.cell(i, 2).text = f"${cost.azure_monthly_cost:,.2f}"
                cost_table.cell(i, 3).text = f"${cost.one_time_migration_cost:,.2f}"
                cost_table.cell(i, 4).text = f"${cost.annual_savings:,.2f}"
    
    def _add_implementation_plan(self, doc, migration_plan):
        """Add implementation plan section."""
        doc.add_heading('8. Implementation Plan', level=1)
        
        # Resource plan
        doc.add_heading('8.1 Resource Plan', level=2)
        if 'project_team' in migration_plan.resource_plan:
            for role, details in migration_plan.resource_plan['project_team'].items():
                doc.add_heading(f"{role.replace('_', ' ').title()}", level=3)
                role_para = doc.add_paragraph()
                if isinstance(details, dict):
                    for key, value in details.items():
                        if key != 'responsibilities':
                            role_para.add_run(f"{key.replace('_', ' ').title()}: {value}\n")
                    if 'responsibilities' in details:
                        role_para.add_run("Responsibilities:\n")
                        for responsibility in details['responsibilities']:
                            role_para.add_run(f"• {responsibility}\n")
        
        # Training plan
        doc.add_heading('8.2 Training Plan', level=2)
        if migration_plan.training_plan:
            for training, details in migration_plan.training_plan.items():
                doc.add_heading(f"{training.replace('_', ' ').title()}", level=3)
                if isinstance(details, dict):
                    for key, value in details.items():
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        # Communication plan
        doc.add_heading('8.3 Communication Plan', level=2)
        if 'stakeholder_meetings' in migration_plan.communication_plan:
            for meeting_type, details in migration_plan.communication_plan['stakeholder_meetings'].items():
                doc.add_heading(f"{meeting_type.replace('_', ' ').title()}", level=3)
                if isinstance(details, dict):
                    for key, value in details.items():
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    
    def _add_governance_compliance(self, doc, migration_plan):
        """Add governance and compliance section."""
        doc.add_heading('9. Governance and Compliance', level=1)
        
        # Security requirements
        doc.add_heading('9.1 Security Requirements', level=2)
        for requirement in migration_plan.security_requirements[:15]:  # Limit to first 15
            doc.add_paragraph(f"• {requirement}")
        
        # Compliance requirements
        doc.add_heading('9.2 Compliance Requirements', level=2)
        for requirement in migration_plan.compliance_requirements[:10]:  # Limit to first 10
            doc.add_paragraph(f"• {requirement}")
        
        # Governance model
        doc.add_heading('9.3 Governance Model', level=2)
        if 'governance_structure' in migration_plan.governance_model:
            for role, description in migration_plan.governance_model['governance_structure'].items():
                doc.add_paragraph(f"• {role.replace('_', ' ').title()}: {description}")
    
    def _add_success_metrics(self, doc, migration_plan):
        """Add success metrics section."""
        doc.add_heading('10. Success Metrics', level=1)
        
        # KPIs
        doc.add_heading('10.1 Key Performance Indicators', level=2)
        if migration_plan.kpis:
            kpi_table = doc.add_table(rows=len(migration_plan.kpis) + 1, cols=4)
            kpi_table.style = 'Table Grid'
            kpi_table.cell(0, 0).text = "KPI"
            kpi_table.cell(0, 1).text = "Target"
            kpi_table.cell(0, 2).text = "Measurement"
            kpi_table.cell(0, 3).text = "Frequency"
            
            for i, kpi in enumerate(migration_plan.kpis, 1):
                kpi_table.cell(i, 0).text = kpi.get('kpi', '')
                kpi_table.cell(i, 1).text = kpi.get('target', '')
                kpi_table.cell(i, 2).text = kpi.get('measurement', '')
                kpi_table.cell(i, 3).text = kpi.get('frequency', '')
        
        # Success criteria
        doc.add_heading('10.2 Success Criteria', level=2)
        for criteria in migration_plan.success_criteria[:15]:  # Limit to first 15
            doc.add_paragraph(f"• {criteria}")
    
    def _add_appendices(self, doc, migration_plan):
        """Add appendices section."""
        doc.add_heading('11. Appendices', level=1)
        
        # Technical specifications
        doc.add_heading('11.1 Technical Specifications', level=2)
        if migration_plan.technical_specifications:
            for category, specs in migration_plan.technical_specifications.items():
                doc.add_heading(f"{category.replace('_', ' ').title()}", level=3)
                if isinstance(specs, dict):
                    for key, value in specs.items():
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
                else:
                    doc.add_paragraph(str(specs))
        
        # Vendor requirements
        doc.add_heading('11.2 Vendor Requirements', level=2)
        for requirement in migration_plan.vendor_requirements:
            doc.add_paragraph(f"• {requirement}")