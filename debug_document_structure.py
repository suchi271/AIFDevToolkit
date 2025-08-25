from docx import Document

def debug_document_structure():
    doc = Document('output/application_assessment_report.docx')
    
    print("=== Document Structure Analysis ===\n")
    
    # Find section 5.4 and see what comes after it
    found_54 = False
    section_54_content = []
    next_section_found = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name
        
        if '5.4' in text and 'Low Level Design' in text:
            found_54 = True
            print(f"Found 5.4 at paragraph {i}: '{text}' (style: {style})")
            continue
        
        if found_54 and not next_section_found:
            # Check if this is the next section
            if text.startswith('6') and 'Architecture Heatmap' in text:
                next_section_found = True
                print(f"Found next section at paragraph {i}: '{text}' (style: {style})")
                break
            
            # Collect first few lines to see what's happening
            if len(section_54_content) < 10:
                section_54_content.append({
                    'index': i,
                    'text': text[:100] + "..." if len(text) > 100 else text,
                    'style': style
                })
    
    if found_54:
        print(f"\nContent immediately after 5.4 heading:")
        for item in section_54_content:
            print(f"  Para {item['index']} ({item['style']}): {item['text']}")
        
        if not next_section_found:
            print(f"\nWARNING: No '6 Architecture Heatmap' section found!")
            print("Looking for any paragraph starting with '6':")
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().startswith('6'):
                    print(f"  Para {i}: '{p.text.strip()[:50]}...' (style: {p.style.name})")
    else:
        print("5.4 section not found!")

if __name__ == "__main__":
    debug_document_structure()
