from docx import Document

def check_low_level_design_content():
    doc = Document('output/application_assessment_report.docx')
    
    found_heading = False
    content_lines = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Look specifically for 5.4 Low Level Design heading
        if '5.4' in text and 'Low Level Design' in text:
            found_heading = True
            print(f"✓ Found heading: {text}")
            continue
        
        if found_heading:
            # Stop if we hit the next major section (6. Architecture Heatmap)
            if text.startswith('6') and 'Architecture Heatmap' in text:
                break
            
            # Collect content lines
            if text:
                content_lines.append(text)
    
    if found_heading:
        if content_lines:
            print(f"✓ Found {len(content_lines)} lines of content:")
            for i, line in enumerate(content_lines[:15]):  # Show first 15 lines
                print(f"  {i+1}: {line[:100]}")
            if len(content_lines) > 15:
                print(f"  ... and {len(content_lines) - 15} more lines")
        else:
            print("✗ Heading found but NO CONTENT under it!")
    else:
        print("✗ 5.4 Low Level Design heading NOT FOUND!")
        print("\nLooking for any headings with 'Low Level Design':")
        for p in doc.paragraphs:
            if 'Low Level Design' in p.text:
                print(f"  Found: {p.text}")

if __name__ == "__main__":
    check_low_level_design_content()
